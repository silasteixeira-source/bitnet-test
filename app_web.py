import streamlit as st
import os
import base64
import re
import json
from datetime import datetime
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import gspread
from google.oauth2.service_account import Credentials as ServiceAccountCredentials
import pandas as pd
import io
from docx import Document
from email.message import EmailMessage
import tempfile
import shutil
import pypdf
from pdf2docx import Converter
try:
    from docx2pdf import convert as docx_to_pdf_convert
except ImportError:
    docx_to_pdf_convert = None
from PIL import Image, ImageOps, ImageDraw, ImageFont
import textwrap

# ==========================================
# CONFIGURAÇÃO GLOBAL DA PÁGINA
# ==========================================
st.set_page_config(page_title="Ferramenta EACE - NOC", page_icon="🪶", layout="wide")

SCOPES_GMAIL = ['https://www.googleapis.com/auth/gmail.readonly']

# ==========================================
# FUNÇÕES GLOBAIS DE AUTENTICAÇÃO
# ==========================================
@st.cache_resource
def authenticate_gmail():
    """Autentica o Google usando arquivos locais ou o Cofre da Nuvem (Secrets)."""
    base_path = os.getcwd()
    token_path = os.path.join(base_path, "token_leitura.json")
    secret_path = os.path.join(base_path, "AuthNOC.json")
    creds = None
    
    # MÁGICA DA NUVEM: Se o arquivo não existir fisicamente, tenta criar a partir dos Secrets
    if not os.path.exists(token_path) and "google_token" in st.secrets:
        with open(token_path, "w") as f:
            f.write(json.dumps(dict(st.secrets["google_token"])))
            
    if os.path.exists(token_path):
        from google.oauth2.credentials import Credentials
        creds = Credentials.from_authorized_user_file(token_path, SCOPES_GMAIL)
        
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(secret_path):
                st.error("⚠️ Credenciais não encontradas. O site precisa do token no Cofre (Secrets)!")
                return None
            from google_auth_oauthlib.flow import InstalledAppFlow
            flow = InstalledAppFlow.from_client_secrets_file(secret_path, SCOPES_GMAIL)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token: 
            token.write(creds.to_json())
            
    return build('gmail', 'v1', credentials=creds)

# ----------------- AUTENTICAÇÃO PLANILHAS -----------------
@st.cache_resource
def authenticate_google_sheets():
    scopes = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    try:
        if "gcp_service_account" in st.secrets:
            # Puxa credencial direto dos secrets
            creds_info = dict(st.secrets["gcp_service_account"])
            creds = ServiceAccountCredentials.from_service_account_info(creds_info, scopes=scopes)
        else:
            # Tenta arquivo local (ambiente de dev)
            base_path = os.getcwd()
            creds_path = os.path.join(base_path, "credentials.json")
            if not os.path.exists(creds_path):
                st.error("⚠️ Arquivo credentials.json não encontrado para leitura da planilha!")
                return None
            creds = ServiceAccountCredentials.from_service_account_file(creds_path, scopes=scopes)
        return gspread.authorize(creds)
    except Exception as e:
        st.error(f"Erro ao autenticar Planilha: {e}")
        return None

@st.cache_data(ttl=3600) # Cache de 1 hora
def fetch_planilha_eace():
    gspread_client = authenticate_google_sheets()
    if not gspread_client: return pd.DataFrame()
    try:
        SPREADSHEET_ID = "1Onw1vaSO2SIQ_OfAoDPI6ycnXWTAZ2ijhtujAOhI9UM"
        spreadsheet = gspread_client.open_by_key(SPREADSHEET_ID)
        worksheet = spreadsheet.worksheet("EACE")
        raw_data = worksheet.get_all_values()
        if not raw_data: return pd.DataFrame()
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
        # Formatar INEP
        col = "Código INEP" if "Código INEP" in df.columns else df.columns[0]
        df[col] = df[col].astype(str).str.split('.').str[0].str.strip()
        return df
    except Exception as e:
        st.error(f"Erro ao baixar dados da aba EACE: {e}")
        return pd.DataFrame()

def extrair_texto_da_mensagem(payload):
    """Extrai e limpa o texto do e-mail (suporta Plain e HTML)."""
    try:
        def buscar_partes(p):
            mime = p.get("mimeType")
            if mime == "text/plain":
                d = p.get("body", {}).get("data")
                if d: return base64.urlsafe_b64decode(d).decode("utf-8", errors="ignore")
                
            if mime == "text/html" and "parts" not in p:
                d = p.get("body", {}).get("data")
                if d:
                    html = base64.urlsafe_b64decode(d).decode("utf-8", errors="ignore")
                    html = re.sub(r'<style.*?>.*?</style>', '', html, flags=re.IGNORECASE|re.DOTALL)
                    html = re.sub(r'<script.*?>.*?</script>', '', html, flags=re.IGNORECASE|re.DOTALL)
                    html = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
                    html = re.sub(r'</p>', '\n\n', html, flags=re.IGNORECASE)
                    texto = re.sub(r'<[^>]+>', ' ', html)
                    texto = texto.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"')
                    texto = re.sub(r'[ \t]+', ' ', texto)
                    texto = re.sub(r'\n\s*\n', '\n\n', texto)
                    return texto.strip()
                    
            if "parts" in p:
                for part in p["parts"]:
                    if part.get("mimeType") == "text/plain":
                        res = buscar_partes(part)
                        if res: return res
                for part in p["parts"]:
                    res = buscar_partes(part)
                    if res: return res
            return None

        texto_final = buscar_partes(payload)
        if texto_final: return texto_final
    except Exception as e:
        return f"[Erro ao extrair corpo da mensagem: {e}]"
    return "Conteúdo da mensagem não suportado ou vazio."

def extrair_arquivos_da_mensagem(service, msg_id, payload):
    """Varre o payload buscando anexos de imagem e PDF e faz o download via API."""
    arquivos = []
    
    def buscar_anexos(p):
        mime = p.get("mimeType", "")
        
        # Se for imagem ou PDF, e tiver um ID de anexo (inline ou real)
        if mime.startswith("image/") or mime == "application/pdf":
            body = p.get("body", {})
            attachment_id = body.get("attachmentId")
            if attachment_id:
                try:
                    att = service.users().messages().attachments().get(
                        userId='me', messageId=msg_id, id=attachment_id).execute()
                    data = att.get("data")
                    if data:
                        file_bytes = base64.urlsafe_b64decode(data)
                        filename = p.get("filename", "")
                        if not filename: 
                            if mime == "application/pdf":
                                filename = "anexo.pdf"
                            else:
                                filename = "imagem_inline.png"
                        arquivos.append({"filename": filename, "bytes": file_bytes, "mimeType": mime})
                except Exception as e:
                    pass
        
        # Chamada recursiva para parts
        if "parts" in p:
            for part in p["parts"]:
                buscar_anexos(part)
                
    try:
        buscar_anexos(payload)
    except Exception as e:
        pass
        
    return arquivos


# ==========================================
# MÓDULOS (TELAS)
# ==========================================

def modulo_robo_gmail():
    st.title("🤖 Robô Leitor do Gmail")
    st.markdown("Consulte o histórico de respostas na caixa de entrada.")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        termo = st.text_input("Termo de Busca (Ex: INEP 22047140):", placeholder="Ex: 22102841")
    with col2:
        st.write("")
        st.write("")
        btn_buscar = st.button("🔍 BUSCAR E-MAILS", use_container_width=True, type="primary")

    st.divider()

    if btn_buscar:
        if not termo:
            st.warning("Por favor, digite um termo para buscar.")
        else:
            with st.spinner("Conectando ao Gmail e varrendo a caixa de entrada..."):
                service = authenticate_gmail()
                if service:
                    try:
                        query = f'subject:("{termo}") OR "{termo}"'
                        results = service.users().messages().list(userId='me', q=query, maxResults=50).execute()
                        messages = results.get('messages', [])
                        
                        if not messages:
                            st.info("Nenhum e-mail encontrado para o termo pesquisado.")
                        else:
                            st.success(f"Foram encontrados {len(messages)} e-mails referentes à busca!")
                            
                            for msg in messages:
                                msg_full = service.users().messages().get(userId='me', id=msg['id'], format='full').execute()
                                
                                headers = msg_full.get("payload", {}).get("headers", [])
                                assunto = next((h['value'] for h in headers if h['name'].lower() == 'subject'), "Sem Assunto")
                                remetente = next((h['value'] for h in headers if h['name'].lower() == 'from'), "Desconhecido")
                                data_str = next((h['value'] for h in headers if h['name'].lower() == 'date'), "")
                                
                                try:
                                    import email.utils
                                    from zoneinfo import ZoneInfo
                                    parsed_date = email.utils.parsedate_to_datetime(data_str)
                                    # Converte a data do e-mail para o fuso horário do Brasil
                                    parsed_date = parsed_date.astimezone(ZoneInfo('America/Sao_Paulo'))
                                    data_formatada = parsed_date.strftime("%d/%m/%Y %H:%M")
                                except Exception as e:
                                    data_formatada = data_str
                                    
                                titulo_expander = f"📅 {data_formatada} | 👤 {remetente} | 📌 {assunto}"
                                
                                with st.expander(titulo_expander):
                                    st.markdown("#### Corpo da Mensagem:")
                                    payload = msg_full.get("payload", {})
                                    texto_corpo = extrair_texto_da_mensagem(payload)
                                    st.text(texto_corpo)
                                    
                                    # Extrair e exibir anexos
                                    arquivos = extrair_arquivos_da_mensagem(service, msg['id'], payload)
                                    if arquivos:
                                        st.markdown("---")
                                        st.markdown("#### 📁 Anexos Deste E-mail:")
                                        
                                        # Separa imagens de PDFs para exibir de forma bonita
                                        imagens = [a for a in arquivos if a["mimeType"].startswith("image/")]
                                        pdfs = [a for a in arquivos if a["mimeType"] == "application/pdf"]
                                        
                                        if pdfs:
                                            st.write("**📄 Documentos (PDF):**")
                                            for p in pdfs:
                                                st.download_button(
                                                    label=f"⬇️ Baixar {p['filename']}",
                                                    data=p['bytes'],
                                                    file_name=p['filename'],
                                                    mime="application/pdf",
                                                    key=f"dl_{msg['id']}_{p['filename']}"
                                                )
                                                
                                        if imagens:
                                            st.write("**🖼️ Imagens / Prints:**")
                                            cols = st.columns(min(len(imagens), 3))
                                            for i, img in enumerate(imagens):
                                                with cols[i % len(cols)]:
                                                    try:
                                                        st.image(img["bytes"], caption=img["filename"], use_container_width=True)
                                                    except Exception as e:
                                                        st.error(f"Erro ao exibir {img['filename']}: {e}")
                                                    
                                    
                    except Exception as e:
                        st.error(f"Ocorreu um erro durante a busca: {e}")

def modulo_consulta_eace():
    st.title("📊 Consulta Rápida — EACE")
    st.markdown("Busque os dados da escola em tempo real e verifique a velocidade ofertada.")

    with st.spinner("Sincronizando com a nuvem..."):
        df = fetch_planilha_eace()
        
    if df.empty:
        st.warning("A base de dados não foi carregada. Verifique suas credenciais.")
        return

    # Busca
    col_busca, col_botao = st.columns([3, 1])
    with col_busca:
        termo = st.text_input("Código INEP ou Nome da Escola:")
    with col_botao:
        st.write("")
        st.write("")
        buscar = st.button("🔍 BUSCAR DADOS", use_container_width=True, type="primary")
        
    st.divider()

    if buscar and termo:
        st.session_state['termo_busca'] = termo

    termo_ativo = st.session_state.get('termo_busca', '')

    if termo_ativo:
        col_inep = "Código INEP" if "Código INEP" in df.columns else df.columns[0]
        
        # Filtro
        if termo_ativo.isdigit() and len(termo_ativo) >= 8:
            res = df[df[col_inep] == termo_ativo.strip()]
        else:
            col_nome = "Nome da Escola" if "Nome da Escola" in df.columns else df.columns[1]
            res = df[df[col_nome].astype(str).str.contains(termo_ativo, case=False, na=False)]
            
        if res.empty:
            st.error(f"Nenhuma escola localizada para '{termo_ativo}'.")
        else:
            # Se encontrar mais de uma, mostrar caixa de seleção
            if len(res) > 1:
                st.info(f"Encontramos {len(res)} escolas com esse termo. Selecione uma:")
                nomes = res[col_nome].tolist()
                escolhida = st.selectbox("Selecione a Escola", nomes)
                escola_data = res[res[col_nome] == escolhida].iloc[0]
            else:
                escola_data = res.iloc[0]
                
            # Exibir Painel da Escola
            st.subheader(f"🏫 {escola_data.get('Nome da Escola', 'Escola')}")
            
            c1, c2 = st.columns(2)
            
            with c1:
                st.markdown("### DADOS DA ESCOLA")
                st.write(f"**INEP:** {escola_data.get('Código INEP', '—')}")
                st.write(f"**Endereço:** {escola_data.get('Endereço', '—')}")
                st.write(f"**Kit Wi-Fi:** {escola_data.get('Kit Wi-Fi', '—')}")
                st.write(f"**Status:** {escola_data.get('Status', '—')}")
                
                # Extrair valores de velocidade para o verificador
                dl_min = str(escola_data.get('Velocidade DL Mínima (Mbps)', '0'))
                dl_ofe = str(escola_data.get('Velocidade DL Ofertada (Mbps)', '0'))
                
                try:
                    dl_min_val = float(''.join(c for c in dl_min if c.isdigit() or c=='.') or 0)
                    dl_ofe_val = float(''.join(c for c in dl_ofe if c.isdigit() or c=='.') or 0)
                except:
                    dl_min_val = 0
                    dl_ofe_val = 0
                    
                st.write(f"**DL Mínima (Requerida):** {dl_min_val} Mbps")
                st.write(f"**DL Ofertada (Contratada):** {dl_ofe_val} Mbps")

            with c2:
                st.markdown("### ⚡ VERIFICADOR DE VELOCIDADE")
                speed_input = st.number_input("Resultado do Speedtest (Mbps):", min_value=0.0, step=1.0)
                if st.button("CALCULAR APROVAÇÃO", use_container_width=True):
                    if dl_ofe_val <= 0:
                        st.warning("Velocidade Ofertada inválida na planilha. Não é possível calcular.")
                    elif speed_input <= 0:
                        st.warning("Insira um valor de speedtest maior que zero.")
                    else:
                        threshold = dl_ofe_val * 0.8
                        diferenca = speed_input - threshold
                        pct_entrega = (speed_input / dl_ofe_val) * 100
                        
                        st.markdown("---")
                        if speed_input >= threshold:
                            st.success(f"### ✅ APROVADO\nEntregou **{pct_entrega:.1f}%** do contratado.")
                            st.info(f"O Mínimo aceitável era {threshold} Mbps (80% de {dl_ofe_val} Mbps). Passou com margem de +{diferenca:.1f} Mbps.")
                        else:
                            st.error(f"### ❌ REPROVADO\nEntregou apenas **{pct_entrega:.1f}%** do contratado.")
                            st.warning(f"O Mínimo aceitável era {threshold} Mbps. Faltou {abs(diferenca):.1f} Mbps para aprovar.")


def preparar_documento(modelo):
    nome_arq = "contrato modeloBIT.docx" if modelo == "BITNET" else "contrato modeloST1.docx"
    caminho_modelo = os.path.join(os.getcwd(), nome_arq)
    if not os.path.exists(caminho_modelo):
        raise FileNotFoundError(f"O modelo '{nome_arq}' não foi encontrado.")
    
    doc = Document(caminho_modelo)
    dados = st.session_state['contrato_dados']
    
    # Substituir tags textuais
    for tag, valor in dados.items():
        marcador = f"{{{{{tag}}}}}"
        for p in doc.paragraphs:
            if marcador in p.text: p.text = p.text.replace(marcador, str(valor))
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if marcador in c.text: c.text = c.text.replace(marcador, str(valor))
                    
    # Preencher tabela de escolas
    target = None
    for i, p in enumerate(doc.paragraphs):
        if "Endereço de instalação do serviço" in p.text:
            if i < len(doc.tables): target = doc.tables[i]
    if not target:
        for t in doc.tables:
            if "inep" in t.rows[0].cells[0].text.lower(): target = t
            
    if target:
        while len(target.rows) > 1:
            target._tbl.remove(target.rows[1]._tr)
            
        total = 0.0
        for esc in st.session_state['ineps_lote']:
            row = target.add_row().cells
            celulas = [esc["inep"], esc["endereco"], esc["latlong"], esc["megas"], f"R$ {esc['mensalidade']:.2f}", esc["data_inst"]]
            for i in range(min(len(row), len(celulas))):
                row[i].text = str(celulas[i])
            total += float(esc['mensalidade'])
            
        # Substituir Total
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if "Total:" in c.text: c.text = f"Total: MENSAL R$ {total:.2f}"
        for p in doc.paragraphs:
            if "Total:" in p.text: p.text = p.text.replace("Total:", f"Total: MENSAL R$ {total:.2f}")
            
    return doc

@st.dialog("👁️ Prévia das Tabelas do Contrato", width="large")
def exibir_popup_previa(doc):
    for idx, t in enumerate(doc.tables):
        tabela_dados = []
        for r in t.rows:
            linha = [c.text.replace('\n', ' ') for c in r.cells]
            tabela_dados.append(linha)
        if tabela_dados:
            st.markdown(f"**Tabela {idx+1}:**")
            st.table(tabela_dados)

def modulo_gerador_contrato():
    col_titulo, col_limpar = st.columns([0.8, 0.2])
    with col_titulo:
        st.title("📝 Gerador de Contratos")
        st.markdown("Preencha as abas abaixo, valide os INEPs e gere o contrato instantaneamente.")
    with col_limpar:
        st.write("")
        st.write("")
        if st.button("🗑️ Limpar Tudo", type="secondary", use_container_width=True):
            st.session_state['contrato_dados'] = {}
            st.session_state['ineps_lote'] = []
            st.rerun()

    # Inicializar dados na sessão se não existirem
    if 'contrato_dados' not in st.session_state:
        st.session_state['contrato_dados'] = {}
    if 'ineps_lote' not in st.session_state:
        st.session_state['ineps_lote'] = []

    t_contratada, t_suporte, t_assinaturas, t_ineps = st.tabs(["🏢 Contratada", "📞 Suporte", "✍️ Assinaturas", "📋 INEPs & Gerar"])

    with t_contratada:
        st.subheader("Dados da Contratada")
        c1, c2 = st.columns(2)
        with c1:
            st.session_state['contrato_dados']['razao'] = st.text_input("Razão Social", autocomplete="off")
            st.session_state['contrato_dados']['endereco'] = st.text_input("Endereço", autocomplete="off")
            st.session_state['contrato_dados']['bairro'] = st.text_input("Bairro", autocomplete="off")
            st.session_state['contrato_dados']['cidade'] = st.text_input("Cidade", autocomplete="off")
            st.session_state['contrato_dados']['telefone'] = st.text_input("Telefone", autocomplete="off")
            st.session_state['contrato_dados']['email'] = st.text_input("E-mail", autocomplete="off")
        with c2:
            st.session_state['contrato_dados']['cnpj'] = st.text_input("CNPJ", autocomplete="off")
            st.session_state['contrato_dados']['uf'] = st.text_input("UF", autocomplete="off")
            st.session_state['contrato_dados']['cep'] = st.text_input("CEP", autocomplete="off")
            st.session_state['contrato_dados']['ie'] = st.text_input("IE", autocomplete="off")
            st.session_state['contrato_dados']['site'] = st.text_input("Site", autocomplete="off")

    with t_suporte:
        st.subheader("Contatos de Suporte")
        st.session_state['contrato_dados']['email_sup'] = st.text_input("E-mail Suporte", autocomplete="off")
        st.session_state['contrato_dados']['tel_sup'] = st.text_input("Telefone Suporte", autocomplete="off")
        st.session_state['contrato_dados']['sistema'] = st.text_input("Sistema Chamado", autocomplete="off")

    with t_assinaturas:
        st.subheader("Dados das Assinaturas")
        st.session_state['contrato_dados']['local'] = st.text_input("Local", autocomplete="off")
        st.session_state['contrato_dados']['data'] = st.text_input("Data", value=datetime.now().strftime("%d/%m/%Y"), autocomplete="off")
        st.session_state['contrato_dados']['contratante'] = st.text_input("Contratante", autocomplete="off")
        st.session_state['contrato_dados']['testemunha'] = st.text_input("Testemunha", autocomplete="off")
        st.session_state['contrato_dados']['contratada'] = st.text_input("Contratada (Assinatura)", autocomplete="off")

    with t_ineps:
        st.subheader("Adicionar INEPs")
        st.info("A planilha NOC já está conectada. Basta digitar os INEPs e buscar.")
        
        texto_ineps = st.text_area("Insira os Códigos INEP (um por linha ou separados por vírgula):", height=100)
        if st.button("➕ Adicionar Lote", type="primary"):
            df = fetch_planilha_eace()
            if df.empty:
                st.error("Não foi possível carregar a planilha da nuvem.")
            else:
                lista_limpa = [i.strip() for i in re.split(r'[,\n]+', texto_ineps) if i.strip()]
                col_inep = "Código INEP" if "Código INEP" in df.columns else df.columns[0]
                
                ineps_existentes = [x["inep"] for x in st.session_state['ineps_lote']]
                novos_ineps = []
                
                for inep in lista_limpa:
                    if inep in ineps_existentes:
                        st.warning(f"⚠️ O INEP {inep} já foi adicionado à lista. Ignorado.")
                        continue
                        
                    res = df[df[col_inep] == inep]
                    if not res.empty:
                        escola = res.iloc[0]
                        dl_min = float(''.join(c for c in str(escola.get('Velocidade DL Mínima (Mbps)', '0')) if c.isdigit() or c=='.') or 0)
                        
                        nome_esc = str(escola.get("Nome da Escola", "Escola"))
                        
                        novos_ineps.append({
                            "inep": inep,
                            "nome": nome_esc,
                            "endereco": str(escola.get("Endereço", "")),
                            "latlong": f"{escola.get('Latitude','')}, {escola.get('Longitude','')}",
                            "megas": f"{dl_min} MEGAS",
                            "mensalidade": 0.0,
                            "data_inst": datetime.now().strftime("%d/%m/%Y")
                        })
                    else:
                        st.error(f"❌ INEP {inep} não encontrado na planilha EACE.")
                        
                if novos_ineps:
                    st.session_state['ineps_lote'].extend(novos_ineps)
                    st.success(f"{len(novos_ineps)} INEPs localizados e adicionados com sucesso!")
        
        if st.session_state['ineps_lote']:
            st.divider()
            st.markdown("### ⚡ Preenchimento Rápido (Lote)")
            col_lote1, col_lote2, col_lote3 = st.columns(3)
            with col_lote1:
                val_lote = st.number_input("Mensalidade (R$):", value=0.0, step=10.0, key="val_lote")
            with col_lote2:
                data_lote = st.text_input("Instalação (DDMMAAAA ou DD/MM/AAAA):", value=datetime.now().strftime("%d/%m/%Y"), key="data_lote")
            with col_lote3:
                st.write("")
                st.write("")
                if st.button("Aplicar a Todos", type="secondary", use_container_width=True):
                    for idx, esc in enumerate(st.session_state['ineps_lote']):
                        esc['mensalidade'] = val_lote
                        esc['data_inst'] = data_lote
                        
                        # Sobrescrevendo o estado interno dos widgets do Streamlit
                        st.session_state[f"val_{idx}"] = float(val_lote)
                        st.session_state[f"data_{idx}"] = data_lote
                    st.rerun()

            st.divider()
            st.markdown("### 🏫 Escolas Adicionadas")
            
            # Edição das escolas na tela
            for idx, esc in enumerate(st.session_state['ineps_lote']):
                with st.expander(f"📍 {esc['inep']} - {esc.get('nome', 'Escola')} | {esc['megas']}"):
                    st.write(f"**Endereço:** {esc['endereco']}")
                    c1, c2 = st.columns(2)
                    with c1:
                        novo_val = st.number_input(f"Mensalidade (R$)", value=float(esc['mensalidade']), key=f"val_{idx}")
                        st.session_state['ineps_lote'][idx]['mensalidade'] = novo_val
                    with c2:
                        nova_data = st.text_input(f"Data Instalação", value=esc['data_inst'], key=f"data_{idx}")
                        st.session_state['ineps_lote'][idx]['data_inst'] = nova_data
                    if st.button(f"🗑️ Remover {esc['inep']}", key=f"rem_{idx}"):
                        st.session_state['ineps_lote'].pop(idx)
                        st.rerun()

            st.divider()
            st.markdown("### ⚡ GERAR CONTRATO FINAL")
            modelo = st.radio("Selecione o Modelo Base:", ["BITNET", "ST1"], horizontal=True)
            
            c_prev, c_gerar = st.columns(2)
            
            with c_prev:
                if st.button("👁️ ABRIR POP-UP DE PRÉVIA", use_container_width=True):
                    try:
                        doc = preparar_documento(modelo)
                        exibir_popup_previa(doc)
                    except Exception as e:
                        st.error(f"Erro ao gerar prévia: {e}")
                        
            with c_gerar:
                try:
                    doc = preparar_documento(modelo)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label=f"⬇️ BAIXAR DOCX ({modelo})",
                        data=buffer,
                        file_name=f"Contrato_{modelo}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                        key=f"download_{modelo}"
                    )
                except Exception as e:
                    st.error(f"Erro ao preparar arquivo: {e}")

# ==========================================
# MÓDULO 3: ENVIADOR DE PLEITOS
# ==========================================

def obter_assunto_email(dados):
    return f"NMA SERVICOS DE TELECOMUNICAÇÕES LTDA - Alteração de RI - INEP {dados['inep']}"

def obter_corpo_email_html(dados):
    adicionais = []
    if dados.get('qtd_switch', 0) > 0: adicionais.append(f"{dados['qtd_switch']}x Switch")
    if dados.get('qtd_rack', 0) > 0: adicionais.append(f"{dados['qtd_rack']}x Rack")
    if dados.get('qtd_nobreak', 0) > 0: adicionais.append(f"{dados['qtd_nobreak']}x Nobreak")
    
    html_adicionais = ""
    if adicionais:
        texto_adic = ", ".join(adicionais)
        html_adicionais = f"""
                <div style="background-color: #fffbeb; border-left: 4px solid #f59e0b; padding: 15px; margin: 25px 0; border-radius: 0 4px 4px 0;">
                    <p style="margin: 0; color: #b45309; font-size: 14px;"><strong>Adicionais de Infraestrutura:</strong><br>
                    Além dos Access Points, a vistoria técnica identificou a necessidade de incluir os seguintes equipamentos para viabilizar a instalação: <strong>{texto_adic}</strong>.</p>
                </div>
        """

    cor_pleito = '#059669' if dados['tipo'] == 'UPGRADE' else '#dc2626'
    return f"""
    <html>
    <body style="font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background-color: #f9fafb; margin: 0; padding: 20px;">
        <div style="max-width: 800px; margin: 0 auto; background-color: #ffffff; border: 1px solid #e5e7eb; border-radius: 8px; overflow: hidden; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);">
            
            <div style="background-color: #1e293b; padding: 20px; text-align: center;">
                <h2 style="color: #ffffff; margin: 0; font-size: 20px; font-weight: 600;">SOLICITAÇÃO DE ALTERAÇÃO DE PROJETO</h2>
            </div>
            
            <div style="padding: 30px;">
                <p style="color: #334155; font-size: 15px; margin-top: 0;">Olá equipe EACE,</p>
                <p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Submetemos para análise e aprovação técnica a solicitação de <strong style="color: {cor_pleito};">{dados['tipo']}</strong> de Access Points para a unidade escolar descrita abaixo.
                </p>

                <div style="background-color: #f8fafc; border-left: 4px solid #3b82f6; padding: 15px; margin: 25px 0; border-radius: 0 4px 4px 0;">
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Escola:</strong> {dados['escola']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Código INEP:</strong> {dados['inep']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Localidade:</strong> {dados['municipio']} - {dados['uf']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Endereço:</strong> {dados['endereco']}</p>
                </div>

                <div style="overflow-x: auto;">
                    <table style="width: 100%; border-collapse: collapse; font-size: 12px; text-align: center; margin-bottom: 25px; border: 1px solid #cbd5e1;">
                        <thead>
                            <tr style="background-color: #f1f5f9; color: #334155;">
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">FASE</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">INEP</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">ESCOLA</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">ESTADO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">CIDADE</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">PREVISTO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">PLEITO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">SUGERIDO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">INFANTIL</th>
                            </tr>
                        </thead>
                        <tbody>
                            <tr style="color: #475569;">
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">5</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['inep']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['escola']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['uf']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['municipio']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold;">{dados['aps_atuais']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold; color: {cor_pleito};">{dados['tipo']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold;">{dados['novos_aps']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['escola_infantil']}</td>
                            </tr>
                        </tbody>
                    </table>
                </div>

                <p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Após vistoria técnica e validação da infraestrutura local, constatamos a necessidade de adequação. O projeto original prevê <strong>{dados['aps_atuais']} AP(s)</strong>, contudo, para garantir o correto funcionamento e cobertura dos ambientes, faz-se necessária a instalação de <strong>{dados['novos_aps']} AP(s)</strong>.
                </p>
                {html_adicionais}
                <p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Ficamos no aguardo da aprovação técnica formal para prosseguirmos com a execução do cronograma.
                </p>
                
                <div style="margin-top: 40px; border-top: 1px solid #e2e8f0; padding-top: 20px;">
                    <p style="margin: 0; color: #64748b; font-size: 14px;">Atenciosamente,</p>
                    <p style="margin: 5px 0 0 0; color: #1e293b; font-weight: bold; font-size: 15px;">Equipe de Projetos</p>
                    <p style="margin: 2px 0 0 0; color: #64748b; font-size: 13px;">NMA SERVIÇOS DE TELECOMUNICAÇÕES LTDA</p>
                </div>
            </div>
        </div>
    </body>
    </html>
    """

def gerar_conteudo_docx_pleito(doc, dados):
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    doc.add_paragraph("Prezados,")
    doc.add_paragraph()
    doc.add_paragraph(f"Gostaria de solicitar a autorização da EACE para realização de {dados['tipo'].lower()} na quantidade de Access Points da escola abaixo:")
    doc.add_paragraph(f"• Escola: {dados['escola']}", style='List Bullet')
    doc.add_paragraph(f"• Código INEP / Identificação: {dados['inep']}", style='List Bullet')
    doc.add_paragraph(f"• Município: {dados['municipio']} - {dados['uf']}", style='List Bullet')
    doc.add_paragraph(f"• Endereço: {dados['endereco']}", style='List Bullet')
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=12)
    table.style = 'Table Grid'
    headers = ["FASE", "INEP", "ESCOLA", "ESTADO", "CIDADE", "KIT PREVISTO", "PLEITO", "KIT SUGERIDO", "LATITUDE", "LONGITUDE", "ESCOLA INFANTIL", "JUSTIFICATIVA"]
    for i, h in enumerate(headers): 
        table.cell(0, i).text = h
        
    cells = table.rows[1].cells
    vals = ["5", dados['inep'], dados['escola'], dados['uf'], dados['municipio'], dados['aps_atuais'], dados['tipo'], dados['novos_aps'], dados.get('latitude','-'), dados.get('longitude','-'), dados['escola_infantil'], "Adequação à necessidade da escola"]
    for i, v in enumerate(vals): 
        cells[i].text = str(v)
        
    doc.add_paragraph()
    doc.add_paragraph(f"Conforme consta na lista atual, a escola possui previsão de {dados['aps_atuais']} APs. Porém, após validação da necessidade do ambiente, identificamos que será necessário {dados['tipo'].lower()} para {dados['novos_aps']} APs, a fim de atender de forma adequada a cobertura e o funcionamento da unidade.")
    doc.add_paragraph()
    doc.add_paragraph(f"Dessa forma, solicitamos a autorização da EACE para atualização da quantidade de equipamentos desta escola, passando de {dados['aps_atuais']} AP para {dados['novos_aps']} APs.")
    doc.add_paragraph()
    doc.add_paragraph("Ficamos no aguardo da aprovação para seguir com o atendimento.")
    doc.add_paragraph()
    doc.add_paragraph("Atenciosamente,")

@st.dialog("👁️ Prévia do E-mail", width="large")
def exibir_popup_email(html_body):
    st.components.v1.html(html_body, height=600, scrolling=True)

def verificar_duplicidade_gmail(inep):
    """Retorna True se um e-mail com este INEP já foi enviado pela conta mestre e barra o processo."""
    try:
        gmail_service = authenticate_gmail()
        if not gmail_service:
            return False # Se não tem token de leitura, ignora a trava
            
        inep_str = str(inep).strip()
        query = f"subject:{inep_str} in:sent"
        
        results = gmail_service.users().messages().list(userId='me', q=query).execute()
        messages = results.get('messages', [])
        
        if messages:
            msg_data = gmail_service.users().messages().get(userId='me', id=messages[0]['id'], format='metadata', metadataHeaders=['Date', 'Subject']).execute()
            headers = msg_data['payload']['headers']
            assunto = next((h['value'] for h in headers if h['name'] == 'Subject'), "Desconhecido")
            data_envio = next((h['value'] for h in headers if h['name'] == 'Date'), "Desconhecida")
            
            st.error(f"🚨 **PLEITO JÁ ENVIADO ANTERIORMENTE!**\n\nA automação foi BLOQUEADA pois o sistema mestre detectou que já existe um pleito enviado para este INEP ({inep_str}).\n\n📌 **Assunto:** {assunto}\n📅 **Data:** {data_envio}")
            return True
            
        return False
    except Exception as e:
        st.warning(f"Aviso: Não foi possível checar a duplicidade de e-mail: {e}")
        return False

def check_oauth_callback():
    from google_auth_oauthlib.flow import Flow
    import os
    import json
    
    web_client_file = "AuthWebLocal.json"
    
    # MÁGICA DA NUVEM: Cria o arquivo a partir do cofre se estiver na web
    if not os.path.exists(web_client_file) and "google_auth_web" in st.secrets:
        auth_web = dict(st.secrets["google_auth_web"])
        if "web" in auth_web:
            auth_web["web"] = dict(auth_web["web"])
        with open(web_client_file, "w") as f:
            f.write(json.dumps(auth_web))
            
    if "code" in st.query_params and "gmail_creds" not in st.session_state:
        code = st.query_params["code"]
        state = st.query_params.get("state", None)
        try:
            if not os.path.exists(web_client_file):
                st.error("Arquivo AuthWebLocal.json não encontrado. Configure-o primeiro.")
                return
            
            is_cloud = "google_auth_web" in st.secrets
            redirect_uri = "https://bitnet-colaboradores.streamlit.app" if is_cloud else "http://localhost:8501" 
            flow = Flow.from_client_secrets_file(
                web_client_file,
                scopes=['https://www.googleapis.com/auth/gmail.send'],
                redirect_uri=redirect_uri
            )
            
            # Restaurar o code_verifier (PKCE) da memoria em disco
            if os.path.exists("oauth_state.json"):
                with open("oauth_state.json", "r") as f:
                    saved = json.load(f)
                    if saved.get("state") == state and saved.get("code_verifier"):
                        flow.code_verifier = saved["code_verifier"]
            
            flow.fetch_token(code=code)
            st.session_state["gmail_creds"] = flow.credentials
            
            st.query_params.clear()
            st.rerun()
        except Exception as e:
            st.error(f"Erro na autenticação: {e}")

def get_oauth_login_url():
    from google_auth_oauthlib.flow import Flow
    import os
    import json
    web_client_file = "AuthWebLocal.json"
    
    # MÁGICA DA NUVEM: Cria o arquivo a partir do cofre se estiver na web
    if not os.path.exists(web_client_file) and "google_auth_web" in st.secrets:
        auth_web = dict(st.secrets["google_auth_web"])
        if "web" in auth_web:
            auth_web["web"] = dict(auth_web["web"])
        with open(web_client_file, "w") as f:
            f.write(json.dumps(auth_web))
            
    if not os.path.exists(web_client_file):
        return None
    try:
        is_cloud = "google_auth_web" in st.secrets
        redirect_uri = "https://bitnet-colaboradores.streamlit.app" if is_cloud else "http://localhost:8501"
        
        flow = Flow.from_client_secrets_file(
            web_client_file,
            scopes=['https://www.googleapis.com/auth/gmail.send'],
            redirect_uri=redirect_uri
        )
        auth_url, state = flow.authorization_url(prompt='consent', access_type='offline')
        
        # Salvar o code_verifier no disco pois a sessão é perdida no redirecionamento
        with open("oauth_state.json", "w") as f:
            json.dump({
                "state": state,
                "code_verifier": getattr(flow, 'code_verifier', None)
            }, f)
            
        return auth_url
    except Exception:
        return None

def modulo_enviador_pleito():
    st.error("🛑 **ACESSO RESTRITO**")
    st.warning("Este módulo é de uso restrito e autorizado **APENAS** para pessoas expressamente designadas para o envio oficial de e-mails do NOC. Se você não possui autorização, não utilize esta funcionalidade.")
    st.markdown("---")
    check_oauth_callback()
    
    st.markdown("### 📄 AUTOMAÇÃO EACE (E-MAIL/DOCX)")
    st.write("")
    
    # AUTENTICAÇÃO OBRIGATÓRIA (OAuth)
    if "gmail_creds" not in st.session_state or not st.session_state["gmail_creds"].valid:
        st.info("🔒 Para continuar, você precisa fazer login com sua conta do Google para enviar os e-mails em seu próprio nome.")
        login_url = get_oauth_login_url()
        if login_url:
            st.markdown(f'<a href="{login_url}" target="_self"><button style="padding:10px 20px; font-weight:bold; background-color:#4285F4; color:white; border:none; border-radius:5px; cursor:pointer;">🔑 Fazer Login com o Google</button></a>', unsafe_allow_html=True)
        else:
            st.error("⚠️ Erro Crítico: O arquivo 'AuthWebLocal.json' (Web Application Client ID) não foi encontrado na pasta. Siga as instruções do GCP para criá-lo.")
        return
        
    st.success("✅ Logado com sucesso! Seus e-mails serão enviados em seu nome.")
    st.write("")
    
    # SELETOR DE MODELOS
    modelo_pleito = st.radio("Selecione o Modelo de Contrato/Pleito:", ["ST1", "BITNET"], horizontal=True)
    if modelo_pleito == "BITNET":
        st.info("🚧 **Em Desenvolvimento** - O modelo BITNET ainda está em processo de criação de layout de e-mail e mensagens.")
        return
    st.divider()
    
    # GERENCIADOR DE EMAILS PADRÃO
    import json
    config_file = "config_pleito.json"
    
    def carregar_emails_salvos():
        if os.path.exists(config_file):
            try:
                with open(config_file, "r") as f:
                    return json.load(f)
            except:
                pass
        return {"to": "planejamento@eace.org.br", "cc": ""}
        
    emails_padrao = carregar_emails_salvos()
    
    # DESTINATÁRIOS
    c_to, c_cc = st.columns([0.2, 0.8])
    with c_to:
        st.write("Para:")
    with c_cc:
        email_to = st.text_input("Para", value=emails_padrao["to"], label_visibility="collapsed", autocomplete="off")
        
    c_to2, c_cc2 = st.columns([0.2, 0.8])
    with c_to2:
        st.write("Cc (separar por vírgula):")
    with c_cc2:
        email_cc = st.text_input("Cc", value=emails_padrao["cc"], label_visibility="collapsed", autocomplete="off")
        
    c_space, c_btn_save = st.columns([0.8, 0.2])
    with c_btn_save:
        if st.button("💾 Salvar Contatos", use_container_width=True):
            with open(config_file, "w") as f:
                json.dump({"to": email_to, "cc": email_cc}, f)
            st.success("E-mails salvos!")
        
    st.write("")
    
    # INEP e APs
    inep_busca = st.text_input("Código INEP da Escola:", autocomplete="off")
    novos_aps_str = st.text_input("Nova Quantidade de APs (Pleito):", autocomplete="off")
    escola_inf = st.checkbox("Escola Infantil?")
    
    st.write("")
    
    # ADICIONAIS
    with st.container(border=True):
        st.markdown("**Adicionais de Infraestrutura (Qtd)**")
        c_sw, c_rk, c_nb = st.columns(3)
        with c_sw:
            qtd_switch = st.number_input("Switch:", min_value=0, max_value=50, value=0)
        with c_rk:
            qtd_rack = st.number_input("Rack:", min_value=0, max_value=50, value=0)
        with c_nb:
            qtd_nobreak = st.number_input("Nobreak:", min_value=0, max_value=50, value=0)
            
    st.write("")
    
    # ANEXOS
    anexos_upload = st.file_uploader("Anexo de Evidência Local (Opcional):", accept_multiple_files=True)
    
    st.write("")
    
    def buscar_dados_locais():
        if not inep_busca or not novos_aps_str:
            st.warning("Por favor, preencha o Código INEP e a Nova Quantidade de APs.")
            return None
            
        inep = inep_busca.split('.')[0].strip()
        
        # Tenta ler da memoria
        df = None
        if 'df_escolas' in st.session_state and not st.session_state['df_escolas'].empty:
            df = st.session_state['df_escolas']
        else:
            # Tenta ler do disco
            if os.path.exists("eace_cache.json"):
                df = pd.read_json("eace_cache.json", dtype=str)
            else:
                st.error("Banco de dados 'eace_cache.json' não encontrado na pasta! Vá na aba Consulta EACE e carregue a planilha para gerar o cache.")
                return None
                
        col_name = "Código INEP" if "Código INEP" in df.columns else df.columns[0]
        df_copy = df.copy()
        df_copy[col_name] = df_copy[col_name].astype(str).str.split('.').str[0].str.strip()
        
        res = df_copy[df_copy[col_name] == inep]
        if res.empty:
            st.error(f"O INEP '{inep}' não foi localizado no banco de dados.")
            return None
            
        d = res.iloc[0]
        
        aps_atuais_str = str(d.get('Kit Wi-Fi', '0')).strip()
        try:
            aps_atuais = int(aps_atuais_str or 0)
            novos_aps = int(novos_aps_str)
            tipo = "UPGRADE" if novos_aps > aps_atuais else "DOWNGRADE"
        except ValueError:
            aps_atuais = aps_atuais_str
            novos_aps = novos_aps_str
            tipo = "UPGRADE"
            
        escola_infantil = "SIM" if escola_inf else "NÃO"
        
        return {
            "qtd_switch": qtd_switch,
            "qtd_rack": qtd_rack,
            "qtd_nobreak": qtd_nobreak,
            "inep": inep,
            "escola": str(d.get('Nome da Escola', '')).strip(),
            "uf": str(d.get('UF', '')).strip(),
            "municipio": str(d.get('Município', '')).strip(),
            "endereco": str(d.get('Endereço', '')).strip(),
            "latitude": str(d.get('Latitude', '')).strip(),
            "longitude": str(d.get('Longitude', '')).strip(),
            "aps_atuais": aps_atuais,
            "novos_aps": novos_aps,
            "tipo": tipo,
            "escola_infantil": escola_infantil
        }

    # BOTOES
    c_baixar, c_prev = st.columns(2)
    with c_baixar:
        btn_baixar = st.button("📄 GERAR DOCX", use_container_width=True)
    with c_prev:
        btn_prev = st.button("👁️ VISUALIZAR PRÉVIA", use_container_width=True)
        
    btn_enviar = st.button("🚀 ENVIAR POR E-MAIL", use_container_width=True)
    
    if btn_prev:
        dados = buscar_dados_locais()
        if dados:
            html_preview = obter_corpo_email_html(dados)
            exibir_popup_email(html_preview)
            
    if btn_baixar:
        dados = buscar_dados_locais()
        if dados:
            try:
                doc_pleito = Document()
                gerar_conteudo_docx_pleito(doc_pleito, dados)
                buffer = io.BytesIO()
                doc_pleito.save(buffer)
                buffer.seek(0)
                st.session_state['pleito_docx_buffer'] = buffer
                st.session_state['pleito_docx_name'] = f"Pleito_{dados['inep']}.docx"
                st.success("✅ DOCX processado com sucesso!")
            except Exception as e:
                st.error(f"Erro ao processar DOCX: {e}")
                
    # Na Web o Download Button real precisa renderizar
    if 'pleito_docx_buffer' in st.session_state:
        st.download_button(
            label="⬇️ CLIQUE AQUI PARA SALVAR O ARQUIVO NO PC",
            data=st.session_state['pleito_docx_buffer'],
            file_name=st.session_state['pleito_docx_name'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
            
    if btn_enviar:
        dados = buscar_dados_locais()
        if dados:
            if verificar_duplicidade_gmail(dados['inep']):
                return
                
            try:
                from googleapiclient.discovery import build
                # Agora usa a credencial da sessão e não o authenticate_gmail antigo centralizado
                gmail_service = build('gmail', 'v1', credentials=st.session_state["gmail_creds"])
                
                # --- TRAVA DE SEGURANÇA DE DOMÍNIO ---
                profile = gmail_service.users().getProfile(userId='me').execute()
                user_email = profile.get('emailAddress', '').lower()
                modelo = dados.get('modelo', '')
                
                if modelo == "BITNET" and "st1.net.br" in user_email:
                    st.error(f"❌ Operação Cancelada! Você está tentando enviar um pleito do modelo **BITNET**, mas o seu e-mail do Google conectado é `{user_email}` (Domínio ST1).")
                    return
                elif modelo == "ST1" and "bitinternet.com.br" in user_email:
                    st.error(f"❌ Operação Cancelada! Você está tentando enviar um pleito do modelo **ST1**, mas o seu e-mail do Google conectado é `{user_email}` (Domínio BITNET).")
                    return
                # -------------------------------------
                
                msg = EmailMessage()
                msg['To'] = email_to
                if email_cc: msg['Cc'] = email_cc
                msg['Subject'] = obter_assunto_email(dados)
                msg.add_alternative(obter_corpo_email_html(dados), subtype='html')
                
                # Anexos
                if anexos_upload:
                    import mimetypes
                    for arquivo in anexos_upload:
                        file_bytes = arquivo.read()
                        content_type, encoding = mimetypes.guess_type(arquivo.name)
                        if content_type is None: content_type = 'application/octet-stream'
                        main_type, sub_type = content_type.split('/', 1)
                        msg.add_attachment(file_bytes, maintype=main_type, subtype=sub_type, filename=arquivo.name)
                        
                raw_msg = base64.urlsafe_b64encode(msg.as_bytes()).decode()
                gmail_service.users().messages().send(userId='me', body={'raw': raw_msg}).execute()
                
                st.success("✅ E-mail enviado com sucesso pelo seu próprio Gmail!")
                st.balloons()
            except Exception as e:
                st.error(f"Erro ao enviar e-mail: {e}")

import platform
import subprocess

def modulo_ferramentas():
    st.title("🛠️ Ferramentas Extras")
    
    tab1, tab2 = st.tabs(["🔄 Conversor Bidirecional", "📚 Unificador de Evidências"])
    
    # --- TAB 1: Conversor Bidirecional ---
    with tab1:
        st.markdown("### Conversor Bidirecional de Documentos")
        sentido = st.radio("Sentido da Conversão:", ["PDF para Word (.docx)", "Word (.docx) para PDF"])
        
        if "PDF para Word" in sentido:
            files = st.file_uploader("Selecione arquivos PDF", type=["pdf"], accept_multiple_files=True, key="pdf2word")
            if files:
                if st.button("Converter para Word", type="primary"):
                    for f in files:
                        with st.spinner(f"Convertendo {f.name}..."):
                            try:
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_in:
                                    tmp_in.write(f.read())
                                    tmp_in_path = tmp_in.name
                                tmp_out_path = tmp_in_path + ".docx"
                                
                                cv = Converter(tmp_in_path)
                                cv.convert(tmp_out_path)
                                cv.close()
                                
                                with open(tmp_out_path, "rb") as fout:
                                    st.download_button(label=f"⬇️ Baixar {f.name.replace('.pdf', '.docx')}",
                                                       data=fout,
                                                       file_name=f.name.replace('.pdf', '.docx'),
                                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                       key=f"d_{f.name}")
                            except Exception as e:
                                st.error(f"Erro ao converter {f.name}: {e}")
        else:
            st.info("💡 A conversão Word para PDF usa recursos nativos do seu sistema. Funciona no Windows e na Nuvem (Streamlit Cloud).")
            files = st.file_uploader("Selecione arquivos Word", type=["docx"], accept_multiple_files=True, key="word2pdf")
            if files:
                if st.button("Converter para PDF", type="primary"):
                    for f in files:
                        with st.spinner(f"Convertendo {f.name}..."):
                            try:
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
                                    tmp_in.write(f.read())
                                    tmp_in_path = tmp_in.name
                                
                                temp_dir = os.path.dirname(tmp_in_path)
                                tmp_out_path = tmp_in_path.replace(".docx", ".pdf")
                                
                                if platform.system() == "Windows":
                                    if docx_to_pdf_convert is None:
                                        st.error("Biblioteca 'docx2pdf' ou MS Word não estão instalados no seu Windows.")
                                        continue
                                    docx_to_pdf_convert(tmp_in_path, tmp_out_path)
                                else:
                                    # Linux / Nuvem
                                    result = subprocess.run(
                                        ['libreoffice', '--headless', '--convert-to', 'pdf', tmp_in_path, '--outdir', temp_dir],
                                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
                                    )
                                    if result.returncode != 0:
                                        st.error(f"Erro no servidor ao converter {f.name}: {result.stderr}")
                                        continue
                                
                                if os.path.exists(tmp_out_path):
                                    with open(tmp_out_path, "rb") as fout:
                                        st.download_button(label=f"⬇️ Baixar {f.name.replace('.docx', '.pdf')}",
                                                           data=fout,
                                                           file_name=f.name.replace('.docx', '.pdf'),
                                                           mime="application/pdf",
                                                           key=f"d2_{f.name}")
                                else:
                                    st.error("Falha ao gerar o arquivo PDF.")
                                    
                            except Exception as e:
                                st.error(f"Erro ao converter {f.name}: {e}")

    # --- TAB 2: Unificador de Evidências ---
    with tab2:
        st.markdown("### Unificador de Evidências Multiformato")
        st.info("Suporta Imagens (PNG/JPG), Texto (TXT) e PDFs. (Suporte a DOCX omitido).")
        
        nome_final = st.text_input("Nome do arquivo PDF final (sem extensão):", value="Evidencias_Unificadas")
        otimizar = st.checkbox("Otimizar tamanho das imagens", value=True)
        
        files = st.file_uploader("Selecione os arquivos na ordem desejada", type=["png", "jpg", "jpeg", "txt", "pdf"], accept_multiple_files=True)
        
        if files:
            st.write(f"**{len(files)} arquivos selecionados para mesclar.**")
            
            if st.button("⚡ GERAR PDF UNIFICADO ⚡", type="primary", use_container_width=True):
                with st.spinner("Unificando evidências..."):
                    try:
                        temp_dir = tempfile.mkdtemp()
                        pdf_parts = []
                        
                        for idx, f in enumerate(files):
                            ext = os.path.splitext(f.name)[1].lower()
                            temp_pdf_path = os.path.join(temp_dir, f"part_{idx}.pdf")
                            
                            if ext in [".png", ".jpg", ".jpeg"]:
                                img = Image.open(f)
                                img = ImageOps.exif_transpose(img).convert('RGB')
                                if otimizar:
                                    max_size = 1920
                                    if img.width > max_size or img.height > max_size:
                                        img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                                img.save(temp_pdf_path, "PDF", resolution=100.0, quality=80)
                                pdf_parts.append(temp_pdf_path)
                                
                            elif ext == ".txt":
                                text = f.getvalue().decode('utf-8', errors='replace')
                                try:
                                    font = ImageFont.truetype("arial.ttf", 20)
                                except:
                                    font = ImageFont.load_default()
                                img = Image.new('RGB', (1240, 1754), color=(255, 255, 255))
                                draw = ImageDraw.Draw(img)
                                lines = []
                                for p in text.split('\n'): lines.extend(textwrap.wrap(p, width=85))
                                y = 50
                                for line in lines:
                                    draw.text((50, y), line, font=font, fill=(0, 0, 0))
                                    y += 30
                                img.save(temp_pdf_path, "PDF", resolution=100.0)
                                pdf_parts.append(temp_pdf_path)
                                
                            elif ext == ".pdf":
                                with open(temp_pdf_path, "wb") as ftemp:
                                    ftemp.write(f.read())
                                pdf_parts.append(temp_pdf_path)
                                
                        if not pdf_parts:
                            st.error("Nenhuma parte do PDF foi gerada.")
                        else:
                            writer = pypdf.PdfWriter()
                            for part in pdf_parts:
                                reader = pypdf.PdfReader(part)
                                for page in reader.pages:
                                    writer.add_page(page)
                                    
                            out_path = os.path.join(temp_dir, f"{nome_final}.pdf")
                            with open(out_path, "wb") as f_out:
                                writer.write(f_out)
                                
                            with open(out_path, "rb") as final_pdf:
                                st.success("✅ PDF unificado gerado com sucesso!")
                                st.download_button(label=f"⬇️ Baixar {nome_final}.pdf",
                                                   data=final_pdf,
                                                   file_name=f"{nome_final}.pdf",
                                                   mime="application/pdf",
                                                   use_container_width=True)
                    except Exception as e:
                        st.error(f"Erro ao unificar: {e}")

# ==========================================
# CAMADA DE SEGURANÇA E ROTEAMENTO
# ==========================================

if 'autenticado' not in st.session_state:
    st.session_state['autenticado'] = False

if not st.session_state['autenticado']:
    # TELA DE LOGIN
    st.title("🔒 Acesso Restrito - Ferramenta EACE")
    st.markdown("Por favor, insira a senha da equipe para acessar o portal.")
    
    with st.form("login_form"):
        senha = st.text_input("Senha:", type="password")
        btn_login = st.form_submit_button("Entrar")
        
    if btn_login:
        # Puxa a lista de senhas EXCLUSIVAMENTE do Secrets da Nuvem
        senhas_permitidas = st.secrets.get("senhas_portal", [])
        
        # Se for só uma string no Secrets, converte pra lista
        if isinstance(senhas_permitidas, str):
            senhas_permitidas = [senhas_permitidas]
            
        # Pega a senha digitada sem espaços perdidos
        senha_limpa = senha.strip()
            
        if senha_limpa in senhas_permitidas: 
            st.session_state['autenticado'] = True
            st.rerun()
        else:
            st.error("Senha incorreta!")
    st.stop()

# --- MENU LATERAL (SIDEBAR) ---
st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/thumb/c/c3/Python-logo-notext.svg/120px-Python-logo-notext.svg.png", width=60) # Placeholder logo
st.sidebar.title("NOC EACE")
st.sidebar.markdown("---")

menu = st.sidebar.radio(
    "Navegação:",
    ("📊 Consulta EACE", "🤖 Robô Leitor do Gmail", "📝 Gerar Contrato", "✉️ Enviador de Pleitos", "🛠️ Ferramentas")
)

st.sidebar.markdown("---")
if st.sidebar.button("Sair (Logout)"):
    st.session_state['autenticado'] = False
    st.rerun()

# --- ROTEAMENTO ---
if menu == "🤖 Robô Leitor do Gmail":
    modulo_robo_gmail()
elif menu == "📊 Consulta EACE":
    modulo_consulta_eace()
elif menu == "📝 Gerar Contrato":
    modulo_gerador_contrato()
elif menu == "✉️ Enviador de Pleitos":
    modulo_enviador_pleito()
elif menu == "🛠️ Ferramentas":
    modulo_ferramentas()
else:
    st.title(menu)
    st.warning("Este módulo ainda está sendo migrado para a versão Web.")
