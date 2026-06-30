import sys

import base64
import webbrowser
from email.message import EmailMessage
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
import os
import io
import json
import re
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import pandas as pd
from PIL import Image, ImageOps, ImageTk, ImageDraw, ImageFont
import textwrap
import pillow_heif
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from pdf2docx import Converter
import multiprocessing
import traceback
import threading
from datetime import datetime
from docx import Document
import tempfile
import shutil

# Tentar importar bibliotecas extras instaladas
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import win32com.client
except ImportError:
    win32com.client = None

try:
    from docx2pdf import convert as docx_to_pdf_convert
except ImportError:
    docx_to_pdf_convert = None

# Registrar suporte a HEIC para Pillow
pillow_heif.register_heif_opener()

# --- CONFIGURAÇÕES GLOBAIS ---
SPREADSHEET_ID = "1Onw1vaSO2SIQ_OfAoDPI6ycnXWTAZ2ijhtujAOhI9UM"
DASHBOARD_SPREADSHEET_ID = "1hoK7MT4ST6dQe5WsA4rfEG3Tk7NTYAovrL6K9cXlxWY"
CONFIG_FILE = "config_apps.json"
CACHE_FILE = "eace_cache.json"
DASHBOARD_CACHE_FILE = "eace_dashboard_cache.json"

# --- CONFIGURAÇÕES DE PALETAS DE CORES (TEMAS) ---
THEMES = {
    "light": {
        "bg": "#f8f9fa",
        "card_bg": "#ffffff",
        "fg": "#2c3e50",
        "fg_muted": "#7f8c8d",
        "primary": "#007bff",
        "primary_hover": "#0056b3",
        "success": "#28a745",
        "info": "#17a2b8",
        "warning": "#ffc107",
        "danger": "#dc3545",
        "border": "#dee2e6",
        "entry_bg": "#ffffff",
        "entry_fg": "#495057",
        "tree_bg": "#ffffff",
        "tree_fg": "#2c3e50",
        "tree_selected": "#007bff"
    },
    "dark": {
        "bg": "#0f172a",
        "card_bg": "#1e293b",
        "fg": "#f8fafc",
        "fg_muted": "#94a3b8",
        "primary": "#6366f1",
        "primary_hover": "#4f46e5",
        "success": "#10b981",
        "info": "#0ea5e9",
        "warning": "#f59e0b",
        "danger": "#ef4444",
        "border": "#334155",
        "entry_bg": "#0f172a",
        "entry_fg": "#f8fafc",
        "tree_bg": "#1e293b",
        "tree_fg": "#f8fafc",
        "tree_selected": "#6366f1"
    }
}

# --- GERENCIADOR DE TEMAS DINÂMICO ---
class ThemeManager:
    current_theme = "dark"

    @classmethod
    def get_colors(cls):
        return THEMES[cls.current_theme]

    @classmethod
    def toggle_theme(cls, root):
        cls.current_theme = "light" if cls.current_theme == "dark" else "dark"
        save_config("theme", cls.current_theme)
        cls.apply_theme_to_all(root)

    @classmethod
    def apply_theme_to_all(cls, root):
        colors = cls.get_colors()

        # Configurar Estilos do TTK
        style = ttk.Style(root)
        style.theme_use("clam")

        # Configurações globais ttk
        style.configure("TNotebook", background=colors["bg"], borderwidth=0)
        style.configure("TNotebook.Tab", background=colors["card_bg"], foreground=colors["fg"], borderwidth=1, padding=8, font=("Segoe UI", 9, "bold"))
        style.map("TNotebook.Tab", background=[("selected", colors["primary"])], foreground=[("selected", "#ffffff")])

        style.configure("TFrame", background=colors["bg"])
        style.configure("TLabel", background=colors["bg"], foreground=colors["fg"], font=("Segoe UI", 10))
        style.configure("TButton", background=colors["card_bg"], foreground=colors["fg"], bordercolor=colors["border"], font=("Segoe UI", 9, "bold"), padding=6)
        style.map("TButton", background=[("active", colors["primary"])], foreground=[("active", "#ffffff")])

        style.configure("TEntry", fieldbackground=colors["entry_bg"], foreground=colors["entry_fg"], insertcolor=colors["fg"], bordercolor=colors["border"])
        style.configure("TCheckbutton", background=colors["bg"], foreground=colors["fg"])

        style.configure("Treeview", background=colors["tree_bg"], foreground=colors["tree_fg"], fieldbackground=colors["tree_bg"], rowheight=24, font=("Segoe UI", 9))
        style.configure("Treeview.Heading", background=colors["card_bg"], foreground=colors["fg"], font=("Segoe UI", 9, "bold"), borderwidth=1)
        style.map("Treeview", background=[("selected", colors["tree_selected"])], foreground=[("selected", "#ffffff")])

        cls._apply_to_widget_recursive(root, colors)

    @classmethod
    def _apply_to_widget_recursive(cls, widget, colors):
        widget_type = widget.winfo_class()
        
        # Verificar se é ModernButton para redesenho com as novas cores do tema
        if hasattr(widget, "update_colors") and hasattr(widget, "bg_type"):
            try:
                widget.update_colors()
            except Exception:
                pass
            return

        try:
            if widget_type == "Frame" or widget_type == "TFrame":
                widget.configure(bg=colors["bg"])
            elif widget_type == "Label" or widget_type == "TLabel":
                # Respeitar cores de alerta/sucesso/info personalizadas
                current_fg = widget.cget("fg") if hasattr(widget, "cget") else ""
                if current_fg not in ["red", "green", "blue", "orange", "#27ae60", "#e74c3c", "#17a2b8", "#2ecc71"]:
                    widget.configure(bg=colors["bg"], fg=colors["fg"])
                else:
                    widget.configure(bg=colors["bg"])
            elif widget_type == "Button" or widget_type == "TButton":
                current_bg = widget.cget("bg") if hasattr(widget, "cget") else ""
                if current_bg not in ["#6c757d", "#007bff", "#17a2b8", "#28a745", "#e67e22", "#3498db", "#e74c3c", "#27ae60", "#2ecc71", "#dc3545"]:
                    widget.configure(bg=colors["card_bg"], fg=colors["fg"], activebackground=colors["primary"], activeforeground="#ffffff", bd=1, relief="groove")
            elif widget_type == "Entry" or widget_type == "TEntry":
                widget.configure(bg=colors["entry_bg"], fg=colors["entry_fg"], insertbackground=colors["fg"], bd=1, relief="solid")
            elif widget_type == "Text":
                widget.configure(bg=colors["entry_bg"], fg=colors["entry_fg"], insertbackground=colors["fg"], bd=1, relief="solid")
            elif widget_type == "Listbox":
                widget.configure(bg=colors["entry_bg"], fg=colors["entry_fg"], selectbackground=colors["primary"], selectforeground="#ffffff", bd=1, relief="solid")
            elif widget_type == "LabelFrame" or widget_type == "TLabelframe":
                widget.configure(bg=colors["bg"], fg=colors["fg"], bd=1, relief="groove")
            elif widget_type == "Checkbutton" or widget_type == "TCheckbutton":
                widget.configure(bg=colors["bg"], fg=colors["fg"], selectcolor=colors["card_bg"], activebackground=colors["bg"], activeforeground=colors["fg"])
        except Exception:
            pass

        for child in widget.winfo_children():
            cls._apply_to_widget_recursive(child, colors)

# --- BOTÃO MODERNO COM CANTOS ARREDONDADOS E MICRO-ANIMAÇÕES ---
class ModernButton(tk.Canvas):
    def __init__(self, parent, text, command, bg_type="primary", font=("Segoe UI", 10, "bold"), radius=12, height=35, width=120):
        super().__init__(parent, height=height, width=width, bg=parent.cget("bg") if hasattr(parent, "cget") else "#0f172a", highlightthickness=0)
        self.text = text
        self.command = command
        self.bg_type = bg_type
        self.font = font
        self.radius = radius
        self.height = height
        
        self.rect_id = None
        self.text_id = None
        
        self.bg_color = ""
        self.fg_color = ""
        self.hover_color = ""
        
        self.update_colors_from_theme()
        
        self.bind("<Configure>", lambda e: self.draw_button(self.bg_color))
        self.bind("<Enter>", lambda e: self.on_enter())
        self.bind("<Leave>", lambda e: self.on_leave())
        self.bind("<ButtonPress-1>", lambda e: self.on_click())
        self.bind("<ButtonRelease-1>", lambda e: self.on_release())

    def update_colors_from_theme(self):
        colors = ThemeManager.get_colors()
        self.fg_color = "#ffffff"
        
        if self.bg_type == "primary":
            self.bg_color = colors["primary"]
            self.hover_color = colors["primary_hover"]
        elif self.bg_type == "success":
            self.bg_color = colors["success"]
            self.hover_color = self.darken_color(colors["success"], 0.12)
        elif self.bg_type == "info":
            self.bg_color = colors["info"]
            self.hover_color = self.darken_color(colors["info"], 0.12)
        elif self.bg_type == "warning":
            self.bg_color = colors["warning"]
            self.hover_color = self.darken_color(colors["warning"], 0.12)
        elif self.bg_type == "danger":
            self.bg_color = colors["danger"]
            self.hover_color = self.darken_color(colors["danger"], 0.12)
        elif self.bg_type == "muted":
            self.bg_color = "#6c757d"
            self.hover_color = "#5a6268"
        else:
            self.bg_color = self.bg_type
            self.hover_color = self.darken_color(self.bg_type, 0.12)
            
        try:
            self.config(bg=self.master.cget("bg"))
        except:
            pass
            
    def draw_button(self, color):
        self.delete("all")
        w = self.winfo_width()
        h = self.height
        
        if w < 10:
            w = 120
            
        r = self.radius
        if r > h // 2:
            r = h // 2
            
        points = [
            r, 0,
            w - r, 0,
            w, 0,
            w, r,
            w, h - r,
            w, h,
            w - r, h,
            r, h,
            0, h,
            0, h - r,
            0, r,
            0, 0
        ]
        
        self.rect_id = self.create_polygon(points, fill=color, smooth=True, splinesteps=36)
        self.text_id = self.create_text(w // 2, h // 2, text=self.text, fill=self.fg_color, font=self.font, justify="center")

    def on_enter(self):
        self.draw_button(self.hover_color)
        self.config(cursor="hand2")

    def on_leave(self):
        self.draw_button(self.bg_color)

    def on_click(self):
        click_color = self.darken_color(self.bg_color, 0.15)
        self.draw_button(click_color)

    def on_release(self):
        self.draw_button(self.hover_color)
        if self.command:
            self.command()

    def update_text(self, new_text):
        self.text = new_text
        self.draw_button(self.bg_color)

    def update_colors(self):
        self.update_colors_from_theme()
        self.draw_button(self.bg_color)

    def darken_color(self, hex_color, factor):
        try:
            hex_color = hex_color.lstrip('#')
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            darkened_rgb = tuple(max(0, int(c * (1 - factor))) for c in rgb)
            return '#{:02x}{:02x}{:02x}'.format(*darkened_rgb)
        except:
            return hex_color

# --- DIÁLOGO DE CARREGAMENTO (MODAL) ---
class ProcessingModal:
    def __init__(self, parent, title, message):
        self.top = tk.Toplevel(parent)
        self.top.title(title)
        self.top.geometry("380x130")
        self.top.transient(parent)
        self.top.grab_set()

        # Centralizar
        self.top.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - 190
        y = parent.winfo_y() + (parent.winfo_height() // 2) - 65
        self.top.geometry(f"+{x}+{y}")

        colors = ThemeManager.get_colors()
        self.top.configure(bg=colors["bg"])

        self.lbl_msg = tk.Label(self.top, text=message, font=("Segoe UI", 11, "bold"), bg=colors["bg"], fg=colors["fg"])
        self.lbl_msg.pack(pady=15)
        
        self.progress = ttk.Progressbar(self.top, orient="horizontal", mode="indeterminate", length=280)
        self.progress.pack(pady=5)
        self.progress.start(12)

        ThemeManager.apply_theme_to_all(self.top)

    def update_message(self, new_message):
        self.top.after(0, lambda: self.lbl_msg.config(text=new_message))

    def close(self):
        self.top.grab_release()
        self.top.destroy()

# --- FUNÇÕES AUXILIARES DE CAMINHO ---
def get_base_path():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = get_base_path()
    return os.path.join(base_path, relative_path)

# --- CONTROLE DE CACHE ---
def get_cache_path():
    return os.path.join(get_base_path(), CACHE_FILE)

def get_dashboard_cache_path():
    return os.path.join(get_base_path(), DASHBOARD_CACHE_FILE)

def load_cache(cache_path=None):
    if cache_path is None:
        cache_path = get_cache_path()
    if os.path.exists(cache_path):
        try:
            return pd.read_json(cache_path, dtype=str)
        except Exception:
            pass
    return pd.DataFrame()

def save_cache(df, cache_path=None):
    if cache_path is None:
        cache_path = get_cache_path()
    try:
        df.to_json(cache_path, force_ascii=False, indent=2)
    except Exception as e:
        print(f"Falha ao salvar cache: {e}")

# --- FUNÇÕES DE CONFIGURAÇÃO (PERSISTÊNCIA) ---
def get_config_path():
    return os.path.join(get_base_path(), CONFIG_FILE)

def load_config():
    path = get_config_path()
    default_dir = os.path.join(os.path.expanduser("~"), "Documents")
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return {"unificador_dest": default_dir, "conversor_dest": default_dir, "theme": "dark"}

def save_config(key, value):
    config = load_config()
    config[key] = value
    try:
        with open(get_config_path(), 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2)
    except Exception as e:
        print(f"Erro ao salvar config: {e}")

def show_detailed_error(title, message, error_obj):
    error_window = tk.Toplevel()
    error_window.title(title)
    error_window.geometry("700x500")
    
    colors = ThemeManager.get_colors()
    error_window.configure(bg=colors["bg"])

    tk.Label(error_window, text=message, font=("Segoe UI", 11, "bold"), fg="red", bg=colors["bg"]).pack(pady=10)
    
    container = tk.Frame(error_window, bg=colors["bg"])
    container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    scrollbar_y = ttk.Scrollbar(container)
    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    
    text_area = tk.Text(container, wrap=tk.NONE, yscrollcommand=scrollbar_y.set, bg=colors["entry_bg"], fg=colors["entry_fg"])
    text_area.pack(fill=tk.BOTH, expand=True)
    scrollbar_y.config(command=text_area.yview)
    
    full_error = f"Erro: {error_obj}\n\nTraceback:\n{traceback.format_exc()}"
    text_area.insert(tk.END, full_error)
    text_area.config(state=tk.DISABLED)
    
    tk.Button(error_window, text="Fechar", command=error_window.destroy, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=10)
    
    ThemeManager.apply_theme_to_all(error_window)

def authenticate_google_sheets():
    try:
        creds_path = os.path.join(get_base_path(), "credentials.json")
        if not os.path.exists(creds_path):
            creds_path = "credentials.json"
        if not os.path.exists(creds_path):
            return None
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_name(creds_path, scope)
        return gspread.authorize(creds)
    except Exception as e:
        print(f"Erro na autenticação Google: {e}")
        return None

def fetch_data_by_id(gspread_client, sheet_id, worksheet_name):
    try:
        spreadsheet = gspread_client.open_by_key(sheet_id)
        worksheet = spreadsheet.worksheet(worksheet_name)
        raw_data = worksheet.get_all_values()
        if not raw_data: return pd.DataFrame()
        return pd.DataFrame(raw_data[1:], columns=raw_data[0])
    except Exception as e: raise Exception(f"Erro na aba '{worksheet_name}': {e}")

# --- CLASSES DAS FERRAMENTAS ---


# ==============================================================================
# CONFIGURACOES E FUNCOES DE EMAIL DO PLEITO
# ==============================================================================
EMAIL_TO_PADRAO = "planejamento@eace.org.br"
EMAIL_CC_PADRAO = ""

def obter_assunto_email(dados):
    return f"NMA SERVICOS DE TELECOMUNICAÇÕES LTDA - Alteração de RI - INEP {dados['inep']}"

def obter_corpo_email_html(dados):
    adicionais = []
    if dados.get('qtd_switch', 0) > 0: adicionais.append(f"{dados['qtd_switch']}x Switch")
    if dados.get('qtd_rack', 0) > 0: adicionais.append(f"{dados['qtd_rack']}x Rack")
    if dados.get('qtd_nobreak', 0) > 0: adicionais.append(f"{dados['qtd_nobreak']}x Nobreak")
    
    html_adicionais = ""
    if adicionais:
        texto_adic = ", ".join(adicionais)
        html_adicionais = f"""
                <div style="background-color: #fffbeb; border-left: 4px solid #f59e0b; padding: 15px; margin: 25px 0; border-radius: 0 4px 4px 0;">
                    <p style="margin: 0; color: #b45309; font-size: 14px;"><strong>Adicionais de Infraestrutura:</strong><br>
                    Além dos Access Points, a vistoria técnica identificou a necessidade de incluir os seguintes equipamentos para viabilizar a instalação: <strong>{texto_adic}</strong>.</p>
                </div>
        """

    cor_pleito = '#059669' if dados['tipo'] == 'UPGRADE' else '#dc2626'
    return f"""
    <html>
    <body style="font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background-color: #f9fafb; margin: 0; padding: 20px;">
        <div style="max-width: 800px; margin: 0 auto; background-color: #ffffff; border: 1px solid #e5e7eb; border-radius: 8px; overflow: hidden; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);">
            
            <div style="background-color: #1e293b; padding: 20px; text-align: center;">
                <h2 style="color: #ffffff; margin: 0; font-size: 20px; font-weight: 600;">SOLICITAÇÃO DE ALTERAÇÃO DE PROJETO</h2>
            </div>
            
            <div style="padding: 30px;">
                <p style="color: #334155; font-size: 15px; margin-top: 0;">Olá equipe EACE,</p>
                <p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Submetemos para análise e aprovação técnica a solicitação de <strong style="color: {cor_pleito};">{dados['tipo']}</strong> de Access Points para a unidade escolar descrita abaixo.
                </p>

                <div style="background-color: #f8fafc; border-left: 4px solid #3b82f6; padding: 15px; margin: 25px 0; border-radius: 0 4px 4px 0;">
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Escola:</strong> {dados['escola']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Código INEP:</strong> {dados['inep']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Localidade:</strong> {dados['municipio']} - {dados['uf']}</p>
                    <p style="margin: 5px 0; color: #1e293b;"><strong>Endereço:</strong> {dados['endereco']}</p>
                </div>

                <div style="overflow-x: auto;">
                    <table style="width: 100%; border-collapse: collapse; font-size: 12px; text-align: center; margin-bottom: 25px; border: 1px solid #cbd5e1;">
                        <thead>
                            <tr style="background-color: #f1f5f9; color: #334155;">
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">FASE</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">INEP</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">ESCOLA</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">ESTADO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">CIDADE</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">PREVISTO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">PLEITO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">SUGERIDO</th>
                                <th style="padding: 10px; border: 1px solid #cbd5e1;">INFANTIL</th>
                            </tr>
                        </thead>
                        <tbody>
                            <tr style="color: #475569;">
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">5</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['inep']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['escola']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['uf']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['municipio']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold;">{dados['aps_atuais']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold; color: {cor_pleito};">{dados['tipo']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1; font-weight: bold;">{dados['novos_aps']}</td>
                                <td style="padding: 10px; border: 1px solid #cbd5e1;">{dados['escola_infantil']}</td>
                            </tr>
                        </tbody>
                    </table>
                </div>

                <p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Após vistoria técnica e validação da infraestrutura local, constatamos a necessidade de adequação. O projeto original prevê <strong>{dados['aps_atuais']} AP(s)</strong>, contudo, para garantir o correto funcionamento e cobertura dos ambientes, faz-se necessária a instalação de <strong>{dados['novos_aps']} AP(s)</strong>.
                </p>
<p style="color: #334155; font-size: 15px; line-height: 1.6;">
                    Ficamos no aguardo da aprovação técnica formal para prosseguirmos com a execução do cronograma.
                </p>
                
                <div style="margin-top: 40px; border-top: 1px solid #e2e8f0; padding-top: 20px;">
                    <p style="margin: 0; color: #64748b; font-size: 14px;">Atenciosamente,</p>
                    <p style="margin: 5px 0 0 0; color: #1e293b; font-weight: bold; font-size: 15px;">Equipe de Projetos</p>
                    <p style="margin: 2px 0 0 0; color: #64748b; font-size: 13px;">NMA SERVIÇOS DE TELECOMUNICAÇÕES LTDA</p>
                </div>
            </div>
        </div>
    </body>
    </html>
    """

def gerar_conteudo_docx(doc, dados):
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    doc.add_paragraph("Prezados,")
    doc.add_paragraph()
    doc.add_paragraph(f"Gostaria de solicitar a autorização da EACE para realização de {dados['tipo'].lower()} na quantidade de Access Points da escola abaixo:")
    doc.add_paragraph(f"• Escola: {dados['escola']}", style='List Bullet')
    doc.add_paragraph(f"• Código INEP / Identificação: {dados['inep']}", style='List Bullet')
    doc.add_paragraph(f"• Município: {dados['municipio']} - {dados['uf']}", style='List Bullet')
    doc.add_paragraph(f"• Endereço: {dados['endereco']}", style='List Bullet')
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=12)
    table.style = 'Table Grid'
    headers = ["FASE", "INEP", "ESCOLA", "ESTADO", "CIDADE", "KIT PREVISTO", "PLEITO", "KIT SUGERIDO", "LATITUDE", "LONGITUDE", "ESCOLA INFANTIL", "JUSTIFICATIVA"]
    for i, h in enumerate(headers): 
        table.cell(0, i).text = h
        
    cells = table.rows[1].cells
    vals = ["5", dados['inep'], dados['escola'], dados['uf'], dados['municipio'], dados['aps_atuais'], dados['tipo'], dados['novos_aps'], dados['latitude'], dados['longitude'], dados['escola_infantil'], "Adequação à necessidade da escola"]
    for i, v in enumerate(vals): 
        cells[i].text = str(v)
        
    doc.add_paragraph()
    doc.add_paragraph(f"Conforme consta na lista atual, a escola possui previsão de {dados['aps_atuais']} APs. Porém, após validação da necessidade do ambiente, identificamos que será necessário {dados['tipo'].lower()} para {dados['novos_aps']} APs, a fim de atender de forma adequada a cobertura e o funcionamento da unidade.")
    doc.add_paragraph()
    doc.add_paragraph(f"Dessa forma, solicitamos a autorização da EACE para atualização da quantidade de equipamentos desta escola, passando de {dados['aps_atuais']} AP para {dados['novos_aps']} APs.")
    doc.add_paragraph()
    doc.add_paragraph("Ficamos no aguardo da aprovação para seguir com o atendimento.")
    doc.add_paragraph()
    doc.add_paragraph("Atenciosamente,")

def authenticate_gmail():
    SCOPES = ['https://www.googleapis.com/auth/gmail.send']
    token_path = os.path.join(get_base_path(), "token.json")
    secret_path = os.path.join(get_base_path(), "AuthNOC.json")
    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(secret_path):
                raise FileNotFoundError(f"Arquivo '{secret_path}' não encontrado.")
            flow = InstalledAppFlow.from_client_secrets_file(secret_path, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token: 
            token.write(creds.to_json())
    return build('gmail', 'v1', credentials=creds)



class EnviadorPleitoApp:
    def __init__(self, parent_frame, main_app_callback, gspread_client=None):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill="both", expand=True)
        self.main_app_callback = main_app_callback
        self.root = parent_frame.winfo_toplevel()
        self.selected_file_paths = []

        tk.Button(self.frame, text="VOLTAR AO MENU PRINCIPAL", command=self.main_app_callback, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=(5, 10))

        self.criar_interface()
        ThemeManager.apply_theme_to_all(self.frame)

    def criar_interface(self):
        # Container Principal
        self.container = tk.Frame(self.frame, padx=20, pady=20)
        self.container.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)
        
        self.lbl_title = tk.Label(self.container, text="📝 AUTOMAÇÃO EACE (E-MAIL/DOCX)", font=("Arial", 16, "bold"))
        self.lbl_title.pack(anchor="w", pady=(0, 20))
        
        # Ocultar campos de email em um frame mais discreto
        self.frame_emails = tk.Frame(self.container)
        self.frame_emails.pack(fill=tk.X, pady=(0, 15))
        self.lbl_to = tk.Label(self.frame_emails, text="Para:", font=("Arial", 9, "bold"))
        self.lbl_to.grid(row=0, column=0, sticky="w")
        self.ent_to = tk.Entry(self.frame_emails, font=("Arial", 9), relief="flat", width=40)
        self.ent_to.grid(row=0, column=1, sticky="w", padx=5, pady=2, ipady=2)
        self.ent_to.insert(0, EMAIL_TO_PADRAO)
        
        self.lbl_cc = tk.Label(self.frame_emails, text="Cc (separar por vírgula):", font=("Arial", 9, "bold"))
        self.lbl_cc.grid(row=1, column=0, sticky="w")
        self.ent_cc = tk.Entry(self.frame_emails, font=("Arial", 9), relief="flat", width=40)
        self.ent_cc.grid(row=1, column=1, sticky="w", padx=5, pady=2, ipady=2)
        
        # Inputs principais (IDêntico à imagem)
        self.lbl_inep = tk.Label(self.container, text="Código INEP da Escola:", font=("Arial", 11, "bold"))
        self.lbl_inep.pack(anchor="w", pady=(10, 2))
        self.ent_inep = tk.Entry(self.container, font=("Arial", 12), relief="flat", width=25)
        self.ent_inep.pack(anchor="w", ipady=4, pady=(0, 15))
        
        self.lbl_aps = tk.Label(self.container, text="Nova Quantidade de APs (Pleito):", font=("Arial", 11, "bold"))
        self.lbl_aps.pack(anchor="w", pady=(10, 2))
        self.ent_aps = tk.Entry(self.container, font=("Arial", 12), relief="flat", width=25)
        self.ent_aps.pack(anchor="w", ipady=4, pady=(0, 15))
        
        self.var_infantil = tk.StringVar(value="NÃO")
        self.chk_infantil = tk.Checkbutton(self.container, text="Escola Infantil?", variable=self.var_infantil, onvalue="SIM", offvalue="NÃO", font=("Arial", 10, "bold"), fg="#eab308", selectcolor="#1e1e24", cursor="hand2")
        self.chk_infantil.pack(anchor="w", pady=(0, 20))
        
        self.frame_adicionais = tk.LabelFrame(self.container, text="Adicionais de Infraestrutura (Qtd)", font=("Arial", 10, "bold"), padx=10, pady=10)
        self.frame_adicionais.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(self.frame_adicionais, text="Switch:", font=("Arial", 9)).grid(row=0, column=0, padx=5, sticky="e")
        self.ent_switch = tk.Entry(self.frame_adicionais, font=("Arial", 9), width=5)
        self.ent_switch.grid(row=0, column=1, padx=5)
        self.ent_switch.insert(0, "0")
        
        tk.Label(self.frame_adicionais, text="Rack:", font=("Arial", 9)).grid(row=0, column=2, padx=5, sticky="e")
        self.ent_rack = tk.Entry(self.frame_adicionais, font=("Arial", 9), width=5)
        self.ent_rack.grid(row=0, column=3, padx=5)
        self.ent_rack.insert(0, "0")
        
        tk.Label(self.frame_adicionais, text="Nobreak:", font=("Arial", 9)).grid(row=0, column=4, padx=5, sticky="e")
        self.ent_nobreak = tk.Entry(self.frame_adicionais, font=("Arial", 9), width=5)
        self.ent_nobreak.grid(row=0, column=5, padx=5)
        self.ent_nobreak.insert(0, "0")
        
        self.lbl_pdf = tk.Label(self.container, text="Anexo de Evidência Local (Opcional):", font=("Arial", 11, "bold"))
        self.lbl_pdf.pack(anchor="w", pady=(10, 2))
        
        self.frame_pdf = tk.Frame(self.container)
        self.frame_pdf.pack(fill=tk.X, pady=(0, 20))
        
        self.ent_pdf = tk.Entry(self.frame_pdf, font=("Arial", 10), relief="flat", state="disabled")
        self.ent_pdf.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4, padx=(0, 5))
        
        self.btn_sel_pdf = tk.Button(self.frame_pdf, text="📎 Selecionar Arquivo...", command=self.selecionar_pdf, font=("Arial", 9, "bold"), relief="flat", cursor="hand2")
        self.btn_sel_pdf.pack(side=tk.LEFT, padx=2, ipady=2)
        
        self.btn_limpar_pdf = tk.Button(self.frame_pdf, text="❌ Limpar", command=self.limpar_pdf, font=("Arial", 9, "bold"), relief="flat", cursor="hand2")
        self.btn_limpar_pdf.pack(side=tk.LEFT, padx=2, ipady=2)
        
        # Botões
        self.frame_botoes_1 = tk.Frame(self.container)
        self.frame_botoes_1.pack(fill=tk.X, pady=(10, 5))
        
        self.btn_docx = tk.Button(self.frame_botoes_1, text="📄 GERAR DOCX", command=self.acao_gerar_docx, font=("Arial", 11, "bold"), relief="flat", cursor="hand2")
        self.btn_docx.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5), ipady=8)
        
        self.btn_preview = tk.Button(self.frame_botoes_1, text="👁 VISUALIZAR PRÉVIA", command=self.acao_visualizar, font=("Arial", 11, "bold"), relief="flat", cursor="hand2")
        self.btn_preview.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=(5, 0), ipady=8)
        
        self.btn_enviar = tk.Button(self.container, text="🚀 ENVIAR POR E-MAIL", command=self.acao_enviar, font=("Arial", 12, "bold"), relief="flat", cursor="hand2")
        self.btn_enviar.pack(fill=tk.X, pady=(5, 10), ipady=8)
        
        # Loading
        self.frame_loading = tk.Frame(self.container)
        self.lbl_loading = tk.Label(self.frame_loading, text="⏳ Autenticando e enviando e-mail, aguarde...", font=("Arial", 10, "italic"))
        self.lbl_loading.pack(pady=(5, 5))
        self.progress = ttk.Progressbar(self.frame_loading, mode="indeterminate")
        self.progress.pack(fill=tk.X, expand=True)
        self.frame_loading.pack_forget() # Oculta no início

    def selecionar_pdf(self):
        caminhos = filedialog.askopenfilenames(filetypes=[("Arquivos Suportados", "*.pdf *.jpg *.jpeg *.png"), ("Todos os arquivos", "*.*")])
        if caminhos:
            self.selected_file_paths = list(caminhos)
            self.ent_pdf.config(state="normal")
            self.ent_pdf.delete(0, tk.END)
            nomes = ", ".join([os.path.basename(c) for c in caminhos])
            self.ent_pdf.insert(0, nomes)
            self.ent_pdf.config(state="disabled")

    def limpar_pdf(self):
        self.selected_file_paths = []
        self.ent_pdf.config(state="normal")
        self.ent_pdf.delete(0, tk.END)
        self.ent_pdf.config(state="disabled")

    def buscar_dados_bd(self):
        inep = self.ent_inep.get().strip()
        novos_aps_str = self.ent_aps.get().strip()
        
        if not inep or not novos_aps_str:
            raise ValueError("Por favor, preencha o Código INEP e a Nova Quantidade de APs.")
            
        cache_file = os.path.join(get_base_path(), "eace_cache.json")
        infantil_cache_file = os.path.join(get_base_path(), "infantil_cache.json")
        
        if not os.path.exists(cache_file):
            raise Exception("Banco de dados 'eace_cache.json' não encontrado na pasta.")
            
        df = pd.read_json(cache_file, dtype=str)
        col = "Código INEP" if "Código INEP" in df.columns else df.columns[0]
        df[col] = df[col].astype(str).str.split('.').str[0].str.strip()
        
        res = df[df[col] == inep]
        if res.empty:
            raise Exception(f"O INEP '{inep}' não foi localizado no banco de dados local.")
            
        d = res.iloc[0]
        
        escola_infantil = self.var_infantil.get()
        if escola_infantil == "NÃO" and os.path.exists(infantil_cache_file):
            try:
                df_inf = pd.read_json(infantil_cache_file, dtype=str)
                col_inf = "Código INEP" if "Código INEP" in df_inf.columns else df_inf.columns[0]
                ineps_infantis = df_inf[col_inf].astype(str).str.split('.').str[0].str.strip().tolist()
                if inep in ineps_infantis:
                    escola_infantil = "SIM"
                    self.var_infantil.set("SIM")
            except Exception:
                pass
                
        aps_atuais_str = str(d.get('Kit Wi-Fi', '0')).strip()
        try:
            aps_atuais = int(aps_atuais_str or 0)
            novos_aps = int(novos_aps_str)
            tipo = "UPGRADE" if novos_aps > aps_atuais else "DOWNGRADE"
        except ValueError:
            aps_atuais = aps_atuais_str
            novos_aps = novos_aps_str
            tipo = "UPGRADE"
            
        try:
            qtd_switch = int(self.ent_switch.get().strip() or "0")
        except: qtd_switch = 0
        try:
            qtd_rack = int(self.ent_rack.get().strip() or "0")
        except: qtd_rack = 0
        try:
            qtd_nobreak = int(self.ent_nobreak.get().strip() or "0")
        except: qtd_nobreak = 0
        
        return {
            "qtd_switch": qtd_switch,
            "qtd_rack": qtd_rack,
            "qtd_nobreak": qtd_nobreak,
            "inep": inep,
            "escola": str(d.get('Nome da Escola', '')).strip(),
            "uf": str(d.get('UF', '')).strip(),
            "municipio": str(d.get('Município', '')).strip(),
            "endereco": str(d.get('Endereço', '')).strip(),
            "latitude": str(d.get('Latitude', '')).strip(),
            "longitude": str(d.get('Longitude', '')).strip(),
            "aps_atuais": aps_atuais,
            "novos_aps": novos_aps,
            "tipo": tipo,
            "escola_infantil": escola_infantil
        }

    def acao_gerar_docx(self):
        try:
            dados = self.buscar_dados_bd()
        except Exception as e:
            messagebox.showwarning("Aviso", str(e))
            return
            
        destino = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"Pleito_{dados['inep']}.docx", filetypes=[("Documento Word", "*.docx")])
        if destino:
            try:
                doc = Document()
                gerar_conteudo_docx(doc, dados)
                doc.save(destino)
                messagebox.showinfo("Sucesso", f"DOCX gerado e salvo com sucesso em:\n{destino}")
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao gerar o DOCX:\n{str(e)}")

    def acao_visualizar(self):
        try:
            dados = self.buscar_dados_bd()
        except Exception as e:
            messagebox.showwarning("Aviso", str(e))
            return
            
        html_body = obter_corpo_email_html(dados)
        assunto = obter_assunto_email(dados)
        
        try:
            with tempfile.NamedTemporaryFile('w', delete=False, suffix='.html', encoding='utf-8') as f:
                html_preview = f"<html><body><div style='background-color: #f1f5f9; padding: 15px; border-bottom: 2px solid #cbd5e1; margin-bottom: 20px; font-family: Arial;'><h3 style='margin:0; color: #1e293b;'>Pré-visualização do E-mail</h3><p style='margin: 5px 0 0 0;'><b>Assunto:</b> {assunto}</p></div>{html_body}</body></html>"
                f.write(html_preview)
                temp_path = f.name
            webbrowser.open('file://' + temp_path)
        except Exception as e:
            messagebox.showwarning("Erro na pré-visualização", f"Não foi possível abrir o navegador: {str(e)}")

    def acao_enviar(self):
        to_email = self.ent_to.get().strip()
        cc_email = self.ent_cc.get().strip()
        if not to_email:
            messagebox.showwarning("Aviso", "O campo 'Para' não pode estar vazio.")
            return
            
        try:
            dados = self.buscar_dados_bd()
        except Exception as e:
            messagebox.showwarning("Aviso", str(e))
            return
            
        html_body = obter_corpo_email_html(dados)
        assunto = obter_assunto_email(dados)
        
        if messagebox.askyesno("Confirmar Envio", "Tem certeza que deseja disparar o e-mail agora?"):
            self.frame_loading.pack(fill=tk.X, pady=10)
            self.progress.start(10)
            
            thread = threading.Thread(target=self._enviar_email_bg, args=(dados, to_email, cc_email, html_body, assunto))
            thread.daemon = True
            thread.start()

    def _enviar_email_bg(self, dados, to_email, cc_email, html_body, assunto):
        sucesso = False
        erro_msg = ""
        try:
            gmail_service = authenticate_gmail()
            
            msg = EmailMessage()
            msg['To'] = to_email
            if cc_email:
                msg['Cc'] = cc_email
            msg['Subject'] = assunto
            
            msg.add_alternative(html_body, subtype='html')
            
            if hasattr(self, 'selected_file_paths') and self.selected_file_paths:
                import mimetypes
                for path in self.selected_file_paths:
                    if os.path.exists(path):
                        content_type, encoding = mimetypes.guess_type(path)
                        if content_type is None:
                            content_type = 'application/octet-stream'
                        main_type, sub_type = content_type.split('/', 1)

                        with open(path, 'rb') as f:
                            msg.add_attachment(f.read(), maintype=main_type, subtype=sub_type, filename=os.path.basename(path))
                    
            raw_msg = base64.urlsafe_b64encode(msg.as_bytes()).decode()
            gmail_service.users().messages().send(userId='me', body={'raw': raw_msg}).execute()
            sucesso = True
            
        except Exception as e:
            erro_msg = str(e)
            
        self.frame.after(0, lambda: self._finalizar_envio(sucesso, erro_msg))

    def _finalizar_envio(self, sucesso, erro_msg):
        self.progress.stop()
        self.frame_loading.pack_forget()
        
        if sucesso:
            messagebox.showinfo("Sucesso", "E-mail enviado com sucesso!")
        else:
            if "AuthNOC.json" in erro_msg or "não encontrado" in erro_msg:
                messagebox.showerror("Erro de Autenticação", f"Arquivo de chaves faltando:\n{erro_msg}")
            else:
                messagebox.showerror("Erro ao Enviar", f"Ocorreu um erro ao enviar o e-mail:\n{erro_msg}")


class GeradorContratoApp:
    def __init__(self, parent_frame, main_app_callback, gspread_client=None):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill="both", expand=True)
        self.main_app_callback = main_app_callback
        self.root = parent_frame.winfo_toplevel()
        
        self.template = None
        self.df = None
        self.escolas = []
        self.col_inep = None    # Coluna INEP detectada na planilha
        self.df_index = None    # Índice O(1) para busca rápida por INEP

        # Botão Voltar
        tk.Button(self.frame, text="VOLTAR AO MENU PRINCIPAL", command=self.main_app_callback, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=(5, 10))

        self.tabs = ttk.Notebook(self.frame)
        self.tabs.pack(fill="both", expand=True)

        self.criar_abas()
        self.auto_detectar_arquivos()
        ThemeManager.apply_theme_to_all(self.frame)

    def criar_abas(self):
        self.aba_arquivos()
        self.aba_contratada()
        self.aba_suporte()
        self.aba_assinaturas()
        self.aba_ineps()
        self.aba_gerar()

    def aba_arquivos(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="Arquivos")
        
        f_conteudo = tk.Frame(aba)
        f_conteudo.pack(pady=40)
        
        ttk.Button(f_conteudo, text="Selecionar Planilha EACE", command=self.carregar_planilha).grid(row=0, column=0, padx=15, pady=10)
        
        self.lbl_planilha = ttk.Label(aba, text="Planilha não carregada", font=("Segoe UI", 10, "italic"))
        self.lbl_planilha.pack(pady=10)

    def auto_detectar_arquivos(self):
        base_path = get_base_path()
        
        xlsx_files = [f for f in os.listdir(base_path) if f.lower().endswith('.xlsx') and not f.startswith('~$')]
        if xlsx_files:
            try:
                path = os.path.join(base_path, xlsx_files[0])
                self.df = pd.read_excel(path)
                self.lbl_planilha.config(text=f"Planilha (Auto): {xlsx_files[0]}")
                self._build_df_index()
            except Exception as e:
                print(f"Erro ao carregar planilha automática: {e}")

    def aba_contratada(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="Contratada")
        
        # Adicionar Scrollbar na aba para garantir visibilidade
        canvas = tk.Canvas(aba, borderwidth=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(aba, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y")
        
        self.contratada = {}
        campos = [
            ("Razão Social", "razao"), ("CNPJ", "cnpj"), ("Endereço", "endereco"),
            ("UF", "uf"), ("Bairro", "bairro"), ("CEP", "cep"),
            ("Cidade", "cidade"), ("IE", "ie"), ("Telefone", "telefone"),
            ("Site", "site"), ("E-mail", "email")
        ]
        for i, (txt, key) in enumerate(campos):
            ttk.Label(scrollable_frame, text=txt).grid(row=i, column=0, padx=10, pady=5, sticky="w")
            e = ttk.Entry(scrollable_frame, width=80)
            e.grid(row=i, column=1, padx=10, pady=5, sticky="w")
            self.contratada[key] = e

    def aba_suporte(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="Suporte")
        self.sup = {}
        campos = [("Email Suporte","email_sup"), ("Telefone Suporte","tel_sup"), ("Sistema Chamado","sistema")]
        for i, (txt, key) in enumerate(campos):
            ttk.Label(aba, text=txt).grid(row=i, column=0, padx=10, pady=10, sticky="w")
            e = ttk.Entry(aba, width=80)
            e.grid(row=i, column=1, padx=10, pady=10, sticky="w")
            self.sup[key] = e

    def aba_assinaturas(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="Assinaturas")
        self.ass = {}
        campos = [
            ("Local", "local"), 
            ("Data", "data"), 
            ("Contratante", "contratante"), 
            ("Testemunha", "testemunha"),
            ("Contratada", "contratada")
        ]
        for i, (txt, key) in enumerate(campos):
            ttk.Label(aba, text=txt).grid(row=i, column=0, padx=10, pady=10, sticky="w")
            e = ttk.Entry(aba, width=80)
            e.grid(row=i, column=1, padx=10, pady=10, sticky="w")
            self.ass[key] = e
        self.ass["data"].insert(0, datetime.now().strftime("%d/%m/%Y"))

    def aba_ineps(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="INEPs")
        
        top = ttk.Frame(aba)
        top.pack(fill="x", padx=10, pady=10)
        
        # Múltiplos INEPs via Textbox
        ttk.Label(top, text="Insira os códigos INEP\n(um por linha ou separados por vírgula):", font=("Segoe UI", 9, "bold")).grid(row=0, column=0, padx=5, sticky="nw")
        
        self.txt_ineps = tk.Text(top, height=4, width=40, font=("Consolas", 10))
        self.txt_ineps.grid(row=0, column=1, padx=5, sticky="w")
        
        btn_frame = ttk.Frame(top)
        btn_frame.grid(row=0, column=2, padx=10, sticky="nw")
        
        ttk.Button(btn_frame, text="Adicionar INEPs", command=self.add_inep_lote).pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="Remover Selecionado", command=self.remover_inep).pack(fill="x", pady=2)

        self.tree = ttk.Treeview(aba, columns=("inep", "endereco", "latlong", "megas", "mensalidade", "instalacao"), show="headings", height=12)
        colunas = {"inep":"INEP", "endereco":"ENDEREÇO", "latlong":"LAT/LONG", "megas":"BANDA", "mensalidade":"MENSALIDADE", "instalacao":"INSTALAÇÃO"}
        for c, t in colunas.items():
            self.tree.heading(c, text=t)
            self.tree.column(c, width=130, anchor="center" if c in ["inep", "megas", "mensalidade", "instalacao"] else "w")
        self.tree.pack(fill="both", expand=True, padx=10, pady=10)

    def remover_inep(self):
        selected_item = self.tree.selection()
        if not selected_item:
            messagebox.showwarning("Aviso", "Selecione um INEP na lista para remover.")
            return
        
        for item in selected_item:
            values = self.tree.item(item, "values")
            inep_to_remove = values[0]
            self.escolas = [esc for esc in self.escolas if esc["inep"] != inep_to_remove]
            self.tree.delete(item)

    def aba_gerar(self):
        aba = ttk.Frame(self.tabs)
        self.tabs.add(aba, text="Gerar Contrato")
        
        f_centro = tk.Frame(aba)
        f_centro.pack(expand=True)
        
        tk.Button(f_centro, text="⚡ GERAR CONTRATO FINAL (WORD) ⚡", command=self.gerar_assincrono, bg="#27ae60", fg="white", font=("Segoe UI", 14, "bold"), padx=25, pady=15).pack()



    def carregar_planilha(self):
        arq = filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")])
        if arq:
            try:
                self.df = pd.read_excel(arq)
                self.lbl_planilha.config(text=f"Planilha: {os.path.basename(arq)}")
                self._build_df_index()
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao carregar planilha: {e}")

    def localizar(self, linha, termos):
        for termo in termos:
            for col in linha.index:
                if termo.lower() in str(col).lower(): return linha[col]
        return ""

    def formatar_data_input(self, data_str):
        apenas_numeros = re.sub(r'\D', '', str(data_str))
        if len(apenas_numeros) == 8:
            return f"{apenas_numeros[:2]}/{apenas_numeros[2:4]}/{apenas_numeros[4:]}"
        if not data_str or str(data_str).strip() == "":
            return "__/__/____"
        return data_str

    def pedir_dados_escola(self, inep):
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Dados INEP: {inep}")
        dialog.geometry("360x160")
        dialog.transient(self.root)
        dialog.grab_set()

        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 180
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 80
        dialog.geometry(f"+{x}+{y}")

        res = {"val": None, "data": None}
        colors = ThemeManager.get_colors()
        dialog.configure(bg=colors["bg"])

        tk.Label(dialog, text="Mensalidade (R$):", bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=0, padx=15, pady=15, sticky="e")
        e_val = ttk.Entry(dialog, width=20)
        e_val.grid(row=0, column=1, padx=15, pady=15)
        e_val.focus()

        tk.Label(dialog, text="Instalação (DDMMAAAA):", bg=colors["bg"], fg=colors["fg"]).grid(row=1, column=0, padx=15, pady=10, sticky="e")
        e_data = ttk.Entry(dialog, width=20)
        e_data.grid(row=1, column=1, padx=15, pady=10)
        e_data.insert(0, datetime.now().strftime("%d%m%Y"))

        def salvar():
            txt_val = e_val.get().strip().replace(',', '.')
            try:
                res["val"] = float(txt_val) if txt_val else 0.0
            except ValueError:
                messagebox.showerror("Erro", "Valor inválido. Digite apenas números.", parent=dialog)
                return
            res["data"] = e_data.get().strip()
            dialog.destroy()

        def cancelar():
            dialog.destroy()

        f_btn = ttk.Frame(dialog)
        f_btn.grid(row=2, column=0, columnspan=2, pady=10)
        ttk.Button(f_btn, text="Confirmar", command=salvar).pack(side="left", padx=10)
        ttk.Button(f_btn, text="Cancelar", command=cancelar).pack(side="left", padx=10)

        ThemeManager.apply_theme_to_all(dialog)
        self.root.wait_window(dialog)
        return res["val"], res["data"]

    def pedir_dados_lote(self, ineps_encontrados):
        dialog = tk.Toplevel(self.root)
        dialog.title("Configuração de INEPs em Lote")
        dialog.geometry("680x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 340
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 250
        dialog.geometry(f"+{x}+{y}")
        
        colors = ThemeManager.get_colors()
        dialog.configure(bg=colors["bg"])
        
        # Frame de lote rápido
        f_lote = tk.LabelFrame(dialog, text="Preenchimento Rápido (Lote)", bg=colors["bg"], fg=colors["fg"], padx=10, pady=10)
        f_lote.pack(fill="x", padx=10, pady=10)
        
        tk.Label(f_lote, text="Mensalidade R$:", bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        e_val_padrao = ttk.Entry(f_lote, width=12)
        e_val_padrao.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(f_lote, text="Instalação (DDMMAAAA):", bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=2, padx=5, pady=5, sticky="e")
        e_data_padrao = ttk.Entry(f_lote, width=12)
        e_data_padrao.grid(row=0, column=3, padx=5, pady=5)
        e_data_padrao.insert(0, datetime.now().strftime("%d%m%Y"))
        
        entries_list = []
        
        # Rolagem
        f_scroll = tk.Frame(dialog, bg=colors["bg"])
        f_scroll.pack(fill="both", expand=True, padx=10, pady=5)
        
        canvas = tk.Canvas(f_scroll, bg=colors["bg"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(f_scroll, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=colors["bg"])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Cabeçalhos
        tk.Label(scrollable_frame, text="INEP - Escola", font=("Segoe UI", 9, "bold"), bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        tk.Label(scrollable_frame, text="Mensalidade R$", font=("Segoe UI", 9, "bold"), bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        tk.Label(scrollable_frame, text="Instalação", font=("Segoe UI", 9, "bold"), bg=colors["bg"], fg=colors["fg"]).grid(row=0, column=2, padx=5, pady=5, sticky="w")
        
        for idx, item in enumerate(ineps_encontrados):
            row_idx = idx + 1
            esc_nome = item["escola"]
            lbl_text = f"{item['inep']} - {esc_nome[:45]}..." if len(esc_nome) > 45 else f"{item['inep']} - {esc_nome}"
            lbl = tk.Label(scrollable_frame, text=lbl_text, bg=colors["bg"], fg=colors["fg"], anchor="w", width=42)
            lbl.grid(row=row_idx, column=0, padx=5, pady=3, sticky="w")
            
            ent_val = ttk.Entry(scrollable_frame, width=12)
            ent_val.grid(row=row_idx, column=1, padx=5, pady=3)
            
            ent_data = ttk.Entry(scrollable_frame, width=12)
            ent_data.grid(row=row_idx, column=2, padx=5, pady=3)
            ent_data.insert(0, datetime.now().strftime("%d%m%Y"))
            
            entries_list.append((item, ent_val, ent_data))
            
        def aplicar_lote():
            val_p = e_val_padrao.get().strip().replace(',', '.')
            data_p = e_data_padrao.get().strip()
            for item, ent_val, ent_data in entries_list:
                if val_p:
                    ent_val.delete(0, "end")
                    ent_val.insert(0, val_p)
                if data_p:
                    ent_data.delete(0, "end")
                    ent_data.insert(0, data_p)
                    
        ttk.Button(f_lote, text="Aplicar a Todos", command=aplicar_lote).grid(row=0, column=4, padx=10, pady=5)
        
        resultado = []
        
        def salvar():
            for item, ent_val, ent_data in entries_list:
                txt_val = ent_val.get().strip().replace(',', '.')
                txt_data = ent_data.get().strip()
                
                try:
                    val = float(txt_val) if txt_val else 0.0
                except ValueError:
                    messagebox.showerror("Erro", f"Valor mensal inválido para o INEP {item['inep']}.", parent=dialog)
                    return
                
                resultado.append({
                    "inep": item["inep"],
                    "linha": item["linha"],
                    "mensalidade": val,
                    "data_inst": txt_data
                })
            dialog.destroy()
            
        def cancelar():
            dialog.destroy()
            
        f_btn = tk.Frame(dialog, bg=colors["bg"])
        f_btn.pack(fill="x", side="bottom", pady=15)
        
        tk.Button(f_btn, text="Confirmar", command=salvar, bg=colors["success"], fg="white", width=12, font=("Segoe UI", 9, "bold")).pack(side="right", padx=10)
        tk.Button(f_btn, text="Cancelar", command=cancelar, bg="#6c757d", fg="white", width=12).pack(side="right", padx=5)
        
        ThemeManager.apply_theme_to_all(dialog)
        self.root.wait_window(dialog)
        return resultado

    def _build_df_index(self):
        """Cria índice O(1) na coluna INEP para buscas ultrarrápidas."""
        if self.df is None or self.df.empty:
            return
        self.col_inep = None
        for c in self.df.columns:
            if "inep" in c.lower():
                self.col_inep = c
                break
        if not self.col_inep:
            self.col_inep = self.df.columns[0]
        # Normalizar a coluna e criar índice hash
        self.df[self.col_inep] = (
            self.df[self.col_inep].astype(str).str.strip().str.split('.').str[0]
        )
        self.df_index = self.df.set_index(self.col_inep, drop=False)

    def add_inep_lote(self):
        if self.df is None:
            messagebox.showerror("Erro", "Carregue a planilha EACE primeiro na aba 'Arquivos'.")
            return
        
        raw_text = self.txt_ineps.get("1.0", "end").strip()
        if not raw_text: return
        
        codigos_inep = re.findall(r'\b\d{8,10}\b', raw_text)
        if not codigos_inep:
            messagebox.showwarning("Aviso", "Nenhum código INEP válido (8 a 10 dígitos) foi encontrado no texto.")
            return
        
        # Filtrar duplicatas já inseridas na lista
        ja_inseridos = {esc["inep"] for esc in self.escolas}
        codigos_novos = [c for c in codigos_inep if c not in ja_inseridos]
        
        if not codigos_novos:
            messagebox.showinfo("Aviso", "Todos os INEPs informados já foram adicionados à lista.")
            return
        
        # Iniciar busca em segundo plano para não travar a UI
        modal = ProcessingModal(
            self.root, "Buscando INEPs",
            f"Pesquisando {len(codigos_novos)} INEP(s) na planilha...\nAguarde."
        )
        threading.Thread(
            target=self._buscar_ineps_thread,
            args=(codigos_novos, modal),
            daemon=True
        ).start()

    def _buscar_ineps_thread(self, codigos_novos, modal):
        """Executa busca de INEPs em thread separada — UI permanece responsiva."""
        try:
            ineps_encontrados = []
            ineps_nao_encontrados = []
            total = len(codigos_novos)
            
            if self.df_index is not None:
                # === Busca via índice hash: O(1) por INEP ===
                for i, inep in enumerate(codigos_novos):
                    if i % 10 == 0:
                        modal.update_message(
                            f"Buscando INEPs...  {i}/{total}\nAguarde."
                        )
                    inep_str = str(inep).strip()
                    if inep_str in self.df_index.index:
                        linha = self.df_index.loc[inep_str]
                        # Se mais de um registro com mesmo INEP, pegar o primeiro
                        if isinstance(linha, pd.DataFrame):
                            linha = linha.iloc[0]
                        esc_nome = str(self.localizar(linha, ["Nome da Escola", "Escola"]))
                        ineps_encontrados.append({"inep": inep_str, "escola": esc_nome, "linha": linha})
                    else:
                        ineps_nao_encontrados.append(inep_str)
            else:
                # === Fallback: busca em lote com isin() — muito mais rápido que loop ===
                modal.update_message(f"Indexando planilha e buscando {total} INEPs...\nAguarde.")
                col = self.col_inep or self.df.columns[0]
                df_norm = self.df.copy()
                df_norm[col] = df_norm[col].astype(str).str.strip().str.split('.').str[0]
                codigos_str = [str(c) for c in codigos_novos]
                resultados = df_norm[df_norm[col].isin(codigos_str)]
                encontrados_set = set(resultados[col].tolist())
                
                for inep in codigos_novos:
                    inep_str = str(inep)
                    if inep_str in encontrados_set:
                        linha = resultados[resultados[col] == inep_str].iloc[0]
                        esc_nome = str(self.localizar(linha, ["Nome da Escola", "Escola"]))
                        ineps_encontrados.append({"inep": inep_str, "escola": esc_nome, "linha": linha})
                    else:
                        ineps_nao_encontrados.append(inep_str)
            
            # Fechar modal e continuar na UI thread
            self.frame.after(0, modal.close)
            self.frame.after(50, lambda: self._pos_busca_ineps(ineps_encontrados, ineps_nao_encontrados))
            
        except Exception as e:
            self.frame.after(0, modal.close)
            self.frame.after(100, lambda: show_detailed_error(
                "Erro de Busca", "Ocorreu um erro ao pesquisar os INEPs na planilha.", e
            ))

    def _pos_busca_ineps(self, ineps_encontrados, ineps_nao_encontrados):
        """Chamado na UI thread após a busca concluir em segundo plano."""
        if ineps_nao_encontrados:
            lista = ', '.join(ineps_nao_encontrados[:15])
            if len(ineps_nao_encontrados) > 15:
                lista += f"\n...e mais {len(ineps_nao_encontrados) - 15} não encontrados."
            messagebox.showwarning(
                "INEPs não encontrados",
                f"Os seguintes INEPs não foram localizados na planilha:\n\n{lista}"
            )
        
        if not ineps_encontrados:
            return
        
        # Abrir popup de configuração de dados
        if len(ineps_encontrados) == 1:
            item = ineps_encontrados[0]
            val, data_raw = self.pedir_dados_escola(item["inep"])
            if val is None: return
            res_list = [{"inep": item["inep"], "linha": item["linha"], "mensalidade": val, "data_inst": data_raw}]
        else:
            res_list = self.pedir_dados_lote(ineps_encontrados)
            if not res_list: return
        
        # Adicionar resultados no Treeview
        for r in res_list:
            linha = r["linha"]
            lat = self.localizar(linha, ["Latitude", "Lat"])
            lon = self.localizar(linha, ["Longitude", "Long"])
            data_i = self.formatar_data_input(r["data_inst"])
            esc = {
                "inep": r["inep"],
                "endereco": str(self.localizar(linha, ["Endereço", "Endereco"])),
                "latlong": f"{lat} / {lon}",
                "megas": f"{self.localizar(linha, ['Velocidade DL Mínima'])} MEGAS",
                "mensalidade": r["mensalidade"],
                "data_inst": data_i
            }
            self.escolas.append(esc)
            self.tree.insert("", "end", values=(
                esc["inep"], esc["endereco"], esc["latlong"],
                esc["megas"], f"R$ {esc['mensalidade']:.2f}", esc["data_inst"]
            ))
        
        self.txt_ineps.delete("1.0", "end")

    def substituir_texto(self, doc, dicionario):
        for tag, entry in dicionario.items():
            valor = entry.get()
            marcador = f"{{{{{tag}}}}}"
            for p in doc.paragraphs:
                if marcador in p.text: p.text = p.text.replace(marcador, valor)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        if marcador in c.text: c.text = c.text.replace(marcador, valor)

    def gerar_assincrono(self):
        if not self.escolas:
            messagebox.showerror("Erro", "Por favor, adicione pelo menos uma escola com INEP.")
            return
            
        dialog = tk.Toplevel(self.root)
        dialog.title("Escolha o Modelo de Contrato")
        dialog.geometry("380x180")
        dialog.transient(self.root)
        dialog.grab_set()
        
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 190
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 90
        dialog.geometry(f"+{x}+{y}")
        
        colors = ThemeManager.get_colors()
        dialog.configure(bg=colors["bg"])
        
        tk.Label(dialog, text="Qual modelo de contrato deseja usar?", font=("Segoe UI", 11, "bold"), bg=colors["bg"], fg=colors["fg"]).pack(pady=15)
        
        f_botoes = tk.Frame(dialog, bg=colors["bg"])
        f_botoes.pack(pady=10)
        
        def usar_modelo(tipo):
            base_path = get_base_path()
            if tipo == "BITNET":
                nome_arq = "contrato modeloBIT.docx"
            else:
                nome_arq = "contrato modeloST1.docx"
                
            caminho_modelo = os.path.join(base_path, nome_arq)
            
            # Tentar achar na pasta atual
            if os.path.exists(caminho_modelo):
                self.template = caminho_modelo
            else:
                # Tentar alternativas case-insensitive na pasta
                encontrado = False
                for f in os.listdir(base_path):
                    if tipo.lower() in f.lower() and f.lower().endswith(".docx") and not f.startswith("~$"):
                        self.template = os.path.join(base_path, f)
                        encontrado = True
                        break
                if not encontrado:
                    messagebox.showinfo("Aviso", f"O arquivo de template '{nome_arq}' não foi encontrado automaticamente na pasta. Por favor, selecione-o manualmente.")
                    arq = filedialog.askopenfilename(filetypes=[("DOCX","*.docx")], title=f"Selecionar Template para {tipo}")
                    if arq:
                        self.template = arq
                    else:
                        dialog.destroy()
                        return
            
            dialog.destroy()
            self._disparar_geracao()
            
        tk.Button(f_botoes, text="Contrato BITNET", command=lambda: usar_modelo("BITNET"), bg="#6366f1", fg="white", font=("Segoe UI", 10, "bold"), width=16).pack(side="left", padx=10)
        tk.Button(f_botoes, text="Contrato ST1", command=lambda: usar_modelo("ST1"), bg="#e67e22", fg="white", font=("Segoe UI", 10, "bold"), width=16).pack(side="left", padx=10)
        
        ThemeManager.apply_theme_to_all(dialog)
        self.root.wait_window(dialog)

    def _disparar_geracao(self):
        if not self.template:
            messagebox.showerror("Erro", "Nenhum arquivo de template foi carregado ou selecionado.")
            return
        modal = ProcessingModal(self.root, "Aguarde", "Gerando contrato... Por favor, aguarde.")
        threading.Thread(target=self._gerar, args=(modal,), daemon=True).start()

    def _gerar(self, modal):
        try:
            doc = Document(self.template)
            
            self.substituir_texto(doc, self.contratada)
            self.substituir_texto(doc, self.sup)
            self.substituir_texto(doc, self.ass)
            
            target = None
            for i, p in enumerate(doc.paragraphs):
                if "Endereço de instalação do serviço" in p.text:
                    if i < len(doc.tables): target = doc.tables[i]
            if not target:
                for t in doc.tables:
                    if "inep" in t.rows[0].cells[0].text.lower(): target = t
            
            if target:
                while len(target.rows) > 1: 
                    target._tbl.remove(target.rows[1]._tr)
                
                total = 0
                for esc in self.escolas:
                    row = target.add_row().cells
                    dados = [esc["inep"], esc["endereco"], esc["latlong"], esc["megas"], f"R$ {esc['mensalidade']:.2f}", esc["data_inst"]]
                    for i in range(min(len(row), len(dados))): row[i].text = str(dados[i])
                    total += esc["mensalidade"]
                
                for t in doc.tables:
                    for r in t.rows:
                        for c in r.cells:
                            if "Total:" in c.text: c.text = f"Total: MENSAL R$ {total:.2f}"
                for p in doc.paragraphs:
                    if "Total:" in p.text: p.text = p.text.replace("Total:", f"Total: MENSAL R$ {total:.2f}")

            # Fechar modal antes do filedialog (evita congelar thread no diálogo)
            self.frame.after(0, modal.close)
            
            def salvar_arquivo():
                salvar = filedialog.asksaveasfilename(defaultextension=".docx", initialfile="Contrato_Finalizado.docx", filetypes=[("DOCX Document", "*.docx")])
                if salvar:
                    try:
                        doc.save(salvar)
                        messagebox.showinfo("Sucesso", "Contrato gerado e salvo com sucesso!")
                    except Exception as ex:
                        messagebox.showerror("Erro de Gravação", f"Não foi possível salvar o contrato:\n{ex}")
            
            self.frame.after(100, salvar_arquivo)
            
        except Exception as e:
            self.frame.after(0, modal.close)
            self.frame.after(100, lambda: show_detailed_error("Erro de Geração", "Ocorreu um erro ao gerar o contrato.", e))

class ConsultaEACE:
    def __init__(self, parent_frame, main_app_callback, gspread_client):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill=tk.BOTH, expand=True)
        self.gspread_client = gspread_client
        self.main_app_callback = main_app_callback
        
        self.df = None
        self.row_atual = None
        self.dl_min_val = None
        self.dl_ofe_val = None
        self.texto_copia = ""

        # Obter cores do ThemeManager
        self.colors = ThemeManager.get_colors()

        self._build()
        self.frame.after(100, self.carregar_so_local)
        ThemeManager.apply_theme_to_all(self.frame)

    def _build(self):
        colors = self.colors
        bg = colors["bg"]
        card = colors["card_bg"]
        border = colors["border"]
        fg = colors["fg"]
        muted = colors["fg_muted"]
        accent = colors["primary"]
        entry_bg = colors["entry_bg"]
        entry_fg = colors["entry_fg"]

        # Configuração do frame principal
        self.frame.configure(bg=bg)

        # ── TOP BAR ──────────────────────────────────────────
        top = tk.Frame(self.frame, bg=card, height=48)
        top.pack(fill="x")
        top.pack_propagate(False)

        tk.Label(top, text="  🔍  CONSULTA RÁPIDA — EACE",
                 font=("Segoe UI", 12, "bold"), bg=card, fg="blue").pack(side="left", padx=8)

        # Botão Voltar
        tk.Button(top, text="VOLTAR AO MENU", command=self.main_app_callback,
                  bg="#6c757d", fg="white", font=("Segoe UI", 8, "bold"),
                  bd=0, cursor="hand2", padx=10).pack(side="right", padx=10, pady=8)

        # ── BARRA DE BUSCA E AÇÕES ───────────────────────────
        bar = tk.Frame(self.frame, bg=bg, pady=10)
        bar.pack(fill="x", padx=16)

        self.lbl_status = tk.Label(bar, text="Carregando cache...", font=("Segoe UI", 9, "italic"), bg=bg, fg="orange")
        self.lbl_status.pack(side="top", anchor="w", pady=(0, 5))

        search_f = tk.Frame(bar, bg=bg)
        search_f.pack(fill="x")

        tk.Label(search_f, text="Código INEP ou Nome:", bg=bg, fg=fg,
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 6))

        self.ent_busca = tk.Entry(search_f, bg=entry_bg, fg=entry_fg, insertbackground=fg,
                                  relief="solid", bd=1, font=("Segoe UI", 11), width=30)
        self.ent_busca.pack(side="left", ipady=4, padx=(0, 8))
        self.ent_busca.bind("<Return>", lambda e: self.buscar_dados())

        # Botão BUSCAR usando TK normal para compatibilidade
        tk.Button(search_f, text="  BUSCAR DADOS  ", command=self.buscar_dados,
                  bg="#007bff", fg="white", font=("Segoe UI", 9, "bold"),
                  bd=0, cursor="hand2", relief="flat").pack(side="left", ipady=4, padx=(0, 8))

        # Botão Sincronizar Nuvem
        tk.Button(search_f, text="  ☁ Atualizar da Nuvem  ", command=self.forcar_atualizacao,
                  bg="#17a2b8", fg="white", font=("Segoe UI", 9, "bold"),
                  bd=0, cursor="hand2", relief="flat").pack(side="left", ipady=4)

        # ── SEPARADOR ────────────────────────────────────────
        tk.Frame(self.frame, bg=border, height=1).pack(fill="x", pady=(5, 5))

        # ── CONTEÚDO PRINCIPAL (2 colunas) ───────────────────
        body = tk.Frame(self.frame, bg=bg)
        body.pack(fill="both", expand=True, padx=16, pady=5)
        body.columnconfigure(0, weight=6)
        body.columnconfigure(1, weight=4)
        body.rowconfigure(0, weight=1)

        # ── COLUNA ESQUERDA: Dados da Escola ─────────────────
        left = tk.Frame(body, bg=card, bd=1, relief="solid", highlightthickness=0)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 8))

        tk.Label(left, text="DADOS DA ESCOLA",
                 font=("Segoe UI", 10, "bold"), bg=card, fg="blue").pack(
                     anchor="w", padx=14, pady=(10, 4))
        tk.Frame(left, bg=border, height=1).pack(fill="x", padx=14)

        # Grid de campos
        self.campos_frame = tk.Frame(left, bg=card)
        self.campos_frame.pack(fill="both", expand=True, padx=14, pady=5)
        self.campos_frame.columnconfigure(1, weight=1)

        self._campo_defs = [
            ("INEP",           "Código INEP",           False),
            ("Escola",         "Nome da Escola",         False),
            ("Município / UF", "__municipio_uf__",       False),
            ("Endereço",       "Endereço",               False),
            ("Latitude",       "Latitude",               False),
            ("Longitude",      "Longitude",              False),
            ("Kit Wi-Fi",      "Kit Wi-Fi",              True),
            ("DL Mínima",      "Velocidade DL Mínima (Mbps)", True),
            ("DL Ofertada",    "Velocidade DL Ofertada (Mbps)", False),
            ("Status",         "Status",                 False),
        ]

        self._vars = {}
        for i, (label, col, destaque) in enumerate(self._campo_defs):
            tk.Label(self.campos_frame, text=label + ":",
                     font=("Segoe UI", 9, "bold" if destaque else "normal"), bg=card, fg=muted,
                     anchor="w", width=14).grid(row=i, column=0, sticky="w", pady=4)

            var = tk.StringVar(value="—")
            cor_val = "orange" if destaque else fg
            peso = "bold" if destaque else "normal"
            lbl = tk.Label(self.campos_frame, textvariable=var,
                           font=("Segoe UI", 9, peso), bg=card, fg=cor_val,
                           anchor="w", wraplength=420, justify="left")
            lbl.grid(row=i, column=1, sticky="w", pady=4, padx=(8, 0))
            self._vars[col] = var

        # Botão copiar
        frame_btn_copy = tk.Frame(left, bg=card)
        frame_btn_copy.pack(fill="x", padx=14, pady=(0, 12))
        self.btn_copiar = tk.Button(frame_btn_copy, text="📋  Copiar Ficha", command=self.copiar_ficha,
                                    bg="#28a745", fg="white", font=("Segoe UI", 9, "bold"),
                                    bd=0, cursor="hand2", state=tk.DISABLED)
        self.btn_copiar.pack(side="left", ipady=4, ipadx=8)

        # ── COLUNA DIREITA: Verificador ───────────────────────
        right = tk.Frame(body, bg=card, bd=1, relief="solid", highlightthickness=0)
        right.grid(row=0, column=1, sticky="nsew")

        tk.Label(right, text="⚡  VERIFICADOR DE VELOCIDADE",
                 font=("Segoe UI", 10, "bold"), bg=card, fg="orange").pack(
                     anchor="w", padx=14, pady=(10, 4))
        tk.Frame(right, bg=border, height=1).pack(fill="x", padx=14)

        inner = tk.Frame(right, bg=card)
        inner.pack(fill="both", expand=True, padx=14, pady=8)

        # Info de velocidade
        for label, attr in [("Velocidade Ofertada:", "lbl_contratada"),
                            ("Mínima Requerida (Planilha):", "lbl_threshold")]:
            row_f = tk.Frame(inner, bg=card)
            row_f.pack(fill="x", pady=2)
            tk.Label(row_f, text=label, font=("Segoe UI", 9), bg=card, fg=muted,
                     width=25, anchor="w").pack(side="left")
            lbl = tk.Label(row_f, text="—", font=("Segoe UI", 9, "bold"),
                           bg=card, fg=fg, anchor="w")
            lbl.pack(side="left")
            setattr(self, attr, lbl)

        tk.Frame(inner, bg=border, height=1).pack(fill="x", pady=8)

        # Input speedtest
        tk.Label(inner, text="Resultado do Speedtest:", font=("Segoe UI", 9, "bold"),
                 bg=card, fg=fg, anchor="w").pack(fill="x")

        speed_row = tk.Frame(inner, bg=card)
        speed_row.pack(fill="x", pady=(4, 0))

        self.ent_speed = tk.Entry(speed_row, bg=entry_bg, fg=entry_fg, insertbackground=fg,
                                  relief="solid", bd=1, font=("Segoe UI", 20, "bold"),
                                  width=8, justify="center")
        self.ent_speed.pack(side="left", ipady=4)
        self.ent_speed.bind("<Return>", lambda e: self.verificar_velocidade())

        tk.Label(speed_row, text=" Mbps", font=("Segoe UI", 11), bg=card,
                 fg=muted).pack(side="left", padx=5)

        tk.Button(inner, text="  CALCULAR  ", command=self.verificar_velocidade,
                  bg="#6366f1", fg="white", font=("Segoe UI", 9, "bold"),
                  bd=0, cursor="hand2").pack(fill="x", pady=(8, 0), ipady=6)

        # Área de resultado
        tk.Frame(inner, bg=border, height=1).pack(fill="x", pady=(10, 6))

        self.frame_res = tk.Frame(inner, bg=card)
        self.frame_res.pack(fill="both", expand=True)

        self.lbl_res_icone = tk.Label(self.frame_res, text="",
                                      font=("Segoe UI", 24, "bold"), bg=card)
        self.lbl_res_icone.pack()

        self.lbl_res_status = tk.Label(self.frame_res, text="",
                                       font=("Segoe UI", 12, "bold"), bg=card)
        self.lbl_res_status.pack()

        self.lbl_res_pct = tk.Label(self.frame_res, text="",
                                    font=("Segoe UI", 9), bg=card, fg=muted)
        self.lbl_res_pct.pack(pady=(2, 6))

        # Mini barra de progresso
        self.canvas_barra = tk.Canvas(self.frame_res, height=16, bg=entry_bg,
                                      bd=0, highlightthickness=0)
        self.canvas_barra.pack(fill="x", padx=2)

        # Detalhes resultado
        self.frame_det = tk.Frame(self.frame_res, bg=card)
        self.frame_det.pack(fill="x", pady=(8, 0))

        self._det_vars = {}
        for label in ["Medido", "Minimo", "Diferenca"]:
            r = tk.Frame(self.frame_det, bg=card)
            r.pack(fill="x", pady=1)
            tk.Label(r, text=label + ":", font=("Segoe UI", 9), bg=card,
                     fg=muted, width=12, anchor="w").pack(side="left")
            v = tk.StringVar(value="")
            lbl = tk.Label(r, textvariable=v, font=("Segoe UI", 9, "bold"),
                           bg=card, fg=fg, anchor="w")
            lbl.pack(side="left")
            self._det_vars[label] = (v, lbl)

        # Nota rodapé
        self.lbl_nota = tk.Label(right,
            text="Busque uma escola e insira o resultado do Speedtest.",
            font=("Segoe UI", 8, "italic"), bg=card, fg=muted, wraplength=260)
        self.lbl_nota.pack(padx=14, pady=(0, 10))

    def carregar_so_local(self):
        self.df = load_cache()
        if not self.df.empty:
            col = "Código INEP" if "Código INEP" in self.df.columns else self.df.columns[0]
            self.df[col] = self.df[col].astype(str).str.split('.').str[0].str.strip()
            self.lbl_status.config(text="✅ Cache local carregado — consultas prontas.", fg="green")
            self.ent_busca.focus()
        else:
            self.lbl_status.config(text="⚠️ Base local vazia. Por favor, clique em 'Atualizar da Nuvem'.", fg="orange")

    def forcar_atualizacao(self):
        self.lbl_status.config(text="⏳ Sincronizando com a nuvem em background...", fg="blue")
        threading.Thread(target=self._sincronizar_nuvem, daemon=True).start()

    def _sincronizar_nuvem(self):
        try:
            if not self.gspread_client:
                self.gspread_client = authenticate_google_sheets()
            if not self.gspread_client:
                raise Exception("Sem credenciais válidas do Google.")
                
            df_nuvem = fetch_data_by_id(self.gspread_client, SPREADSHEET_ID, "EACE")
            if not df_nuvem.empty:
                col = "Código INEP" if "Código INEP" in df_nuvem.columns else df_nuvem.columns[0]
                df_nuvem[col] = df_nuvem[col].astype(str).str.split('.').str[0].str.strip()
                self.df = df_nuvem
                save_cache(df_nuvem)
                self.frame.after(0, lambda: self.lbl_status.config(text="✅ BASE ATUALIZADA COM SUCESSO!", fg="green"))
                self.frame.after(0, lambda: self.ent_busca.focus())
        except Exception as e:
            if self.df is not None and not self.df.empty:
                self.frame.after(0, lambda: self.lbl_status.config(text="⚠️ Falha na rede. Operando com cache local anterior.", fg="orange"))
            else:
                self.frame.after(0, lambda: self.lbl_status.config(text="❌ ERRO DE CONEXÃO. Sem cache e sem rede.", fg="red"))
                self.frame.after(0, lambda: show_detailed_error("Erro de Sincronização", "Não foi possível acessar a planilha em nuvem.", e))

    def buscar_dados(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Aviso", "Não há base de dados carregada para busca.")
            return
            
        termo = self.ent_busca.get().strip()
        if not termo: return
        
        self._limpar_escola()
        
        col_inep = "Código INEP" if "Código INEP" in self.df.columns else self.df.columns[0]
        
        # Verificar se é busca exata de INEP (apenas dígitos)
        if termo.isdigit() and len(termo) >= 8:
            res = self.df[self.df[col_inep] == termo]
            if res.empty:
                messagebox.showinfo("Não Encontrado", f"Código INEP '{termo}' não encontrado.")
            else:
                self.exibir_escola(res.iloc[0])
        else:
            # Busca textual parcial no Nome da Escola
            col_nome = "Nome da Escola" if "Nome da Escola" in self.df.columns else ""
            if not col_nome:
                for c in self.df.columns:
                    if "escola" in c.lower() or "nome" in c.lower():
                        col_nome = c
                        break
            
            if not col_nome:
                messagebox.showerror("Erro", "Coluna de nome da escola não encontrada na base.")
                return
                
            res = self.df[self.df[col_nome].astype(str).str.contains(termo, case=False, na=False)]
            
            if res.empty:
                messagebox.showinfo("Não Encontrado", f"Nenhuma escola localizada para '{termo}'.")
            elif len(res) == 1:
                self.exibir_escola(res.iloc[0])
            else:
                self.selecionar_escola_popup(res, col_nome)

    def selecionar_escola_popup(self, df_resultados, col_nome):
        popup = tk.Toplevel(self.frame)
        popup.title("Múltiplos Resultados")
        popup.geometry("550x340")
        popup.transient(self.frame)
        popup.grab_set()
        
        popup.update_idletasks()
        x = self.frame.winfo_toplevel().winfo_x() + (self.frame.winfo_toplevel().winfo_width() // 2) - 275
        y = self.frame.winfo_toplevel().winfo_y() + (self.frame.winfo_toplevel().winfo_height() // 2) - 170
        popup.geometry(f"+{x}+{y}")
        
        colors = ThemeManager.get_colors()
        popup.configure(bg=colors["bg"])
        
        tk.Label(popup, text="Selecione a Escola na lista abaixo:", font=("Segoe UI", 10, "bold"), bg=colors["bg"], fg=colors["fg"]).pack(pady=10)
        
        list_frame = tk.Frame(popup, bg=colors["bg"])
        list_frame.pack(fill="both", expand=True, padx=15, pady=5)
        
        scy = ttk.Scrollbar(list_frame)
        scy.pack(side="right", fill="y")
        
        lb = tk.Listbox(list_frame, font=("Segoe UI", 10), yscrollcommand=scy.set, bg=colors["entry_bg"], fg=colors["entry_fg"])
        lb.pack(fill="both", expand=True)
        scy.config(command=lb.yview)
        
        col_inep = "Código INEP" if "Código INEP" in df_resultados.columns else df_resultados.columns[0]
        
        # Preencher a listbox
        escolas_list = []
        for index, row in df_resultados.iterrows():
            txt = f"{row[col_inep]} - {row[col_nome]}"
            lb.insert(tk.END, txt)
            escolas_list.append(row)
            
        def confirmar():
            sel = lb.curselection()
            if sel:
                row_selecionada = escolas_list[sel[0]]
                self.exibir_escola(row_selecionada)
                popup.destroy()
                
        def cancelar():
            popup.destroy()
            
        btn_frame = ttk.Frame(popup)
        btn_frame.pack(pady=15)
        ttk.Button(btn_frame, text="Visualizar Escola", command=confirmar).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="Cancelar", command=cancelar).pack(side="left", padx=10)
        
        ThemeManager.apply_theme_to_all(popup)

    def exibir_escola(self, row):
        self.row_atual = row
        col_inep = "Código INEP" if "Código INEP" in row.index else row.index[0]
        col_nome = "Nome da Escola" if "Nome da Escola" in row.index else ""
        if not col_nome:
            for c in row.index:
                if "escola" in c.lower() or "nome" in c.lower():
                    col_nome = c
                    break

        mun = row.get("Município", "N/A")
        uf  = row.get("UF", "N/A")

        # Ficha formatada para cópia tradicional
        self.texto_copia = (
            f"============================================================\n"
            f"             FICHA DE INFORMAÇÕES DA ESCOLA - EACE          \n"
            f"============================================================\n"
            f"INEP:         {row.get(col_inep, 'N/A')}\n"
            f"ESCOLA:       {row.get(col_nome, 'N/A')}\n"
            f"MUNICÍPIO:    {row.get('Município', 'N/A')}\n"
            f"UF:           {row.get('UF', 'N/A')}\n"
            f"ENDEREÇO:     {row.get('Endereço', 'N/A')}\n"
            f"LATITUDE:     {row.get('Latitude', 'N/A')}\n"
            f"LONGITUDE:    {row.get('Longitude', 'N/A')}\n"
            f"KIT PREVISTO: {row.get('Kit Wi-Fi', 'N/A')}\n"
            f"------------------------------------------------------------\n"
            f"DL MÍNIMA:    {row.get('Velocidade DL Mínima (Mbps)', 'N/A')} Mbps\n"
            f"DL OFERTADA:  {row.get('Velocidade DL Ofertada (Mbps)', 'N/A')} Mbps\n"
            f"STATUS:       {row.get('Status', 'Ativo')}\n"
            f"============================================================\n"
        )

        mapping = {
            "Código INEP":    row.get(col_inep, "N/A"),
            "Nome da Escola": row.get(col_nome, "N/A"),
            "__municipio_uf__": f"{mun} / {uf}",
            "Endereço":       row.get("Endereço", "N/A"),
            "Latitude":       row.get("Latitude", "N/A"),
            "Longitude":      row.get("Longitude", "N/A"),
            "Kit Wi-Fi":      row.get("Kit Wi-Fi", "N/A"),
            "Velocidade DL Mínima (Mbps)": str(row.get("Velocidade DL Mínima (Mbps)", "N/A")) + " Mbps",
            "Velocidade DL Ofertada (Mbps)": str(row.get("Velocidade DL Ofertada (Mbps)", "N/A")) + " Mbps",
            "Status":         row.get("Status", "Ativo"),
        }

        for col, var in self._vars.items():
            var.set(mapping.get(col, "N/A"))

        self.btn_copiar.config(state=tk.NORMAL)

        # Atualizar painel de velocidade
        dl_raw = row.get("Velocidade DL Mínima (Mbps)", "N/A")
        
        def extrair_mbps(v):
            try:
                nums = re.findall(r'[\d.,]+', str(v).replace(',', '.'))
                return float(nums[0]) if nums else None
            except Exception:
                return None
                
        self.dl_min_val = extrair_mbps(dl_raw)
        self.dl_ofe_val = extrair_mbps(row.get("Velocidade DL Ofertada (Mbps)", "N/A"))

        fg = self.colors["fg"]
        if self.dl_ofe_val and self.dl_ofe_val > 0:
            threshold_disp = self.dl_ofe_val * 0.8
            self.lbl_contratada.config(text=f"{self.dl_ofe_val:.1f} Mbps", fg=fg)
            self.lbl_threshold.config(text=f"{threshold_disp:.1f} Mbps", fg="orange")
            self.lbl_nota.config(text="Insira o resultado do Speedtest e clique em CALCULAR.")
            self.ent_speed.focus()
        elif self.dl_min_val is not None:
            self.lbl_contratada.config(text="N/A", fg=fg)
            self.lbl_threshold.config(text=f"{self.dl_min_val:.1f} Mbps", fg="orange")
            self.lbl_nota.config(text="Insira o resultado do Speedtest e clique em CALCULAR.")
            self.ent_speed.focus()
        else:
            self.lbl_contratada.config(text="N/A", fg="orange")
            self.lbl_threshold.config(text=str(dl_raw) + " (não numérico)", fg="orange")
            self.lbl_nota.config(text="Velocidade não numérica — verificação manual necessária.")

        self._limpar_resultado()

    def verificar_velocidade(self):
        if self.row_atual is None:
            messagebox.showwarning("Aviso", "Busque uma escola primeiro.")
            return
        if self.dl_ofe_val and self.dl_ofe_val > 0:
            threshold = self.dl_ofe_val * 0.8
        elif self.dl_min_val is not None:
            threshold = self.dl_min_val
        else:
            messagebox.showwarning("Aviso", "Velocidade não numérica. Verifique manualmente.")
            return

        raw = self.ent_speed.get().strip().replace(',', '.')
        if not raw:
            messagebox.showwarning("Aviso", "Digite o resultado do Speedtest.")
            return
        try:
            medido = float(raw)
        except ValueError:
            messagebox.showerror("Erro", "Valor inválido. Use números (ex: 45.3).")
            return

        aprovado   = medido >= threshold
        diferenca  = medido - threshold
        
        if self.dl_ofe_val and self.dl_ofe_val > 0:
            pct = (medido / self.dl_ofe_val) * 100
        else:
            pct = (medido / threshold) * 100

        self._mostrar_resultado(medido, threshold, pct, aprovado, diferenca)

    def _mostrar_resultado(self, medido, threshold, pct, aprovado, diferenca):
        cor    = "green" if aprovado else "red"
        status = "APROVADO" if aprovado else "REPROVADO"
        icone  = "✅" if aprovado else "❌"

        self.lbl_res_icone.config(text=icone, fg=cor)
        self.lbl_res_status.config(text=status, fg=cor)
        
        if self.dl_ofe_val and self.dl_ofe_val > 0:
            self.lbl_res_pct.config(text=f"{pct:.1f}% da velocidade ofertada")
        else:
            self.lbl_res_pct.config(text=f"{pct:.1f}% da velocidade mínima")

        # Barra de progresso
        self.canvas_barra.update_idletasks()
        w = self.canvas_barra.winfo_width() or 240
        self.canvas_barra.delete("all")
        
        entry_bg = self.colors["entry_bg"]
        self.canvas_barra.create_rectangle(0, 0, w, 16, fill=entry_bg, outline="")
        
        fill_w = int(w * min(pct, 100) / 100)
        self.canvas_barra.create_rectangle(0, 0, fill_w, 16, fill=cor, outline="")
        
        if self.dl_ofe_val and self.dl_ofe_val > 0:
            threshold_pct = (self.dl_min_val / self.dl_ofe_val) * 100
        else:
            threshold_pct = 100
            
        x_thresh = int(w * threshold_pct / 100)
        self.canvas_barra.create_line(x_thresh, 0, x_thresh, 16, fill="orange", width=2, dash=(3, 2))
        self.canvas_barra.create_text(w // 2, 8, text=f"{pct:.1f}%",
                                      fill="white", font=("Segoe UI", 8, "bold"))

        # Detalhes
        sinal = "+" if diferenca >= 0 else ""
        vals = {
            "Medido":    (f"{medido:.1f} Mbps", self.colors["fg"]),
            "Minimo":    (f"{threshold:.1f} Mbps", "orange"),
            "Diferenca": (f"{sinal}{diferenca:.1f} Mbps",
                           "green" if diferenca >= 0 else "red"),
        }
        for key, (texto, cor_det) in vals.items():
            var, lbl = self._det_vars[key]
            var.set(texto)
            lbl.config(fg=cor_det)

        self.lbl_nota.config(text="")

    def _limpar_escola(self):
        for var in self._vars.values():
            var.set("—")
        self.btn_copiar.config(state=tk.DISABLED)
        self.lbl_contratada.config(text="—", fg=self.colors["fg"])
        self.lbl_threshold.config(text="—", fg=self.colors["fg"])
        self.dl_min_val = None
        self.dl_ofe_val = None
        self.row_atual = None
        self._limpar_resultado()

    def _limpar_resultado(self):
        self.lbl_res_icone.config(text="")
        self.lbl_res_status.config(text="")
        self.lbl_res_pct.config(text="")
        self.canvas_barra.delete("all")
        for var, lbl in self._det_vars.values():
            var.set("")
            lbl.config(fg=self.colors["fg"])

    def copiar_ficha(self):
        if self.texto_copia:
            self.frame.clipboard_clear()
            self.frame.clipboard_append(self.texto_copia)
            messagebox.showinfo("Sucesso", "A ficha de informações foi copiada para a área de transferência!")

class UnificadorApp:
    def __init__(self, parent_frame, main_app_callback, gspread_client=None):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill=tk.BOTH, expand=True)
        self.main_app_callback = main_app_callback
        self.root = parent_frame.winfo_toplevel()
        
        config = load_config()
        self.target_dir = config.get("unificador_dest")
        self.selected_files = []
        self.img_preview_ref = None # Evitar coleta de lixo da imagem do Tkinter

        tk.Button(self.frame, text="VOLTAR AO MENU PRINCIPAL", command=self.main_app_callback, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=(5, 10))
        
        main_frame = tk.Frame(self.frame, padx=15, pady=5)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(main_frame, text="Unificador de Evidências Multiformato", font=("Segoe UI", 18, "bold"), fg="#2c3e50").pack(pady=(0, 2))
        tk.Label(main_frame, text="Bitnet Telecom - Suporte a Imagens, TXT, PDF e DOCX", font=("Segoe UI", 9, "bold"), fg="#27ae60").pack(pady=(0, 10))
        
        # Destino
        path_frame = tk.LabelFrame(main_frame, text="Pasta de Destino do PDF Unificado", padx=10, pady=5, font=("Segoe UI", 9, "bold"))
        path_frame.pack(fill=tk.X, pady=5)
        self.lbl_dest = tk.Label(path_frame, text=self.target_dir, wraplength=550, justify=tk.LEFT, fg="#2980b9")
        self.lbl_dest.pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Button(path_frame, text="Alterar", command=self.mudar_destino, font=("Segoe UI", 8, "bold")).pack(side=tk.RIGHT)

        # Container Principal de Lista e Preview
        layout_container = tk.Frame(main_frame)
        layout_container.pack(fill=tk.BOTH, expand=True, pady=5)

        # Lado Esquerdo (Lista de Arquivos)
        files_frame = tk.LabelFrame(layout_container, text="Lista de Arquivos Selecionados", padx=10, pady=5, font=("Segoe UI", 9, "bold"))
        files_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        list_container = tk.Frame(files_frame)
        list_container.pack(fill=tk.BOTH, expand=True)
        
        sc_y = ttk.Scrollbar(list_container)
        sc_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.lb = tk.Listbox(list_container, height=10, font=("Consolas", 9), yscrollcommand=sc_y.set)
        self.lb.pack(fill=tk.BOTH, expand=True)
        sc_y.config(command=self.lb.yview)
        
        self.lb.bind("<<ListboxSelect>>", self.mostrar_preview)
        
        # Botões de Ordenação e Ação de Arquivos
        btn_files_frame = tk.Frame(files_frame)
        btn_files_frame.pack(fill="x", pady=5)
        
        tk.Button(btn_files_frame, text="+ Adicionar", command=self.select_files, bg="#3498db", fg="white", font=("Segoe UI", 9, "bold"), width=12).grid(row=0, column=0, padx=2)
        tk.Button(btn_files_frame, text="⬆️ Subir", command=self.subir_arquivo, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold"), width=9).grid(row=0, column=1, padx=2)
        tk.Button(btn_files_frame, text="⬇️ Descer", command=self.descer_arquivo, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold"), width=9).grid(row=0, column=2, padx=2)
        tk.Button(btn_files_frame, text="Limpar Tudo", command=self.limpar, bg="#e74c3c", fg="white", font=("Segoe UI", 9, "bold"), width=12).grid(row=0, column=3, padx=2)

        # Lado Direito (Área de Pré-visualização)
        self.preview_frame = tk.LabelFrame(layout_container, text="Pré-visualização da Evidência", padx=10, pady=5, font=("Segoe UI", 9, "bold"), width=280)
        self.preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(5, 0))
        self.preview_frame.pack_propagate(False) # Impedir encolhimento do frame de preview

        self.preview_canvas = tk.Canvas(self.preview_frame, bg="#eaeaea", highlightthickness=1)
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        
        self.preview_text = tk.Text(self.preview_frame, font=("Segoe UI", 9), wrap=tk.WORD, state=tk.DISABLED)
        # O text fica oculto por padrão, usaremos pack de acordo com a necessidade

        # Opções Extras e Nome
        opt_frame = tk.Frame(main_frame)
        opt_frame.pack(fill=tk.X, pady=5)

        self.var_otimizar = tk.BooleanVar(value=True)
        self.chk_otimizar = ttk.Checkbutton(opt_frame, text="Otimizar tamanho das imagens (Reduzir tamanho final do PDF)", variable=self.var_otimizar)
        self.chk_otimizar.pack(side=tk.LEFT)

        name_frame = tk.Frame(main_frame, pady=5)
        name_frame.pack(fill=tk.X)
        tk.Label(name_frame, text="Nome do PDF Final:", font=("Segoe UI", 10, "bold")).pack(side=tk.LEFT)
        self.ne = ttk.Entry(name_frame, font=("Segoe UI", 11))
        self.ne.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))

        tk.Button(main_frame, text="⚡ GERAR PDF UNIFICADO AGORA ⚡", bg="#27ae60", fg="white", font=("Segoe UI", 12, "bold"), command=self.processar_assincrono, height=2).pack(fill=tk.X, pady=5)
        
        footer_frame = tk.Frame(self.frame, padx=10, pady=2)
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X)
        tk.Label(footer_frame, text="BITNET TELECOM - Unificador de Evidências", font=("Segoe UI", 8, "italic"), fg="#bdc3c7").pack(side=tk.RIGHT, pady=5)
        
        ThemeManager.apply_theme_to_all(self.frame)

    def mudar_destino(self):
        new_dir = filedialog.askdirectory(initialdir=self.target_dir)
        if new_dir:
            self.target_dir = new_dir
            self.lbl_dest.config(text=self.target_dir)
            save_config("unificador_dest", self.target_dir)

    def select_files(self):
        fs = filedialog.askopenfilenames(filetypes=[
            ("Arquivos de Evidências", "*.png;*.jpg;*.jpeg;*.txt;*.pdf;*.docx"),
            ("Imagens", "*.png;*.jpg;*.jpeg"),
            ("Documentos de Texto", "*.txt"),
            ("Arquivos PDF", "*.pdf"),
            ("Documentos Word", "*.docx"),
            ("Todos os arquivos", "*.*")
        ])
        for f in fs:
            if f not in self.selected_files:
                self.selected_files.append(f)
                self.lb.insert(tk.END, os.path.basename(f))

    def limpar(self):
        self.selected_files = []
        self.lb.delete(0, tk.END)
        self.limpar_preview()

    def limpar_preview(self):
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        self.preview_text.pack_forget()
        self.preview_canvas.delete("all")
        self.img_preview_ref = None

    def subir_arquivo(self):
        sel = self.lb.curselection()
        if not sel: return
        idx = sel[0]
        if idx == 0: return # Já está no topo
        
        # Trocar na lista interna
        self.selected_files[idx], self.selected_files[idx-1] = self.selected_files[idx-1], self.selected_files[idx]
        
        # Atualizar a Listbox
        txt_ativo = self.lb.get(idx)
        txt_vizinho = self.lb.get(idx-1)
        self.lb.delete(idx)
        self.lb.delete(idx-1)
        self.lb.insert(idx-1, txt_ativo)
        self.lb.insert(idx, txt_vizinho)
        self.lb.selection_set(idx-1)
        self.mostrar_preview()

    def descer_arquivo(self):
        sel = self.lb.curselection()
        if not sel: return
        idx = sel[0]
        if idx == len(self.selected_files) - 1: return # Já está no final
        
        # Trocar na lista interna
        self.selected_files[idx], self.selected_files[idx+1] = self.selected_files[idx+1], self.selected_files[idx]
        
        # Atualizar a Listbox
        txt_ativo = self.lb.get(idx)
        txt_vizinho = self.lb.get(idx+1)
        self.lb.delete(idx+1)
        self.lb.delete(idx)
        self.lb.insert(idx, txt_vizinho)
        self.lb.insert(idx+1, txt_ativo)
        self.lb.selection_set(idx+1)
        self.mostrar_preview()

    def mostrar_preview(self, event=None):
        sel = self.lb.curselection()
        if not sel: 
            self.limpar_preview()
            return
            
        path = self.selected_files[sel[0]]
        ext = os.path.splitext(path)[1].lower()
        
        self.preview_canvas.delete("all")
        
        if ext in [".png", ".jpg", ".jpeg"]:
            self.preview_canvas.pack(fill=tk.BOTH, expand=True)
            self.preview_text.pack_forget()
            
            try:
                img = Image.open(path)
                img = ImageOps.exif_transpose(img)
                
                # Proporções do preview canvas
                cw = self.preview_canvas.winfo_width()
                ch = self.preview_canvas.winfo_height()
                if cw < 50: cw = 250
                if ch < 50: ch = 250
                
                img.thumbnail((cw - 10, ch - 10))
                
                self.img_preview_ref = ImageTk.PhotoImage(img)
                self.preview_canvas.create_image(cw // 2, ch // 2, image=self.img_preview_ref, anchor="center")
            except Exception as e:
                self.preview_canvas.create_text(10, 10, text=f"Erro ao ler imagem:\n{e}", anchor="nw", fill="red")
        
        elif ext == ".txt":
            self.preview_canvas.pack_forget()
            self.preview_text.pack(fill=tk.BOTH, expand=True)
            
            self.preview_text.config(state=tk.NORMAL)
            self.preview_text.delete(1.0, tk.END)
            
            try:
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    conteudo = f.read(800) # Mostrar apenas o início
                self.preview_text.insert(tk.END, conteudo + "\n\n[Texto truncado para pré-visualização...]" if len(conteudo) >= 800 else conteudo)
            except Exception as e:
                self.preview_text.insert(tk.END, f"Erro ao abrir arquivo:\n{e}")
                
            self.preview_text.config(state=tk.DISABLED)
            
        elif ext == ".pdf":
            self.preview_canvas.pack(fill=tk.BOTH, expand=True)
            self.preview_text.pack_forget()
            
            cw = self.preview_canvas.winfo_width()
            ch = self.preview_canvas.winfo_height()
            if cw < 50: cw = 250
            if ch < 50: ch = 250
            
            # Desenhar ícone representativo do PDF
            self.preview_canvas.create_rectangle(cw//2 - 50, ch//2 - 60, cw//2 + 50, ch//2 + 60, fill="#dc3545", outline="#b21f2d", width=2)
            self.preview_canvas.create_text(cw//2, ch//2, text="📄 PDF", fill="white", font=("Segoe UI", 16, "bold"), anchor="center")
            self.preview_canvas.create_text(cw//2, ch//2 + 85, text=os.path.basename(path)[:25] + "...", fill="black", font=("Segoe UI", 9), anchor="center")
            
        elif ext == ".docx":
            self.preview_canvas.pack(fill=tk.BOTH, expand=True)
            self.preview_text.pack_forget()
            
            cw = self.preview_canvas.winfo_width()
            ch = self.preview_canvas.winfo_height()
            if cw < 50: cw = 250
            if ch < 50: ch = 250
            
            # Desenhar ícone representativo do DOCX
            self.preview_canvas.create_rectangle(cw//2 - 50, ch//2 - 60, cw//2 + 50, ch//2 + 60, fill="#0056b3", outline="#002d62", width=2)
            self.preview_canvas.create_text(cw//2, ch//2, text="📝 DOCX", fill="white", font=("Segoe UI", 16, "bold"), anchor="center")
            self.preview_canvas.create_text(cw//2, ch//2 + 85, text=os.path.basename(path)[:25] + "...", fill="black", font=("Segoe UI", 9), anchor="center")

    def criar_imagem_texto(self, txt_path):
        with open(txt_path, 'r', encoding='utf-8', errors='replace') as f: 
            text = f.read()
        try: font = ImageFont.truetype("arial.ttf", 20)
        except: font = ImageFont.load_default()
        img = Image.new('RGB', (1240, 1754), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        lines = []
        for p in text.split('\n'): lines.extend(textwrap.wrap(p, width=85))
        y = 50
        for line in lines:
            draw.text((50, y), line, font=font, fill=(0, 0, 0))
            y += 30
        return img

    def processar_assincrono(self):
        nome = self.ne.get().strip()
        if not self.selected_files:
            messagebox.showwarning("Aviso", "Selecione pelo menos um arquivo para unificar.")
            return
        if not nome:
            messagebox.showwarning("Aviso", "Digite o nome da pasta/PDF final.")
            return
            
        modal = ProcessingModal(self.root, "Aguarde", "Unificando evidências e gerando PDF... Por favor, aguarde.")
        threading.Thread(target=self._processar, args=(nome, modal), daemon=True).start()

    def _processar(self, nome, modal):
        temp_dir = tempfile.mkdtemp()
        out_path = None
        try:
            pdf_parts = []
            total = len(self.selected_files)

            # Garantir existência da biblioteca pypdf
            if pypdf is None:
                raise ImportError("A biblioteca 'pypdf' é necessária para mesclar PDFs no Unificador de Evidências.")

            # Validar e criar pasta de destino se não existir
            if not os.path.isdir(self.target_dir):
                os.makedirs(self.target_dir, exist_ok=True)

            for idx, f in enumerate(self.selected_files):
                modal.update_message(
                    f"Processando arquivo {idx + 1} de {total}:\n{os.path.basename(f)}"
                )
                ext = os.path.splitext(f)[1].lower()
                temp_pdf_path = os.path.join(temp_dir, f"part_{idx}.pdf")

                if ext in [".png", ".jpg", ".jpeg"]:
                    img = Image.open(f)
                    img = ImageOps.exif_transpose(img).convert('RGB')
                    if self.var_otimizar.get():
                        max_size = 1920
                        if img.width > max_size or img.height > max_size:
                            img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                    img.save(temp_pdf_path, "PDF", resolution=100.0, quality=80)
                    pdf_parts.append(temp_pdf_path)

                elif ext == ".txt":
                    img_txt = self.criar_imagem_texto(f)
                    img_txt.save(temp_pdf_path, "PDF", resolution=100.0)
                    pdf_parts.append(temp_pdf_path)

                elif ext == ".pdf":
                    shutil.copy2(f, temp_pdf_path)
                    pdf_parts.append(temp_pdf_path)

                elif ext == ".docx":
                    if win32com.client is None and docx_to_pdf_convert is None:
                        raise ImportError("É necessário ter o Microsoft Word instalado ou a biblioteca docx2pdf para importar arquivos .docx.")
                    success_docx = False
                    if win32com.client is not None:
                        try:
                            word = win32com.client.Dispatch("Word.Application")
                            word.Visible = False
                            doc = word.Documents.Open(os.path.abspath(f))
                            doc.SaveAs(os.path.abspath(temp_pdf_path), FileFormat=17)
                            doc.Close()
                            word.Quit()
                            pdf_parts.append(temp_pdf_path)
                            success_docx = True
                        except Exception as ex_word:
                            print(f"Falha DOCX via COM: {ex_word}")
                    if not success_docx and docx_to_pdf_convert is not None:
                        try:
                            docx_to_pdf_convert(f, temp_pdf_path)
                            pdf_parts.append(temp_pdf_path)
                            success_docx = True
                        except Exception as ex_docx:
                            print(f"Falha DOCX via docx2pdf: {ex_docx}")
                    if not success_docx:
                        raise Exception(f"Não foi possível converter o arquivo DOCX:\n{f}")

            if not pdf_parts:
                raise Exception("Nenhuma parte do PDF foi gerada. Verifique os arquivos selecionados.")

            modal.update_message(f"Mesclando {len(pdf_parts)} parte(s) no PDF final...\nAguarde.")

            out_path = os.path.join(self.target_dir, f"{nome}.pdf")

            # Usar PdfWriter + PdfReader (API moderna do pypdf — compatível com 4.x e 5.x)
            writer = pypdf.PdfWriter()
            for part in pdf_parts:
                reader = pypdf.PdfReader(part)
                for page in reader.pages:
                    writer.add_page(page)

            with open(out_path, "wb") as f_out:
                writer.write(f_out)

            # Confirmar que o arquivo foi gerado
            if not os.path.isfile(out_path) or os.path.getsize(out_path) == 0:
                raise Exception(f"O arquivo foi gerado mas parece estar vazio ou corrompido:\n{out_path}")

            self.frame.after(0, modal.close)
            self.frame.after(100, lambda: messagebox.showinfo(
                "Sucesso ✅",
                f"PDF Unificado gerado com sucesso!\n\nSalvo em:\n{out_path}"
            ))
            self.frame.after(200, self.limpar)
            self.frame.after(300, lambda: self.ne.delete(0, tk.END))

        except Exception as e:
            self.frame.after(0, modal.close)
            self.frame.after(100, lambda: show_detailed_error(
                "Erro de Unificação", "Não foi possível gerar o PDF final das evidências.", e
            ))
        finally:
            # Limpar diretório temporário (arquivos intermediários)
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass

class ConversorBidirecionalApp:
    def __init__(self, parent_frame, main_app_callback, gspread_client=None):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill=tk.BOTH, expand=True)
        self.main_app_callback = main_app_callback
        self.root = parent_frame.winfo_toplevel()
        
        config = load_config()
        self.target_dir = config.get("conversor_dest")
        self.active_file_paths = []
        self.sentido_conversao = "pdf2docx" # ou "docx2pdf"

        tk.Button(self.frame, text="VOLTAR AO MENU PRINCIPAL", command=self.main_app_callback, bg="#6c757d", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=(5, 10))
        
        main_frame = tk.Frame(self.frame, padx=20, pady=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(main_frame, text="Conversor Bidirecional de Documentos", font=("Segoe UI", 18, "bold"), fg="#2c3e50").pack(pady=(0, 10))
        
        # Destino
        path_frame = tk.LabelFrame(main_frame, text="Pasta de Salvamento do Arquivo Convertido", padx=10, pady=5, font=("Segoe UI", 9, "bold"))
        path_frame.pack(fill=tk.X, pady=5)
        self.lbl_dest = tk.Label(path_frame, text=self.target_dir, wraplength=500, justify=tk.LEFT, fg="#2980b9")
        self.lbl_dest.pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Button(path_frame, text="Alterar", command=self.mudar_destino, font=("Segoe UI", 8, "bold")).pack(side=tk.RIGHT)

        # Seleção de Modo
        mode_frame = tk.LabelFrame(main_frame, text="Sentido da Conversão", padx=15, pady=10, font=("Segoe UI", 9, "bold"))
        mode_frame.pack(fill=tk.X, pady=10)
        
        self.var_modo = tk.StringVar(value="pdf2docx")
        self.r_pdf2docx = ttk.Radiobutton(mode_frame, text="Convert PDF para Word (.docx)", variable=self.var_modo, value="pdf2docx", command=self.atualizar_modo)
        self.r_pdf2docx.grid(row=0, column=0, padx=15)
        
        self.r_docx2pdf = ttk.Radiobutton(mode_frame, text="Convert Word (.docx) para PDF", variable=self.var_modo, value="docx2pdf", command=self.atualizar_modo)
        self.r_docx2pdf.grid(row=0, column=1, padx=15)

        # Selecionar Arquivos (Listbox)
        self.lbl_info = tk.Label(main_frame, text="Nenhum arquivo selecionado", fg="gray", font=("Segoe UI", 10, "italic"))
        self.lbl_info.pack(pady=5)
        
        f_listbox = tk.Frame(main_frame)
        f_listbox.pack(fill=tk.BOTH, expand=True, pady=5)
        
        sc_y = ttk.Scrollbar(f_listbox)
        sc_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.lb_arquivos = tk.Listbox(f_listbox, height=6, font=("Consolas", 9), yscrollcommand=sc_y.set)
        self.lb_arquivos.pack(fill=tk.BOTH, expand=True)
        sc_y.config(command=self.lb_arquivos.yview)
        
        f_botoes_sel = tk.Frame(main_frame)
        f_botoes_sel.pack(pady=5)
        
        tk.Button(f_botoes_sel, text="+ Selecionar Arquivos", command=self.select_files, bg="#3498db", fg="white", font=("Segoe UI", 9, "bold"), width=20).grid(row=0, column=0, padx=5)
        tk.Button(f_botoes_sel, text="Limpar Lista", command=self.limpar, bg="#e74c3c", fg="white", font=("Segoe UI", 9, "bold"), width=15).grid(row=0, column=1, padx=5)
        
        self.btn_conv = tk.Button(main_frame, text="2. EXECUTAR CONVERSÃO EM LOTE AGORA", command=self.converter_assincrono, bg="#e67e22", fg="white", font=("Segoe UI", 12, "bold"), height=2, state=tk.DISABLED)
        self.btn_conv.pack(pady=15, fill=tk.X)

        ThemeManager.apply_theme_to_all(self.frame)

    def mudar_destino(self):
        new_dir = filedialog.askdirectory(initialdir=self.target_dir)
        if new_dir:
            self.target_dir = new_dir
            self.lbl_dest.config(text=self.target_dir)
            save_config("conversor_dest", self.target_dir)

    def atualizar_modo(self):
        self.sentido_conversao = self.var_modo.get()
        self.limpar()

    def select_files(self):
        if self.sentido_conversao == "pdf2docx":
            tipos = [("Arquivos PDF", "*.pdf")]
        else:
            tipos = [("Arquivos Word DOCX", "*.docx")]
            
        fs = filedialog.askopenfilenames(filetypes=tipos)
        for f in fs:
            if f not in self.active_file_paths:
                self.active_file_paths.append(f)
                self.lb_arquivos.insert(tk.END, os.path.basename(f))
                
        if self.active_file_paths:
            fg_color = "black" if ThemeManager.current_theme == "light" else "white"
            self.lbl_info.config(text=f"Total de {len(self.active_file_paths)} arquivos selecionados:", fg=fg_color, font=("Segoe UI", 10, "bold"))
            self.btn_conv.config(state=tk.NORMAL)
        else:
            self.lbl_info.config(text="Nenhum arquivo selecionado", fg="gray", font=("Segoe UI", 10, "italic"))
            self.btn_conv.config(state=tk.DISABLED)

    def limpar(self):
        self.active_file_paths = []
        self.lb_arquivos.delete(0, tk.END)
        self.lbl_info.config(text="Nenhum arquivo selecionado", fg="gray", font=("Segoe UI", 10, "italic"))
        self.btn_conv.config(state=tk.DISABLED)

    def converter_assincrono(self):
        if not self.active_file_paths: return
        
        sentido = self.sentido_conversao
        origens = list(self.active_file_paths)
        
        modal = ProcessingModal(self.root, "Processando", f"Convertendo arquivo 1 de {len(origens)}... Por favor, aguarde.")
        threading.Thread(target=self._converter, args=(origens, sentido, modal), daemon=True).start()

    def _converter(self, origens, sentido, modal):
        erros = []
        concluidos = 0
        total = len(origens)
        
        word = None
        if sentido == "docx2pdf" and win32com.client is not None:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
            except Exception as e_word_init:
                print(f"Não foi possível iniciar o Word COM em lote: {e_word_init}")
                
        try:
            for idx, origem in enumerate(origens):
                modal.update_message(f"Convertendo ({idx+1}/{total}):\n{os.path.basename(origem)}")
                
                try:
                    if sentido == "pdf2docx":
                        out = os.path.join(self.target_dir, os.path.splitext(os.path.basename(origem))[0] + ".docx")
                        cv = Converter(origem)
                        cv.convert(out)
                        cv.close()
                        concluidos += 1
                    else:
                        out = os.path.join(self.target_dir, os.path.splitext(os.path.basename(origem))[0] + ".pdf")
                        
                        success = False
                        if word is not None:
                            try:
                                doc = word.Documents.Open(os.path.abspath(origem))
                                doc.SaveAs(os.path.abspath(out), FileFormat=17)
                                doc.Close()
                                success = True
                                concluidos += 1
                            except Exception as e_com:
                                print(f"Falha na conversão do arquivo via COM: {e_com}")
                                
                        if not success and docx_to_pdf_convert is not None:
                            try:
                                docx_to_pdf_convert(origem, out)
                                success = True
                                concluidos += 1
                            except Exception as e_lib:
                                print(f"Falha na conversão do arquivo via docx2pdf: {e_lib}")
                                
                        if not success:
                            raise Exception("A conversão necessita do Microsoft Word instalado ou da biblioteca 'docx2pdf'.")
                            
                except Exception as ex_arq:
                    erros.append(f"{os.path.basename(origem)}: {ex_arq}")
            
            if word is not None:
                try:
                    word.Quit()
                except:
                    pass
            
            self.frame.after(0, modal.close)
            
            def exibir_resultado_lote():
                if erros:
                    lista_erros = "\n".join(erros[:5])
                    if len(erros) > 5:
                        lista_erros += f"\n...e mais {len(erros)-5} arquivos."
                    messagebox.showwarning(
                        "Conversão Concluída com Alertas",
                        f"Processo concluído!\n\nConvertidos com sucesso: {concluidos} de {total} arquivos.\n\nFalhas encontradas:\n{lista_erros}"
                    )
                else:
                    messagebox.showinfo("Sucesso", f"Todos os {total} arquivos foram convertidos com absoluto sucesso!")
                self.limpar()
                
            self.frame.after(100, exibir_resultado_lote)
            
        except Exception as e_geral:
            if word is not None:
                try: word.Quit()
                except: pass
            self.frame.after(0, modal.close)
            self.frame.after(100, lambda: show_detailed_error("Erro de Lote", "Ocorreu um erro geral no lote de conversão.", e_geral))

class DashboardEACE:
    def __init__(self, parent_frame, main_app_callback, gspread_client=None):
        self.frame = tk.Frame(parent_frame)
        self.frame.pack(fill=tk.BOTH, expand=True)
        self.main_app_callback = main_app_callback
        self.gspread_client = gspread_client
        self.root = parent_frame.winfo_toplevel()
        
        self.df = None
        self.colors = ThemeManager.get_colors()
        self._build()
        
        self.frame.after(100, self.carregar_so_local)
        ThemeManager.apply_theme_to_all(self.frame)
        
    def _build(self):
        colors = self.colors
        bg = colors["bg"]
        card = colors["card_bg"]
        border = colors["border"]
        fg = colors["fg"]
        muted = colors["fg_muted"]
        entry_bg = colors["entry_bg"]
        entry_fg = colors["entry_fg"]
        
        self.frame.configure(bg=bg)
        
        # ── TOP BAR ──────────────────────────────────────────
        top = tk.Frame(self.frame, bg=card, height=48)
        top.pack(fill="x")
        top.pack_propagate(False)
        
        tk.Label(top, text="  📊  DASHBOARD DE MONITORAMENTO — EACE",
                 font=("Segoe UI", 12, "bold"), bg=card, fg="blue").pack(side="left", padx=8)
                 
        tk.Button(top, text="VOLTAR AO MENU", command=self.main_app_callback,
                  bg="#6c757d", fg="white", font=("Segoe UI", 8, "bold"),
                  bd=0, cursor="hand2", padx=10).pack(side="right", padx=10, pady=8)
                  
        # ── BARRA DE STATUS E ATUALIZAÇÃO ───────────────────
        bar = tk.Frame(self.frame, bg=bg, pady=10)
        bar.pack(fill="x", padx=16)
        
        self.lbl_status = tk.Label(bar, text="Verificando dados...", font=("Segoe UI", 9, "italic"), bg=bg, fg="orange")
        self.lbl_status.pack(side="left")
        
        tk.Button(bar, text="🔄 Sincronizar Nuvem (Sheets)", command=self.forcar_atualizacao,
                  bg="#17a2b8", fg="white", font=("Segoe UI", 9, "bold"),
                  bd=0, cursor="hand2", padx=12, pady=4).pack(side="right")
                  
        # Separador
        tk.Frame(self.frame, bg=border, height=1).pack(fill="x", pady=2)
        
        # Conteúdo Principal
        body = tk.Frame(self.frame, bg=bg)
        body.pack(fill="both", expand=True, padx=16, pady=10)
        
        # ── CARDS DE MÉTRICAS (Grid) ──────────────────────────
        self.cards_frame = tk.Frame(body, bg=bg)
        self.cards_frame.pack(fill="x", pady=(0, 10))
        self.cards_frame.columnconfigure((0, 1, 2, 3), weight=1, uniform="equal")

        self.cards_data = [
            ("Total de Escolas", "—", "lbl_card_total", "#3498db"),
            ("Instaladas / Concluídas", "—", "lbl_card_concluidas", "green"),
            ("Em Progresso", "—", "lbl_card_progresso", "orange"),
            ("Taxa de Conclusão", "—", "lbl_card_taxa", "#9b59b6")
        ]

        self.card_labels = {}
        for idx, (title, value, name, color) in enumerate(self.cards_data):
            f = tk.Frame(self.cards_frame, bg=card, bd=1, relief="solid", highlightthickness=0)
            f.grid(row=0, column=idx, padx=5, sticky="nsew")
            
            dec = tk.Frame(f, bg=color, height=3)
            dec.pack(fill="x")
            
            tk.Label(f, text=title.upper(), font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(pady=(8, 2))
            lbl = tk.Label(f, text=value, font=("Segoe UI", 16, "bold"), bg=card, fg=color)
            lbl.pack(pady=(0, 8))
            self.card_labels[name] = lbl
            
        # ── FILTROS (Grid de 5 colunas) ──────────────────────
        filters_frame = tk.Frame(body, bg=card, bd=1, relief="solid", highlightthickness=0)
        filters_frame.pack(fill="x", pady=5, ipady=5)
        
        for i in range(7):
            filters_frame.columnconfigure(i, weight=1)
            
        # Regional
        f_reg = tk.Frame(filters_frame, bg=card)
        f_reg.grid(row=0, column=0, padx=8, pady=5, sticky="ew")
        tk.Label(f_reg, text="REGIONAL:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_reg = ttk.Combobox(f_reg, state="readonly", font=("Segoe UI", 9))
        self.cb_reg.pack(fill="x", pady=(2, 0))
        self.cb_reg.bind("<<ComboboxSelected>>", self.on_reg_changed)
            
        # UF
        f_uf = tk.Frame(filters_frame, bg=card)
        f_uf.grid(row=0, column=1, padx=8, pady=5, sticky="ew")
        tk.Label(f_uf, text="UF:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_uf = ttk.Combobox(f_uf, state="readonly", font=("Segoe UI", 9))
        self.cb_uf.pack(fill="x", pady=(2, 0))
        self.cb_uf.bind("<<ComboboxSelected>>", self.on_uf_changed)
        
        # Município
        f_mun = tk.Frame(filters_frame, bg=card)
        f_mun.grid(row=0, column=2, padx=8, pady=5, sticky="ew")
        tk.Label(f_mun, text="MUNICÍPIO:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_mun = ttk.Combobox(f_mun, state="readonly", font=("Segoe UI", 9))
        self.cb_mun.pack(fill="x", pady=(2, 0))
        self.cb_mun.bind("<<ComboboxSelected>>", self.on_filter_changed)
        
        # Status
        f_status = tk.Frame(filters_frame, bg=card)
        f_status.grid(row=0, column=3, padx=8, pady=5, sticky="ew")
        tk.Label(f_status, text="STATUS:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_status = ttk.Combobox(f_status, state="readonly", font=("Segoe UI", 9))
        self.cb_status.pack(fill="x", pady=(2, 0))
        self.cb_status.bind("<<ComboboxSelected>>", self.on_filter_changed)
        
        # Busca texto
        f_busca = tk.Frame(filters_frame, bg=card)
        f_busca.grid(row=0, column=4, padx=8, pady=5, sticky="ew")
        tk.Label(f_busca, text="BUSCA (INEP/NOME):", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.ent_search = tk.Entry(f_busca, bg=entry_bg, fg=entry_fg, insertbackground=fg, relief="solid", bd=1, font=("Segoe UI", 10))
        self.ent_search.pack(fill="x", ipady=2, pady=(2, 0))
        self.ent_search.bind("<KeyRelease>", self.on_filter_changed)
        
        # Ordem de Data
        f_ordem = tk.Frame(filters_frame, bg=card)
        f_ordem.grid(row=0, column=5, padx=8, pady=5, sticky="ew")
        tk.Label(f_ordem, text="ORDENAR DATA:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_ordem_data = ttk.Combobox(f_ordem, state="readonly", font=("Segoe UI", 9))
        self.cb_ordem_data["values"] = ["Padrão", "Mais Antigas Primeiro", "Mais Recentes Primeiro"]
        self.cb_ordem_data.set("Padrão")
        self.cb_ordem_data.pack(fill="x", pady=(2, 0))
        self.cb_ordem_data.bind("<<ComboboxSelected>>", self.on_filter_changed)

        # Mês do Filtro
        f_mes = tk.Frame(filters_frame, bg=card)
        f_mes.grid(row=0, column=6, padx=8, pady=5, sticky="ew")
        tk.Label(f_mes, text="MÊS/ANO:", font=("Segoe UI", 8, "bold"), bg=card, fg=muted).pack(anchor="w")
        self.cb_mes_data = ttk.Combobox(f_mes, state="readonly", font=("Segoe UI", 9))
        self.cb_mes_data["values"] = ["Todos"]
        self.cb_mes_data.set("Todos")
        self.cb_mes_data.pack(fill="x", pady=(2, 0))
        self.cb_mes_data.bind("<<ComboboxSelected>>", self.on_filter_changed)
        
        # ── TABELA (Treeview) ─────────────────────────────────
        table_frame = tk.Frame(body, bg=bg)
        table_frame.pack(fill="both", expand=True, pady=5)
        
        sc_y = ttk.Scrollbar(table_frame)
        sc_y.pack(side="right", fill="y")
        
        sc_x = ttk.Scrollbar(table_frame, orient="horizontal")
        sc_x.pack(side="bottom", fill="x")

        self.tree = ttk.Treeview(table_frame, columns=("inep", "escola", "municipio", "uf", "regional", "status", "data_inst", "observacao"),
                                 show="headings", yscrollcommand=sc_y.set, xscrollcommand=sc_x.set)
        
        self.tree.heading("inep", text="Código INEP")
        self.tree.heading("escola", text="Nome da Escola")
        self.tree.heading("municipio", text="Município")
        self.tree.heading("uf", text="UF")
        self.tree.heading("regional", text="Regional")
        self.tree.heading("status", text="Status")
        self.tree.heading("data_inst", text="Instalação")
        self.tree.heading("observacao", text="Observação")
        
        self.tree.column("inep", width=90, anchor="center")
        self.tree.column("escola", width=230, anchor="w")
        self.tree.column("municipio", width=140, anchor="w")
        self.tree.column("uf", width=50, anchor="center")
        self.tree.column("regional", width=90, anchor="center")
        self.tree.column("status", width=90, anchor="center")
        self.tree.column("data_inst", width=90, anchor="center")
        self.tree.column("observacao", width=180, anchor="w")
        
        self.tree.pack(fill="both", expand=True)
        sc_y.config(command=self.tree.yview)
        sc_x.config(command=self.tree.xview)
        
        # ── BOTÕES DE AÇÕES ───────────────────────────────────
        actions_frame = tk.Frame(body, bg=bg)
        actions_frame.pack(fill="x", pady=(5, 0))
        
        tk.Button(actions_frame, text="📋 Copiar Registro Selecionado", command=self.copiar_registro,
                  bg="#28a745", fg="white", font=("Segoe UI", 9, "bold"), bd=0, cursor="hand2", padx=12, pady=6).pack(side="left", padx=(0, 10))
                  
        tk.Button(actions_frame, text="📊 Copiar Tabela Filtrada (Excel)", command=self.copiar_tabela,
                  bg="#17a2b8", fg="white", font=("Segoe UI", 9, "bold"), bd=0, cursor="hand2", padx=12, pady=6).pack(side="left")

    def consolidar_matriz_raw(self, raw_data):
        if len(raw_data) < 2:
            return pd.DataFrame()
            
        reg_line = raw_data[0]
        header_line = raw_data[1]
        rows = raw_data[2:]
        
        regionais_map = []
        current_reg = "Desconhecido"
        for item in reg_line:
            if item.strip():
                current_reg = item.strip()
            regionais_map.append(current_reg)
            
        while len(regionais_map) < len(header_line):
            regionais_map.append(current_reg)
            
        records = []
        num_cols = len(header_line)
        
        for group_idx in range(0, num_cols, 5):
            if group_idx + 2 >= num_cols:
                break
                
            col_inep_name = header_line[group_idx].strip()
            if not col_inep_name or "inep" not in col_inep_name.lower():
                continue
                
            uf = col_inep_name.split()[-1] if len(col_inep_name.split()) > 1 else "N/A"
            regional = regionais_map[group_idx]
            regional_curta = regional.split()[0] if regional else "Desconhecido"
            
            for r in rows:
                if group_idx < len(r):
                    inep = str(r[group_idx]).strip()
                    inep = re.sub(r'\D', '', inep)
                    
                    if not inep:
                        continue
                        
                    data_inst = str(r[group_idx+1]).strip() if (group_idx+1 < len(r) and r[group_idx+1]) else ""
                    status = str(r[group_idx+2]).strip() if (group_idx+2 < len(r) and r[group_idx+2]) else ""
                    obs = str(r[group_idx+3]).strip() if (group_idx+3 < len(r) and r[group_idx+3]) else ""
                    links = str(r[group_idx+4]).strip() if (group_idx+4 < len(r) and r[group_idx+4]) else ""
                    
                    records.append({
                        "INEP": inep,
                        "UF": uf,
                        "REGIONAL": regional_curta,
                        "DATA_INSTALAÇÃO": data_inst,
                        "STATUS": status if status else "PENDENTE",
                        "OBSERVAÇÃO": obs,
                        "LINKS": links
                    })
                    
        return pd.DataFrame(records)

    def enriquecer_dados_escola(self, df_consolidado):
        if df_consolidado.empty:
            return df_consolidado
            
        df_cache = load_cache()
        if df_cache.empty:
            df_consolidado["ESCOLA"] = "—"
            df_consolidado["MUNICÍPIO"] = "—"
            return df_consolidado
            
        col_inep_c = None
        col_nome_c = None
        col_mun_c = None
        
        for c in df_cache.columns:
            c_low = c.lower()
            if "inep" in c_low and not col_inep_c:
                col_inep_c = c
            elif ("escola" in c_low or "nome" in c_low or "unidade" in c_low) and not col_nome_c:
                col_nome_c = c
            elif ("munic" in c_low or "cidade" in c_low) and not col_mun_c:
                col_mun_c = c
                
        if not col_inep_c or not col_nome_c:
            df_consolidado["ESCOLA"] = "—"
            df_consolidado["MUNICÍPIO"] = "—"
            return df_consolidado
            
        df_cache_clean = df_cache.copy()
        df_cache_clean[col_inep_c] = df_cache_clean[col_inep_c].astype(str).str.replace(r'\D', '', regex=True).str.strip()
        
        map_nome = dict(zip(df_cache_clean[col_inep_c], df_cache_clean[col_nome_c]))
        map_mun = dict(zip(df_cache_clean[col_inep_c], df_cache_clean[col_mun_c])) if col_mun_c else {}
        
        df_consolidado["ESCOLA"] = df_consolidado["INEP"].map(map_nome).fillna("—")
        df_consolidado["MUNICÍPIO"] = df_consolidado["INEP"].map(map_mun).fillna("—")
        
        return df_consolidado

    def importar_do_csv_local(self, path):
        if not os.path.exists(path):
            return pd.DataFrame()
        try:
            import csv
            raw_data = []
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f, delimiter=',')
                for row in reader:
                    raw_data.append(row)
            if len(raw_data) < 2:
                return pd.DataFrame()
            
            df_cons = self.consolidar_matriz_raw(raw_data)
            return self.enriquecer_dados_escola(df_cons)
        except Exception as e:
            print(f"Erro ao ler CSV local: {e}")
            return pd.DataFrame()

    def carregar_so_local(self):
        cache_p = get_dashboard_cache_path()
        
        self.df = load_cache(cache_p)
        
        if self.df.empty:
            csv_path = os.path.join(get_base_path(), "Validação Instação - Plan1.csv")
            if os.path.exists(csv_path):
                self.df = self.importar_do_csv_local(csv_path)
                if not self.df.empty:
                    save_cache(self.df, cache_p)
                    self.lbl_status.config(text="✅ Importado de Validação Instação - Plan1.csv local.", fg="green")
            
        if not self.df.empty:
            self.popular_filtros(reset_municipios=True)
            self.aplicar_filtros()
            self.lbl_status.config(text=f"✅ Dados locais carregados — {len(self.df)} escolas monitoradas.", fg="green")
        else:
            self.lbl_status.config(text="⚠️ Sem base de dados local. Clique em 'Sincronizar Nuvem' para carregar.", fg="orange")

    def forcar_atualizacao(self):
        self.lbl_status.config(text="⏳ Conectando à nuvem e baixando dados... Aguarde.", fg="blue")
        threading.Thread(target=self._sincronizar_nuvem, daemon=True).start()

    def _sincronizar_nuvem(self):
        try:
            if not self.gspread_client:
                self.gspread_client = authenticate_google_sheets()
            if not self.gspread_client:
                raise Exception("Sem credenciais válidas do Google (credentials.json ausente).")
            
            spreadsheet = self.gspread_client.open_by_key(DASHBOARD_SPREADSHEET_ID)
            worksheet = None
            for w in spreadsheet.worksheets():
                if str(w.id) == "1113059829":
                    worksheet = w
                    break
            if not worksheet:
                worksheet = spreadsheet.get_worksheet(0)
            
            raw_data = worksheet.get_all_values()
            if not raw_data:
                raise Exception("A planilha está vazia.")
                
            df_cons = self.consolidar_matriz_raw(raw_data)
            df_nuvem = self.enriquecer_dados_escola(df_cons)
            
            self.df = df_nuvem
            save_cache(df_nuvem, get_dashboard_cache_path())
            
            self.frame.after(0, lambda: self.popular_filtros(reset_municipios=True))
            self.frame.after(0, self.aplicar_filtros)
            self.frame.after(0, lambda: self.lbl_status.config(text=f"✅ Sincronizado com sucesso! {len(self.df)} registros.", fg="green"))
            
        except Exception as e:
            if self.df is not None and not self.df.empty:
                self.frame.after(0, lambda: self.lbl_status.config(text="⚠️ Falha na rede. Operando com dados locais salvos.", fg="orange"))
            else:
                self.frame.after(0, lambda: self.lbl_status.config(text="❌ ERRO DE CONEXÃO. Sem cache e sem rede.", fg="red"))
                self.frame.after(0, lambda: show_detailed_error("Erro de Conexão", "Não foi possível acessar a planilha em nuvem.", e))

    def popular_filtros(self, reset_municipios=True):
        if self.df is None or self.df.empty: return

        # Regional
        regs = sorted(self.df["REGIONAL"].dropna().unique())
        self.cb_reg["values"] = ["Todas"] + [str(r).strip() for r in regs if str(r).strip()]
        self.cb_reg.set("Todas")

        # UF
        ufs = sorted(self.df["UF"].dropna().unique())
        self.cb_uf["values"] = ["Todas"] + [str(u).strip() for u in ufs if str(u).strip()]
        self.cb_uf.set("Todas")

        # Status
        status = sorted(self.df["STATUS"].dropna().unique())
        self.cb_status["values"] = ["Todos"] + [str(s).strip() for s in status if str(s).strip()]
        self.cb_status.set("Todos")
        
        # Mês/Ano
        if hasattr(self, "cb_mes_data"):
            datas_validas = pd.to_datetime(self.df["DATA_INSTALAÇÃO"], format="%d/%m/%Y", errors="coerce").dropna()
            meses_unicos = sorted(datas_validas.dt.strftime("%m/%Y").unique())
            self.cb_mes_data["values"] = ["Todos"] + meses_unicos
            self.cb_mes_data.set("Todos")

        if reset_municipios:
            muns = sorted(self.df["MUNICÍPIO"].dropna().unique())
            self.cb_mun["values"] = ["Todos"] + [str(m).strip() for m in muns if str(m).strip()]
            self.cb_mun.set("Todos")

    def on_reg_changed(self, event=None):
        reg_sel = self.cb_reg.get()
        if self.df is not None and not self.df.empty:
            if reg_sel and reg_sel != "Todas":
                df_reg = self.df[self.df["REGIONAL"] == reg_sel]
                ufs = sorted(df_reg["UF"].dropna().unique())
            else:
                ufs = sorted(self.df["UF"].dropna().unique())
                
            self.cb_uf["values"] = ["Todas"] + [str(u).strip() for u in ufs if str(u).strip()]
            self.cb_uf.set("Todas")
        else:
            self.cb_uf["values"] = ["Todas"]
            self.cb_uf.set("Todas")
            
        self.on_uf_changed()

    def on_uf_changed(self, event=None):
        uf_sel = self.cb_uf.get()
        reg_sel = self.cb_reg.get()
        
        if self.df is not None and not self.df.empty:
            df_filtered = self.df.copy()
            if reg_sel and reg_sel != "Todas":
                df_filtered = df_filtered[df_filtered["REGIONAL"] == reg_sel]
            if uf_sel and uf_sel != "Todas":
                df_filtered = df_filtered[df_filtered["UF"] == uf_sel]
                
            muns = sorted(df_filtered["MUNICÍPIO"].dropna().unique())
            self.cb_mun["values"] = ["Todos"] + [str(m).strip() for m in muns if str(m).strip()]
            self.cb_mun.set("Todos")
        else:
            self.cb_mun["values"] = ["Todos"]
            self.cb_mun.set("Todos")
            
        self.aplicar_filtros()

    def on_filter_changed(self, event=None):
        self.aplicar_filtros()

    def aplicar_filtros(self):
        if self.df is None or self.df.empty: return

        df_filtrado = self.df.copy()

        # Regional
        reg_sel = self.cb_reg.get()
        if reg_sel and reg_sel != "Todas":
            df_filtrado = df_filtrado[df_filtrado["REGIONAL"] == reg_sel]

        # UF
        uf_sel = self.cb_uf.get()
        if uf_sel and uf_sel != "Todas":
            df_filtrado = df_filtrado[df_filtrado["UF"] == uf_sel]

        # Município
        mun_sel = self.cb_mun.get()
        if mun_sel and mun_sel != "Todos":
            df_filtrado = df_filtrado[df_filtrado["MUNICÍPIO"] == mun_sel]

        # Status
        status_sel = self.cb_status.get()
        if status_sel and status_sel != "Todos":
            df_filtrado = df_filtrado[df_filtrado["STATUS"] == status_sel]

        # Busca Texto
        busca = self.ent_search.get().strip().lower()
        if busca:
            condicoes = [
                df_filtrado["INEP"].astype(str).str.lower().str.contains(busca),
                df_filtrado["ESCOLA"].astype(str).str.lower().str.contains(busca),
                df_filtrado["MUNICÍPIO"].astype(str).str.lower().str.contains(busca)
            ]
            combined = condicoes[0]
            for cond in condicoes[1:]:
                combined = combined | cond
            df_filtrado = df_filtrado[combined]

        # Processamento de Data (Filtro e Ordenação)
        mes_sel = getattr(self, "cb_mes_data", None)
        ordem_sel = getattr(self, "cb_ordem_data", None)
        
        precisa_data = (mes_sel and mes_sel.get() != "Todos") or (ordem_sel and ordem_sel.get() != "Padrão")
        
        if precisa_data:
            df_filtrado["DATA_PARSED"] = pd.to_datetime(df_filtrado["DATA_INSTALAÇÃO"], format="%d/%m/%Y", errors="coerce")
            
            # 1. Filtrar por Mês
            if mes_sel and mes_sel.get() != "Todos":
                df_filtrado = df_filtrado[df_filtrado["DATA_PARSED"].dt.strftime("%m/%Y") == mes_sel.get()]
                
            # 2. Ordenar
            if ordem_sel and ordem_sel.get() != "Padrão":
                ordem = ordem_sel.get()
                if ordem == "Mais Antigas Primeiro":
                    df_filtrado = df_filtrado.sort_values(by="DATA_PARSED", ascending=True, na_position="last")
                elif ordem == "Mais Recentes Primeiro":
                    df_filtrado = df_filtrado.sort_values(by="DATA_PARSED", ascending=False, na_position="last")
                    
            df_filtrado = df_filtrado.drop(columns=["DATA_PARSED"])

        # Atualizar Treeview
        self.tree.delete(*self.tree.get_children())
        
        for idx, row in df_filtrado.iterrows():
            self.tree.insert("", "end", values=(
                row.get("INEP", ""),
                row.get("ESCOLA", ""),
                row.get("MUNICÍPIO", ""),
                row.get("UF", ""),
                row.get("REGIONAL", ""),
                row.get("STATUS", ""),
                row.get("DATA_INSTALAÇÃO", ""),
                row.get("OBSERVAÇÃO", "")
            ))

        self.atualizar_cards(df_filtrado)

    def atualizar_cards(self, df_filtrado):
        if df_filtrado.empty:
            self.card_labels["lbl_card_total"].config(text="0")
            self.card_labels["lbl_card_concluidas"].config(text="0")
            self.card_labels["lbl_card_progresso"].config(text="0")
            self.card_labels["lbl_card_taxa"].config(text="0.0%")
            return

        total = len(df_filtrado)
        concluidas = 0
        progresso = 0
        
        status_series = df_filtrado["STATUS"].astype(str).str.lower().str.strip()
        
        termos_concluidos = ["ok", "concluido", "concluída", "concluidos", "concluídas", "finalizado", "finalizada", "instalado", "instalada", "ativo", "ativa", "aprovado", "aprovada"]
        termos_progresso = ["andamento", "progresso", "instalando", "execucao", "execução", "iniciado", "iniciada", "vistoria", "agendado", "pendente"]
        
        for st in status_series:
            if any(t in st for t in termos_concluidos):
                concluidas += 1
            elif any(t in st for t in termos_progresso):
                progresso += 1
            else:
                if st and st != "nan" and st != "":
                    progresso += 1

        taxa = (concluidas / total) * 100 if total > 0 else 0.0
        
        self.card_labels["lbl_card_total"].config(text=str(total))
        self.card_labels["lbl_card_concluidas"].config(text=str(concluidas))
        self.card_labels["lbl_card_progresso"].config(text=str(progresso))
        self.card_labels["lbl_card_taxa"].config(text=f"{taxa:.1f}%")

    def copiar_registro(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("Aviso", "Selecione um registro na tabela para copiar.")
            return
            
        item = self.tree.item(sel[0])
        vals = item["values"]
        
        texto = (
            f"INEP:      {vals[0]}\n"
            f"ESCOLA:    {vals[1]}\n"
            f"MUNICÍPIO: {vals[2]}\n"
            f"UF:        {vals[3]}\n"
            f"REGIONAL:  {vals[4]}\n"
            f"STATUS:    {vals[5]}\n"
            f"INSTALAÇÃO:{vals[6]}\n"
            f"OBSERVAÇÃO:{vals[7]}"
        )
        self.frame.clipboard_clear()
        self.frame.clipboard_append(texto)
        messagebox.showinfo("Sucesso ✅", "Registro copiado para a área de transferência!")

    def copiar_tabela(self):
        children = self.tree.get_children()
        if not children:
            messagebox.showwarning("Aviso", "Tabela vazia. Nada para copiar.")
            return
            
        linhas = []
        linhas.append("Código INEP\tNome da Escola\tMunicípio\tUF\tStatus")
        
        for child in children:
            vals = self.tree.item(child)["values"]
            linhas.append(f"{vals[0]}\t{vals[1]}\t{vals[2]}\t{vals[3]}\t{vals[4]}")
            
        texto = "\n".join(linhas)
        self.frame.clipboard_clear()
        self.frame.clipboard_append(texto)
        messagebox.showinfo("Sucesso ✅", f"Copiadas {len(children)} linhas da tabela no formato Excel (Tabulado)!")

class MainApp:
    def __init__(self, root, gs):
        self.root = root
        self.gs = gs
        self.current_frame = None
        
        # Carregar Tema Inicial da Configuração
        config = load_config()
        ThemeManager.current_theme = config.get("theme", "dark")
        
        self.show_main_menu()

    def show_main_menu(self):
        if self.current_frame: self.current_frame.destroy()
        self.current_frame = tk.Frame(self.root)
        self.current_frame.pack(fill=tk.BOTH, expand=True)
        
        # Barra superior com título e toggle de tema
        header_frame = tk.Frame(self.current_frame, pady=10)
        header_frame.pack(fill="x", padx=20)
        
        title_frame = tk.Frame(header_frame)
        title_frame.pack(side="left")
        
        tk.Label(title_frame, text="FERRAMENTAS DE APOIO", font=("Segoe UI", 20, "bold")).pack(anchor="w")
        tk.Label(title_frame, text="BITNET TELECOM - Automações", font=("Segoe UI", 9, "bold"), fg="#27ae60").pack(anchor="w")
        
        # Botão de Toggle de Tema
        self.btn_tema = ModernButton(
            header_frame, 
            text="☀️ Modo Claro" if ThemeManager.current_theme == "dark" else "🌙 Modo Escuro", 
            command=self.alternar_tema,
            bg_type="#34495e", 
            font=("Segoe UI", 9, "bold"),
            radius=12,
            height=30,
            width=110
        )
        self.btn_tema.pack(side="right", pady=5)
        
        f_botoes = tk.Frame(self.current_frame, pady=20)
        f_botoes.pack(expand=True)
        
        buttons = [
            ("Gerador de Contrato EACE", lambda: self.open(GeradorContratoApp), "#007bff"),
            ("Envio de Pleito (E-mail)", lambda: self.open(EnviadorPleitoApp), "#eab308"),
            ("Consulta EACE", lambda: self.open(ConsultaEACE), "#17a2b8"),
            ("Unificador de Evidências", lambda: self.open(UnificadorApp), "#28a745"),
            ("Conversor Bidirecional PDF ⇄ Word", lambda: self.open(ConversorBidirecionalApp), "#e67e22"),
            ("Dashboard EACE", lambda: self.open(DashboardEACE), "#8a2be2")
        ]
        
        for text, cmd, color in buttons:
            ModernButton(f_botoes, text=text, command=cmd, bg_type=color, font=("Segoe UI", 11, "bold"), radius=16, height=46, width=380).pack(pady=10)
        
        footer = tk.Frame(self.current_frame, padx=10, pady=5)
        footer.pack(side="bottom", fill="x")
        tk.Label(footer, text="v3.0 - Otimizado e Unificado", font=("Segoe UI", 8), fg="gray").pack(side="left", pady=10)
        tk.Label(footer, text="desenvolvido por Silas Duarte", font=("Segoe UI", 8, "italic"), fg="#bdc3c7").pack(side="right", pady=10)
        
        ThemeManager.apply_theme_to_all(self.root)
        
    def alternar_tema(self):
        ThemeManager.toggle_theme(self.root)
        self.btn_tema.update_text("☀️ Modo Claro" if ThemeManager.current_theme == "dark" else "🌙 Modo Escuro")
        ThemeManager.apply_theme_to_all(self.root)

    def open(self, cls):
        if self.current_frame: self.current_frame.destroy()
        tool_instance = cls(self.root, self.show_main_menu, self.gs)
        self.current_frame = tool_instance.frame
        ThemeManager.apply_theme_to_all(self.root)

if __name__ == "__main__":
    multiprocessing.freeze_support()
    root = tk.Tk()
    root.withdraw()
    gs = authenticate_google_sheets()
    root.deiconify()
    root.title("Ferramentas de Apoio - EACE")
    root.geometry("900x780")
    app = MainApp(root, gs)
    root.mainloop()
