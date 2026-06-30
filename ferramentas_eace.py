import sys
import os
import json
import io
import base64
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageOps, ImageTk, ImageDraw, ImageFont
import pillow_heif
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from email.message import EmailMessage
from google_auth_oauthlib.flow import InstalledAppFlow
import multiprocessing
import traceback
import shutil
import textwrap
import tempfile
import threading
from datetime import datetime

# Registrar suporte a HEIC para Pillow
pillow_heif.register_heif_opener()

# --- CAMINHOS BÁSICOS ---
def get_base_path():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = get_base_path()
    return os.path.join(base_path, relative_path)

CONFIG_FILE = os.path.join(get_base_path(), "config.json")
LOG_FILE = os.path.join(get_base_path(), "eace_app.log")
CACHE_FILE = os.path.join(get_base_path(), "eace_cache.json")
CACHE_FILE_INFANTIL = os.path.join(get_base_path(), "infantil_cache.json")

# --- SISTEMA DE LOGS ---
def log_message(level, message):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_line = f"[{timestamp}] [{level.upper()}] {message}\n"
    print(log_line.strip())
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(log_line)
    except Exception as e:
        print(f"Erro ao escrever log: {e}")

# --- GERENCIAMENTO DE CONFIGURAÇÕES ---
DEFAULT_CONFIG = {
    "EMAIL_TO": "planejamento@eace.org.br",
    "EMAIL_CC": "demostenes.freitas@eace.org.br,nildo.sobral@eace.org.br,salatecnica@eace.org.br,douglas.bicudo@eace.org.br,nielsen.bezerra@bitinternet.com.br,brenomatheus936@gmail.com,fabiano.nascimento@bitinternet.com.br,matheus.campos@bitinternet.com.br",
    "ID_PASTA_EVIDENCIAS_DRIVE": "1UbfeHCcQfafevuHO5lDUs7fBwu7o4IOQ",
    "SPREADSHEET_ID": "1Onw1vaSO2SIQ_OfAoDPI6ycnXWTAZ2ijhtujAOhI9UM",
    "GOOGLE_DRIVE_FOLDER_ID": "1RSsspBObpQe5S30o0L3YzpyY9AjpfwKj",
    "NETWORK_PATH": r"\\Desktop-noc-1\noc\Revisado-Enviar EACE",
    "unificador_save_path": os.path.join(os.path.expanduser("~"), "Documents")
}

def load_config():
    config = DEFAULT_CONFIG.copy()
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                config.update(loaded)
        except Exception as e:
            log_message("WARNING", f"Erro ao ler config.json, usando padroes: {e}")
    else:
        save_config(config)
    return config

def save_config(config):
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=4, ensure_ascii=False)
        log_message("INFO", "Configuracoes gravadas com sucesso.")
    except Exception as e:
        log_message("ERROR", f"Falha ao salvar config.json: {e}")

# Instancia de configuracoes ativa globalmente
APP_CONFIG = load_config()

# --- CACHE LOCAL ---
def load_cache(cache_path=None) -> pd.DataFrame:
    """Carrega o DataFrame do cache local. Retorna vazio se nao existir ou falhar."""
    if cache_path is None:
        cache_path = CACHE_FILE
    if os.path.exists(cache_path):
        try:
            return pd.read_json(cache_path, dtype=str)
        except Exception as e:
            log_message("WARNING", f"Erro ao ler cache em {cache_path}: {e}")
            return pd.DataFrame()
    return pd.DataFrame()

def save_cache(df: pd.DataFrame, cache_path=None):
    """Persiste o DataFrame no cache local."""
    if cache_path is None:
        cache_path = CACHE_FILE
    try:
        df.to_json(cache_path, force_ascii=False, indent=2)
        log_message("INFO", f"Cache persistido com sucesso em {cache_path}.")
    except Exception as e:
        log_message("WARNING", f"Aviso: nao foi possivel salvar cache: {e}")

# --- JANELA DE ERRO DETALHADA ---
def show_detailed_error(title, message, error_obj):
    log_message("ERROR", f"{title}: {message} - {error_obj}")
    error_window = tk.Toplevel()
    error_window.title(title)
    error_window.geometry("700x500")
    error_window.configure(bg="#121214")
    
    tk.Label(error_window, text=message, font=("Arial", 12, "bold"), fg="#ef4444", bg="#121214").pack(pady=10)
    
    container = tk.Frame(error_window, bg="#202024")
    container.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
    
    scrollbar_y = ttk.Scrollbar(container)
    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_x = ttk.Scrollbar(container, orient=tk.HORIZONTAL)
    scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    
    text_area = tk.Text(
        container, 
        wrap=tk.NONE, 
        bg="#1e1e24", 
        fg="#e2e8f0", 
        insertbackground="white", 
        font=("Courier New", 10),
        yscrollcommand=scrollbar_y.set, 
        xscrollcommand=scrollbar_x.set
    )
    text_area.pack(fill=tk.BOTH, expand=True)
    
    scrollbar_y.config(command=text_area.yview)
    scrollbar_x.config(command=text_area.xview)
    
    full_error = f"Erro: {error_obj}\n\nTraceback Completo:\n{traceback.format_exc()}"
    text_area.insert(tk.END, full_error)
    text_area.config(state=tk.DISABLED)
    
    btn_fechar = tk.Button(
        error_window, 
        text="Fechar", 
        command=error_window.destroy, 
        bg="#6c757d", 
        fg="white", 
        font=("Arial", 10, "bold"),
        relief="flat",
        cursor="hand2"
    )
    btn_fechar.pack(pady=10)
    btn_fechar.bind("<Enter>", lambda e: btn_fechar.config(bg="#5a6268"))
    btn_fechar.bind("<Leave>", lambda e: btn_fechar.config(bg="#6c757d"))

# --- AUTENTICAÇÃO E CONEXÃO ASSÍNCRONA ---
GOOGLE_CONNECTION = {
    "gspread_client": None,
    "drive_service": None,
    "status": "Nao Conectado",  # "Nao Conectado", "Conectando...", "Conectado", "Offline (Erro)"
    "error_message": "",
    "lock": threading.Lock()
}

def get_credentials():
    creds_path = os.path.join(get_base_path(), "credentials.json")
    if not os.path.exists(creds_path): 
        creds_path = "credentials.json"
    if not os.path.exists(creds_path):
        raise FileNotFoundError("Arquivo 'credentials.json' nao encontrado para Service Account.")
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    return ServiceAccountCredentials.from_json_keyfile_name(creds_path, scope)

def thread_autenticar_google(on_complete_callback=None):
    """Executa a autenticacao dos servicos Google em background (thread-safe)."""
    global GOOGLE_CONNECTION
    with GOOGLE_CONNECTION["lock"]:
        GOOGLE_CONNECTION["status"] = "Conectando..."
    
    log_message("INFO", "Iniciando autenticacao Google APIs em background...")
    try:
        creds = get_credentials()
        gspread_client = gspread.authorize(creds)
        drive_service = build('drive', 'v3', credentials=creds)
        
        with GOOGLE_CONNECTION["lock"]:
            GOOGLE_CONNECTION["gspread_client"] = gspread_client
            GOOGLE_CONNECTION["drive_service"] = drive_service
            GOOGLE_CONNECTION["status"] = "Conectado"
            GOOGLE_CONNECTION["error_message"] = ""
            
        log_message("SUCCESS", "Autenticacao Google concluida com sucesso (Sheets & Drive).")
        
    except Exception as e:
        error_msg = str(e)
        with GOOGLE_CONNECTION["lock"]:
            GOOGLE_CONNECTION["status"] = "Offline (Erro)"
            GOOGLE_CONNECTION["error_message"] = error_msg
            
        log_message("ERROR", f"Falha na autenticacao Google: {error_msg}")
        
    if on_complete_callback:
        try:
            on_complete_callback()
        except Exception as cb_err:
            log_message("ERROR", f"Erro no callback de autenticacao: {cb_err}")

def authenticate_gmail():
    SCOPES = ['https://www.googleapis.com/auth/gmail.send']
    token_path = os.path.join(get_base_path(), "token.json")
    secret_path = os.path.join(get_base_path(), "AuthNOC.json")
    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(secret_path):
                raise FileNotFoundError(f"Arquivo '{secret_path}' nao encontrado.")
            flow = InstalledAppFlow.from_client_secrets_file(secret_path, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token: 
            token.write(creds.to_json())
    return build('gmail', 'v1', credentials=creds)

def authenticate_gmail_readonly():
    SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
    token_path = os.path.join(get_base_path(), "token_leitura.json")
    secret_path = os.path.join(get_base_path(), "AuthNOC.json")
    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(secret_path):
                raise FileNotFoundError(f"Arquivo '{secret_path}' nao encontrado.")
            flow = InstalledAppFlow.from_client_secrets_file(secret_path, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token: 
            token.write(creds.to_json())
    return build('gmail', 'v1', credentials=creds)

def fetch_data_by_id(gspread_client, sheet_id, worksheet_name):
    if gspread_client is None:
        raise Exception("Cliente Google Sheets nao autenticado.")
    try:
        spreadsheet = gspread_client.open_by_key(sheet_id)
        worksheet = spreadsheet.worksheet(worksheet_name)
        raw_data = worksheet.get_all_values()
        if not raw_data: 
            return pd.DataFrame()
        return pd.DataFrame(raw_data[1:], columns=raw_data[0])
    except Exception as e: 
        raise Exception(f"Erro na aba '{worksheet_name}': {e}")

# --- DIALOGO DE CONFIGURAÇÕES GRÁFICAS ---
class ConfiguracoesDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Configuracoes Globais")
        self.geometry("600x520")
        self.configure(bg="#121214")
        self.resizable(False, False)
        
        self.transient(parent)
        self.grab_set()
        
        tk.Label(self, text="⚙️ CONFIGURAÇÕES DO SISTEMA", font=("Arial", 16, "bold"), fg="#e2e8f0", bg="#121214").pack(pady=(15, 10))
        
        form_frame = tk.Frame(self, bg="#202024", padx=20, pady=15)
        form_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        self.entries = {}
        fields = [
            ("ID_PASTA_EVIDENCIAS_DRIVE", "ID Pasta Drive de Evidencias (HEIC/PDF):"),
            ("SPREADSHEET_ID", "ID Planilha Google (EACE / UP_DOWN):"),
            ("GOOGLE_DRIVE_FOLDER_ID", "ID Pasta Drive Destino DOCX:"),
            ("NETWORK_PATH", "Caminho da Pasta de Rede Compartilhada:"),
            ("unificador_save_path", "Pasta de Destino padrao do Unificador (PDF):")
        ]
        
        config_atual = load_config()
        
        for i, (key, label_text) in enumerate(fields):
            lbl = tk.Label(form_frame, text=label_text, font=("Arial", 9, "bold"), fg="#94a3b8", bg="#202024", anchor="w")
            lbl.pack(fill=tk.X, pady=(6, 2))
            
            entry = tk.Entry(form_frame, font=("Arial", 10), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2)
            entry.pack(fill=tk.X, ipady=4)
            entry.insert(0, config_atual.get(key, ""))
            self.entries[key] = entry
            
            if key in ["NETWORK_PATH", "unificador_save_path"]:
                btn_browse = tk.Button(
                    form_frame,
                    text="Selecionar Pasta...",
                    font=("Arial", 8, "bold"),
                    bg="#4b5563",
                    fg="white",
                    relief="flat",
                    cursor="hand2",
                    command=lambda k=key: self.browse_directory(k)
                )
                btn_browse.pack(anchor="e", pady=(2, 0))
                btn_browse.bind("<Enter>", lambda e, b=btn_browse: b.config(bg="#374151"))
                btn_browse.bind("<Leave>", lambda e, b=btn_browse: b.config(bg="#4b5563"))
                
        btn_frame = tk.Frame(self, bg="#121214")
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=15)
        
        btn_salvar = tk.Button(
            btn_frame, 
            text="SALVAR CONFIGURAÇÕES", 
            font=("Arial", 10, "bold"),
            bg="#10b981", 
            fg="white", 
            relief="flat", 
            width=22, 
            height=2,
            cursor="hand2",
            command=self.salvar
        )
        btn_salvar.pack(side=tk.RIGHT, padx=(10, 20))
        btn_salvar.bind("<Enter>", lambda e: btn_salvar.config(bg="#059669"))
        btn_salvar.bind("<Leave>", lambda e: btn_salvar.config(bg="#10b981"))
        
        btn_cancelar = tk.Button(
            btn_frame, 
            text="CANCELAR", 
            font=("Arial", 10, "bold"),
            bg="#ef4444", 
            fg="white", 
            relief="flat", 
            width=12, 
            height=2,
            cursor="hand2",
            command=self.destroy
        )
        btn_cancelar.pack(side=tk.RIGHT, padx=5)
        btn_cancelar.bind("<Enter>", lambda e: btn_cancelar.config(bg="#dc2626"))
        btn_cancelar.bind("<Leave>", lambda e: btn_cancelar.config(bg="#ef4444"))

    def browse_directory(self, key):
        caminho_atual = self.entries[key].get()
        new_path = filedialog.askdirectory(initialdir=caminho_atual if os.path.exists(caminho_atual) else os.path.expanduser("~"))
        if new_path:
            self.entries[key].delete(0, tk.END)
            self.entries[key].insert(0, new_path)

    def salvar(self):
        global APP_CONFIG
        novas_configs = {}
        for key, entry in self.entries.items():
            novas_configs[key] = entry.get().strip()
            
        save_config(novas_configs)
        APP_CONFIG = load_config()
        messagebox.showinfo("Sucesso", "Configuracoes atualizadas e salvas com sucesso!", parent=self)
        self.destroy()

# --- CLASSE DE CONSULTA EACE ---
class ConsultaEACE:
    def __init__(self, parent_frame, main_app_callback):
        self.parent_frame = parent_frame
        self.main_app_callback = main_app_callback
        self.frame = tk.Frame(parent_frame, bg="#121214")
        self.frame.pack(fill=tk.BOTH, expand=True)
        
        header_frame = tk.Frame(self.frame, bg="#121214")
        header_frame.pack(fill=tk.X, pady=(15, 10))
        
        tk.Label(header_frame, text="🔍 CONSULTA RÁPIDA - EACE", font=("Arial", 16, "bold"), fg="#e2e8f0", bg="#121214").pack(side=tk.LEFT, padx=15)
        
        btn_voltar = tk.Button(
            header_frame, 
            text="⬅️ VOLTAR AO MENU", 
            command=self.main_app_callback, 
            bg="#4b5563", 
            fg="white", 
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=10,
            pady=5
        )
        btn_voltar.pack(side=tk.RIGHT, padx=15)
        btn_voltar.bind("<Enter>", lambda e: btn_voltar.config(bg="#374151"))
        btn_voltar.bind("<Leave>", lambda e: btn_voltar.config(bg="#4b5563"))
        
        card_frame = tk.Frame(self.frame, bg="#202024", padx=25, pady=20)
        card_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        self.lbl_status = tk.Label(card_frame, text="Carregando status local...", font=("Arial", 10, "italic"), fg="#94a3b8", bg="#202024", anchor="w")
        self.lbl_status.pack(fill=tk.X, pady=(0, 10))
        
        input_frame = tk.Frame(card_frame, bg="#202024")
        input_frame.pack(fill=tk.X, pady=10)
        
        tk.Label(input_frame, text="Código INEP:", font=("Arial", 11, "bold"), fg="#e2e8f0", bg="#202024").pack(side=tk.LEFT, padx=(0, 10))
        
        self.ent_inep = tk.Entry(input_frame, font=("Arial", 14), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", width=18, relief="flat", bd=2)
        self.ent_inep.pack(side=tk.LEFT, padx=5, ipady=3)
        self.ent_inep.bind("<Return>", lambda e: self.buscar_dados())
        
        btn_buscar = tk.Button(
            input_frame, 
            text="BUSCAR DADOS", 
            command=self.buscar_dados, 
            bg="#3b82f6", 
            fg="white", 
            font=("Arial", 11, "bold"), 
            relief="flat",
            padx=15,
            cursor="hand2"
        )
        btn_buscar.pack(side=tk.LEFT, padx=10, ipady=2)
        btn_buscar.bind("<Enter>", lambda e: btn_buscar.config(bg="#2563eb"))
        btn_buscar.bind("<Leave>", lambda e: btn_buscar.config(bg="#3b82f6"))
        
        self.btn_nuvem = tk.Button(
            input_frame, 
            text="🔄 Atualizar da Nuvem", 
            command=self._forcar_atualizacao_nuvem, 
            bg="#10b981", 
            fg="white", 
            font=("Arial", 10, "bold"),
            relief="flat",
            padx=15,
            cursor="hand2"
        )
        self.btn_nuvem.pack(side=tk.RIGHT, padx=5, ipady=2)
        self.btn_nuvem.bind("<Enter>", lambda e: self.btn_nuvem.config(bg="#059669"))
        self.btn_nuvem.bind("<Leave>", lambda e: self.btn_nuvem.config(bg="#10b981"))
        
        tk.Label(card_frame, text="Resultados da Pesquisa:", font=("Arial", 10, "bold"), fg="#94a3b8", bg="#202024", anchor="w").pack(fill=tk.X, pady=(15, 5))
        
        self.txt_resultado = tk.Text(
            card_frame, 
            font=("Courier New", 12), 
            bg="#1e1e24", 
            fg="#e2e8f0", 
            insertbackground="white",
            height=8, 
            relief="flat",
            bd=5,
            state=tk.DISABLED
        )
        self.txt_resultado.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Inicializa variável da velocidade mínima
        self.velocidade_minima_atual = 0.0
        
        # Frame da Calculadora de 80% (Speedtest)
        self.calc_frame = tk.LabelFrame(
            card_frame, 
            text="⚡ CÁLCULO DE CONFORMIDADE (MÍNIMO 80%)", 
            font=("Arial", 10, "bold"), 
            bg="#202024", 
            fg="#6366f1", 
            padx=15, 
            pady=10,
            relief="solid",
            bd=1
        )
        self.calc_frame.pack(fill=tk.X, pady=(10, 0))
        
        calc_input_frame = tk.Frame(self.calc_frame, bg="#202024")
        calc_input_frame.pack(fill=tk.X)
        
        tk.Label(calc_input_frame, text="Velocidade registrada no Speedtest (Mbps):", font=("Arial", 10, "bold"), fg="#e2e8f0", bg="#202024").pack(side=tk.LEFT, padx=(0, 10))
        
        self.ent_speedtest = tk.Entry(calc_input_frame, font=("Arial", 11), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", width=12, relief="flat", bd=2)
        self.ent_speedtest.pack(side=tk.LEFT, padx=5, ipady=2)
        self.ent_speedtest.bind("<Return>", lambda e: self.calcular_conformidade())
        
        btn_calcular = tk.Button(
            calc_input_frame, 
            text="🧮 CALCULAR", 
            command=self.calcular_conformidade, 
            bg="#6366f1", 
            fg="white", 
            font=("Arial", 10, "bold"), 
            relief="flat",
            padx=15,
            cursor="hand2"
        )
        btn_calcular.pack(side=tk.LEFT, padx=10, ipady=1)
        btn_calcular.bind("<Enter>", lambda e: btn_calcular.config(bg="#4f46e5"))
        btn_calcular.bind("<Leave>", lambda e: btn_calcular.config(bg="#6366f1"))
        
        self.lbl_resultado_calculo = tk.Label(
            self.calc_frame, 
            text="Insira o Código INEP acima, depois digite a velocidade do Speedtest e clique em Calcular.", 
            font=("Arial", 10, "italic"), 
            fg="#94a3b8", 
            bg="#202024", 
            anchor="w"
        )
        self.lbl_resultado_calculo.pack(fill=tk.X, pady=(8, 0))
        
        self.df = None
        self.frame.after(100, self.carregar_so_local)

    def carregar_so_local(self):
        """Carrega somente o cache local. Nao faz nenhuma requisicao a nuvem."""
        df_local = load_cache(CACHE_FILE)
        if not df_local.empty:
            col = "Código INEP" if "Código INEP" in df_local.columns else df_local.columns[0]
            df_local[col] = df_local[col].astype(str).str.split('.').str[0].str.strip()
            self.df = df_local
            self.lbl_status.config(text="✅ Cache local carregado offline. Atualizado via cache local.", fg="#10b981")
            self.ent_inep.focus()
        else:
            self.lbl_status.config(text="⚠️ Sem dados de cache local detectados. Clique em '🔄 Atualizar da Nuvem'.", fg="#eab308")

    def _sincronizar_nuvem(self):
        """Roda em background; atualiza UI de forma segura."""
        global GOOGLE_CONNECTION
        
        with GOOGLE_CONNECTION["lock"]:
            status = GOOGLE_CONNECTION["status"]
            gspread_client = GOOGLE_CONNECTION["gspread_client"]
            
        if status != "Conectado" or gspread_client is None:
            self.frame.after(0, lambda: self.lbl_status.config(
                text="❌ Falha: Google APIs nao autenticadas no background. Verifique sua rede/conexao.", fg="#ef4444"
            ))
            self.frame.after(0, lambda: self.btn_nuvem.config(state=tk.NORMAL))
            return
            
        try:
            log_message("INFO", "Sincronizando aba 'EACE' do Sheets com cache local...")
            config = load_config()
            df_nuvem = fetch_data_by_id(gspread_client, config["SPREADSHEET_ID"], "EACE")
            if not df_nuvem.empty:
                col = "Código INEP" if "Código INEP" in df_nuvem.columns else df_nuvem.columns[0]
                df_nuvem[col] = df_nuvem[col].astype(str).str.split('.').str[0].str.strip()
                self.df = df_nuvem
                save_cache(df_nuvem, CACHE_FILE)
                self.frame.after(0, lambda: self.lbl_status.config(
                    text="✅ ONLINE — Dados sincronizados e atualizados da Planilha Google!", fg="#10b981"
                ))
                self.frame.after(0, lambda: self.ent_inep.focus())
            else:
                self.frame.after(0, lambda: self.lbl_status.config(
                    text="⚠️ Nuvem retornou uma tabela vazia.", fg="#eab308"
                ))
        except Exception as e:
            log_message("ERROR", f"Erro ao sincronizar com a nuvem: {e}")
            if self.df is not None:
                self.frame.after(0, lambda: self.lbl_status.config(
                    text=f"⚠️ Falha na nuvem: usando cache local. Detalhe: {e}", fg="#ef4444"
                ))
            else:
                self.frame.after(0, lambda: self.lbl_status.config(
                    text="❌ Sem dados locais e falha critica na sincronizacao com a nuvem.", fg="#ef4444"
                ))
        self.frame.after(0, lambda: self.btn_nuvem.config(state=tk.NORMAL))

    def _forcar_atualizacao_nuvem(self):
        """Forca sincronizacao manual com a nuvem ao clicar no botão."""
        self.btn_nuvem.config(state=tk.DISABLED)
        self.lbl_status.config(text="⏳ Conectando e atualizando dados da planilha do Google...", fg="#3b82f6")
        threading.Thread(target=self._sincronizar_nuvem, daemon=True).start()

    def buscar_dados(self):
        if self.df is None or self.df.empty: 
            messagebox.showwarning("Aviso", "Nenhum banco de dados ou cache carregado.")
            return
        inep = self.ent_inep.get().strip()
        if not inep: 
            return
        
        self.txt_resultado.config(state=tk.NORMAL)
        self.txt_resultado.delete(1.0, tk.END)
        
        col_inep = "Código INEP" if "Código INEP" in self.df.columns else self.df.columns[0]
        res = self.df[self.df[col_inep] == inep]
        
        if res.empty:
            self.txt_resultado.insert(tk.END, f"❌ ERRO: Codigo INEP '{inep}' nao localizado no banco de dados.")
            log_message("INFO", f"Busca EACE: Codigo INEP '{inep}' nao localizado.")
            self.velocidade_minima_atual = 0.0
            self.lbl_resultado_calculo.config(
                text="Busque um INEP válido primeiro.",
                fg="#94a3b8"
            )
        else:
            dados_escola = res.iloc[0]
            
            # Obter a velocidade contratada/mínima
            vel_min_str = str(dados_escola.get('Velocidade DL Mínima (Mbps)', '0')).strip()
            # Se for N/A ou vazio, tenta a DL Ofertada
            if not vel_min_str or vel_min_str == 'N/A':
                vel_min_str = str(dados_escola.get('Velocidade DL Ofertada (Mbps)', '0')).strip()
                
            # Limpar string para converter em float
            vel_cleaned = ""
            for char in vel_min_str.replace(',', '.'):
                if char.isdigit() or char == '.':
                    vel_cleaned += char
                elif char.isspace() and vel_cleaned:
                    break
                    
            try:
                self.velocidade_minima_atual = float(vel_cleaned) if vel_cleaned else 0.0
            except ValueError:
                self.velocidade_minima_atual = 0.0

            texto = (
                f"IDENTIFICAÇÃO DO REGISTRO EACE:\n"
                f"===========================================================\n"
                f"CÓDIGO INEP:  {dados_escola.get('Código INEP', 'N/A')}\n"
                f"ESCOLA:       {dados_escola.get('Nome da Escola', 'N/A')}\n"
                f"UF / CIDADE:  {dados_escola.get('UF', 'N/A')} - {dados_escola.get('Município', 'N/A')}\n"
                f"ENDEREÇO:     {dados_escola.get('Endereço', 'N/A')}\n"
                f"LATITUDE:     {dados_escola.get('Latitude', 'N/A')}\n"
                f"LONGITUDE:    {dados_escola.get('Longitude', 'N/A')}\n"
                f"DL MÍNIMA:    {dados_escola.get('Velocidade DL Mínima (Mbps)', 'N/A')} Mbps\n"
                f"DL OFERTADA:  {dados_escola.get('Velocidade DL Ofertada (Mbps)', 'N/A')} Mbps\n"
                f"KIT WI-FI:    {dados_escola.get('Kit Wi-Fi', 'N/A')} APs\n"
                f"===========================================================\n"
            )
            self.txt_resultado.insert(tk.END, texto)
            log_message("INFO", f"Busca EACE efetuada: INEP {inep} - {dados_escola.get('Nome da Escola', 'N/A')}")
            
            minimo_80 = self.velocidade_minima_atual * 0.8
            self.lbl_resultado_calculo.config(
                text=f"✅ INEP Localizado! Meta: {self.velocidade_minima_atual:.1f} Mbps | Mínimo exigido (80%): {minimo_80:.1f} Mbps. Insira o Speedtest e clique em Calcular.",
                fg="#3b82f6"
            )
            
        self.txt_resultado.config(state=tk.DISABLED)
        self.ent_inep.delete(0, tk.END)

    def calcular_conformidade(self):
        if not hasattr(self, 'velocidade_minima_atual') or self.velocidade_minima_atual <= 0:
            messagebox.showwarning("Aviso", "Por favor, busque por um INEP válido primeiro para obter a meta de velocidade.")
            return
            
        speed_str = self.ent_speedtest.get().strip()
        if not speed_str:
            messagebox.showwarning("Aviso", "Por favor, insira a velocidade medida no Speedtest.")
            return
            
        # Limpar string do speedtest
        speed_cleaned = ""
        for char in speed_str.replace(',', '.'):
            if char.isdigit() or char == '.':
                speed_cleaned += char
                
        try:
            speed_val = float(speed_cleaned) if speed_cleaned else 0.0
        except ValueError:
            messagebox.showerror("Erro", "Formato de velocidade inválido. Insira apenas números.")
            return
            
        minimo_80 = self.velocidade_minima_atual * 0.8
        
        if speed_val >= minimo_80:
            self.lbl_resultado_calculo.config(
                text=f"🟢 APROVADO! O teste atingiu {speed_val:.1f} Mbps (mínimo de 80% exigido: {minimo_80:.1f} Mbps para a meta de {self.velocidade_minima_atual:.1f} Mbps).",
                fg="#10b981"
            )
        else:
            diferenca = minimo_80 - speed_val
            self.lbl_resultado_calculo.config(
                text=f"🔴 REPROVADO! O teste atingiu apenas {speed_val:.1f} Mbps (abaixo dos 80% exigidos de {minimo_80:.1f} Mbps. Falta {diferenca:.1f} Mbps).",
                fg="#ef4444"
            )

# --- CLASSE DE AUTOMAÇÃO EACE ---
class AutomacaoEACE:
    def __init__(self, parent_frame, main_app_callback):
        self.parent_frame = parent_frame
        self.main_app_callback = main_app_callback
        self.frame = tk.Frame(parent_frame, bg="#121214")
        self.frame.pack(fill=tk.BOTH, expand=True)
        
        if hasattr(self.parent_frame, 'winfo_toplevel'):
            self.parent_frame.winfo_toplevel().geometry("900x880")
        
        header_frame = tk.Frame(self.frame, bg="#121214")
        header_frame.pack(fill=tk.X, pady=(15, 10))
        
        tk.Label(header_frame, text="📝 AUTOMAÇÃO EACE (GERADOR DOCX)", font=("Arial", 16, "bold"), fg="#e2e8f0", bg="#121214").pack(side=tk.LEFT, padx=15)
        
        btn_voltar = tk.Button(
            header_frame, 
            text="⬅️ VOLTAR AO MENU", 
            command=self.main_app_callback, 
            bg="#4b5563", 
            fg="white", 
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=10,
            pady=5
        )
        btn_voltar.pack(side=tk.RIGHT, padx=15)
        btn_voltar.bind("<Enter>", lambda e: btn_voltar.config(bg="#374151"))
        btn_voltar.bind("<Leave>", lambda e: btn_voltar.config(bg="#4b5563"))
        
        self.local_output_dir = os.path.join(os.path.expanduser("~"), "Documents", "Revisado-Enviar EACE")
        self.df_principal = None
        self.infantil_ineps = []
        
        canvas_container = tk.Frame(self.frame, bg="#121214")
        canvas_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.canvas = tk.Canvas(canvas_container, bg="#202024", highlightthickness=0)
        scrollbar = ttk.Scrollbar(canvas_container, orient="vertical", command=self.canvas.yview)
        
        container = tk.Frame(self.canvas, bg="#202024", padx=15, pady=15)
        self.canvas_window = self.canvas.create_window((0, 0), window=container, anchor="nw")
        
        container.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfig(self.canvas_window, width=e.width))
        self.canvas.configure(yscrollcommand=scrollbar.set)
        
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        left_pane = tk.Frame(container, bg="#202024")
        left_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        right_pane = tk.Frame(container, bg="#202024", width=300)
        right_pane.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(10, 0))
        
        self.lbl_status = tk.Label(left_pane, text="Carregando base local...", font=("Arial", 10, "italic"), fg="#94a3b8", bg="#202024", anchor="w")
        self.lbl_status.pack(fill=tk.X, pady=(0, 10))
        
        # Carregar config inicial
        config = load_config()
        
        # Ocultar campos de email em um frame mais discreto
        self.frame_emails = tk.Frame(left_pane, bg="#202024")
        self.frame_emails.pack(fill=tk.X, pady=(0, 15))
        tk.Label(self.frame_emails, text="Para:", font=("Arial", 9, "bold"), fg="#94a3b8", bg="#202024").grid(row=0, column=0, sticky="w")
        self.ent_to = tk.Entry(self.frame_emails, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, width=42)
        self.ent_to.grid(row=0, column=1, sticky="w", padx=5, pady=2, ipady=2)
        self.ent_to.insert(0, config.get("EMAIL_TO", ""))
        
        tk.Label(self.frame_emails, text="Cc:", font=("Arial", 9, "bold"), fg="#94a3b8", bg="#202024").grid(row=1, column=0, sticky="nw", pady=2)
        self.ent_cc = tk.Text(self.frame_emails, height=3, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, width=42, wrap=tk.WORD)
        self.ent_cc.grid(row=1, column=1, sticky="w", padx=5, pady=2)
        self.ent_cc.insert("1.0", config.get("EMAIL_CC", ""))
        
        btn_salvar_emails = tk.Button(
            self.frame_emails,
            text="💾 Salvar",
            command=self.salvar_emails_padrao,
            bg="#3b82f6",
            fg="white",
            font=("Arial", 8, "bold"),
            relief="flat",
            cursor="hand2"
        )
        btn_salvar_emails.grid(row=0, column=2, rowspan=2, padx=(10, 0), sticky="ns", pady=2)
        btn_salvar_emails.bind("<Enter>", lambda e, b=btn_salvar_emails: b.config(bg="#2563eb"))
        btn_salvar_emails.bind("<Leave>", lambda e, b=btn_salvar_emails: b.config(bg="#3b82f6"))
        
        tk.Label(left_pane, text="Código INEP da Escola:", font=("Arial", 10, "bold"), fg="#e2e8f0", bg="#202024", anchor="w").pack(fill=tk.X, pady=(10, 2))
        self.ent_inep = tk.Entry(left_pane, font=("Arial", 12), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2)
        self.ent_inep.pack(fill=tk.X, ipady=4, pady=(0, 10))
        
        tk.Label(left_pane, text="Nova Quantidade de APs (Pleito):", font=("Arial", 10, "bold"), fg="#e2e8f0", bg="#202024", anchor="w").pack(fill=tk.X, pady=(10, 2))
        self.ent_novos_aps = tk.Entry(left_pane, font=("Arial", 12), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2)
        self.ent_novos_aps.pack(fill=tk.X, ipady=4, pady=(0, 15))
        
        self.frame_adicionais = tk.LabelFrame(left_pane, text="Adicionais de Infraestrutura (Qtd)", font=("Arial", 10, "bold"), bg="#202024", fg="#e2e8f0", padx=10, pady=10)
        self.frame_adicionais.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(self.frame_adicionais, text="Switch:", font=("Arial", 9), bg="#202024", fg="#e2e8f0").grid(row=0, column=0, padx=5, pady=2, sticky="e")
        self.ent_switch = tk.Entry(self.frame_adicionais, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, width=10)
        self.ent_switch.grid(row=0, column=1, padx=5, pady=2, sticky="w")
        self.ent_switch.insert(0, "0")
        
        tk.Label(self.frame_adicionais, text="Rack:", font=("Arial", 9), bg="#202024", fg="#e2e8f0").grid(row=1, column=0, padx=5, pady=2, sticky="e")
        self.ent_rack = tk.Entry(self.frame_adicionais, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, width=10)
        self.ent_rack.grid(row=1, column=1, padx=5, pady=2, sticky="w")
        self.ent_rack.insert(0, "0")
        
        tk.Label(self.frame_adicionais, text="Nobreak:", font=("Arial", 9), bg="#202024", fg="#e2e8f0").grid(row=2, column=0, padx=5, pady=2, sticky="e")
        self.ent_nobreak = tk.Entry(self.frame_adicionais, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, width=10)
        self.ent_nobreak.grid(row=2, column=1, padx=5, pady=2, sticky="w")
        self.ent_nobreak.insert(0, "0")
        
        # Seleção de PDF local
        self.selected_pdf_path = ""
        tk.Label(left_pane, text="PDF de Evidência Local (Opcional):", font=("Arial", 10, "bold"), fg="#e2e8f0", bg="#202024", anchor="w").pack(fill=tk.X, pady=(10, 2))
        pdf_select_frame = tk.Frame(left_pane, bg="#202024")
        pdf_select_frame.pack(fill=tk.X, pady=(0, 15))
        
        self.ent_pdf_path = tk.Entry(pdf_select_frame, font=("Arial", 10), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, state="disabled")
        self.ent_pdf_path.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4, padx=(0, 5))
        
        btn_sel_pdf = tk.Button(
            pdf_select_frame,
            text="📎 Selecionar PDF...",
            command=self.selecionar_pdf_local,
            bg="#3b82f6",
            fg="white",
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2"
        )
        btn_sel_pdf.pack(side=tk.LEFT, padx=2, ipady=2)
        btn_sel_pdf.bind("<Enter>", lambda e: btn_sel_pdf.config(bg="#2563eb"))
        btn_sel_pdf.bind("<Leave>", lambda e: btn_sel_pdf.config(bg="#3b82f6"))
        
        btn_limpar_pdf = tk.Button(
            pdf_select_frame,
            text="❌ Limpar",
            command=self.limpar_pdf_local,
            bg="#ef4444",
            fg="white",
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2"
        )
        btn_limpar_pdf.pack(side=tk.LEFT, padx=(2, 0), ipady=2)
        btn_limpar_pdf.bind("<Enter>", lambda e: btn_limpar_pdf.config(bg="#dc2626"))
        btn_limpar_pdf.bind("<Leave>", lambda e: btn_limpar_pdf.config(bg="#ef4444"))
        
        self.var_enviar_email = tk.BooleanVar(value=True)
        self.chk_email = tk.Checkbutton(
            left_pane, 
            text="Enviar e-mail automático com PDF anexado", 
            variable=self.var_enviar_email, 
            font=("Arial", 11, "bold"), 
            bg="#202024",
            fg="#ef4444",
            selectcolor="#1e1e24",
            activebackground="#202024",
            activeforeground="#ef4444"
        )
        self.chk_email.pack(anchor="w", pady=10)
        
        self.btn_previa = tk.Button(
            left_pane,
            text="👀 VISUALIZAR PRÉVIA DO E-MAIL",
            command=self.mostrar_previa,
            bg="#f59e0b",
            fg="white",
            font=("Arial", 11, "bold"),
            height=1,
            relief="flat",
            cursor="hand2"
        )
        self.btn_previa.pack(fill=tk.X, pady=(0, 10))
        
        self.btn_gerar = tk.Button(
            left_pane, 
            text="🚀 INICIAR AUTOMAÇÃO DOCX/EMAIL", 
            command=self.gerar_documento, 
            bg="#10b981", 
            fg="white", 
            font=("Arial", 12, "bold"), 
            height=2,
            relief="flat",
            cursor="hand2"
        )
        self.btn_gerar.pack(fill=tk.X, pady=15)
        self.btn_gerar.bind("<Enter>", lambda e: self.btn_gerar.config(bg="#059669"))
        self.btn_gerar.bind("<Leave>", lambda e: self.btn_gerar.config(bg="#10b981"))
        
        info_box = tk.LabelFrame(right_pane, text=" Ações & Informações ", font=("Arial", 10, "bold"), bg="#202024", fg="#94a3b8", padx=10, pady=10, relief="solid", bd=1)
        info_box.pack(fill=tk.BOTH, expand=True)
        
        self.btn_nuvem = tk.Button(
            info_box, 
            text="🔄 Sincronizar com a Nuvem", 
            command=self._forcar_atualizacao_nuvem, 
            bg="#3b82f6", 
            fg="white", 
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2"
        )
        self.btn_nuvem.pack(fill=tk.X, pady=(5, 15), ipady=3)
        self.btn_nuvem.bind("<Enter>", lambda e: self.btn_nuvem.config(bg="#2563eb"))
        self.btn_nuvem.bind("<Leave>", lambda e: self.btn_nuvem.config(bg="#3b82f6"))
        
        instrucoes = (
            "Como funciona a automacao:\n\n"
            "1. Busca a escola no cache local.\n"
            "2. Identifica se e escola infantil.\n"
            "3. Cria o Word (DOCX) com a estrutura "
            "e a tabela comparativa do pleito.\n"
            "4. Salva localmente, tenta na Rede Local "
            "e faz upload no Google Drive.\n"
            "5. Registra o pleito na aba 'UP_DOWN' "
            "do Sheets.\n"
            "6. Anexa a evidencia PDF baixada da pasta "
            "do Drive e envia o e-mail formal."
        )
        lbl_instrucoes = tk.Label(info_box, text=instrucoes, font=("Arial", 9), fg="#94a3b8", bg="#202024", justify=tk.LEFT, anchor="nw")
        lbl_instrucoes.pack(fill=tk.BOTH)
        
        tk.Label(info_box, text="Bloco de Notas (Contatos/Rascunho):", font=("Arial", 9, "bold"), fg="#e2e8f0", bg="#202024", anchor="w").pack(fill=tk.X, pady=(15, 5))
        self.txt_notas = tk.Text(info_box, height=8, font=("Arial", 9), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2, wrap=tk.WORD)
        self.txt_notas.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        self.txt_notas.insert("1.0", config.get("BLOCO_NOTAS", ""))
        
        self.inicializar_sistema()

    def salvar_emails_padrao(self):
        to_email = self.ent_to.get().strip()
        cc_email = self.ent_cc.get("1.0", tk.END).strip().replace("\n", ",")
        bloco_notas = self.txt_notas.get("1.0", tk.END).strip()
        
        try:
            config = load_config()
            config["EMAIL_TO"] = to_email
            config["EMAIL_CC"] = cc_email
            config["BLOCO_NOTAS"] = bloco_notas
            save_config(config)
            log_message("SUCCESS", "E-mails e bloco de notas atualizados pelo botao rapido da interface.")
            messagebox.showinfo("Sucesso", "Configurações e Bloco de Notas salvos com sucesso!")
        except Exception as e:
            log_message("ERROR", f"Falha ao salvar configuracoes: {e}")
            messagebox.showerror("Erro", f"Ocorreu um erro ao tentar salvar:\n{e}")

    def selecionar_pdf_local(self):
        caminho = filedialog.askopenfilename(filetypes=[("Arquivos PDF", "*.pdf")])
        if caminho:
            self.selected_pdf_path = caminho
            self.ent_pdf_path.config(state="normal")
            self.ent_pdf_path.delete(0, tk.END)
            self.ent_pdf_path.insert(0, os.path.basename(caminho))
            self.ent_pdf_path.config(state="disabled")
            log_message("INFO", f"PDF local de evidencia selecionado: {caminho}")

    def limpar_pdf_local(self):
        self.selected_pdf_path = ""
        self.ent_pdf_path.config(state="normal")
        self.ent_pdf_path.delete(0, tk.END)
        self.ent_pdf_path.config(state="disabled")
        log_message("INFO", "PDF local de evidencia desmarcado.")

    def get_field_safely(self, dados, key, default="não informado"):
        if dados is None:
            return default
        val = dados.get(key)
        if pd.isna(val) or val is None:
            return default
        val_str = str(val).strip()
        if not val_str or val_str.lower() in ("nan", "null", "n/a"):
            return default
        return val_str

    def inicializar_sistema(self):
        """Carrega somente o cache local. Nao faz nenhuma requisicao a nuvem."""
        df_local = load_cache(CACHE_FILE)
        if not df_local.empty:
            col = "Código INEP" if "Código INEP" in df_local.columns else df_local.columns[0]
            df_local[col] = df_local[col].astype(str).str.split('.').str[0].str.strip()
            self.df_principal = df_local

        df_inf_local = load_cache(CACHE_FILE_INFANTIL)
        if not df_inf_local.empty:
            col_inf = "Código INEP" if "Código INEP" in df_inf_local.columns else df_inf_local.columns[0]
            self.infantil_ineps = df_inf_local[col_inf].astype(str).str.split('.').str[0].str.strip().tolist()

        if not df_local.empty:
            self.lbl_status.config(text="✅ Base de dados do cache local carregada.", fg="#10b981")
        else:
            self.lbl_status.config(text="⚠️ Sem dados de cache local. Clique em '🔄 Sincronizar com a Nuvem'.", fg="#eab308")

    def _sincronizar_nuvem_automacao(self):
        """Roda em background; atualiza UI via thread."""
        global GOOGLE_CONNECTION
        
        with GOOGLE_CONNECTION["lock"]:
            status = GOOGLE_CONNECTION["status"]
            gspread_client = GOOGLE_CONNECTION["gspread_client"]
            
        if status != "Conectado" or gspread_client is None:
            self.frame.after(0, lambda: self.lbl_status.config(
                text="❌ Falha: Google APIs offline. Autenticacao pendente.", fg="#ef4444"
            ))
            self.frame.after(0, lambda: self.btn_nuvem.config(state=tk.NORMAL))
            return
            
        try:
            log_message("INFO", "Sincronizando planilhas EACE e ESCOLA_INFANTIL...")
            config = load_config()
            df_n = fetch_data_by_id(gspread_client, config["SPREADSHEET_ID"], "EACE")
            df_inf_n = fetch_data_by_id(gspread_client, config["SPREADSHEET_ID"], "ESCOLA_INFANTIL")
            
            if not df_n.empty:
                col = "Código INEP" if "Código INEP" in df_n.columns else df_n.columns[0]
                df_n[col] = df_n[col].astype(str).str.split('.').str[0].str.strip()
                self.df_principal = df_n
                save_cache(df_n, CACHE_FILE)
                
            if not df_inf_n.empty:
                col_inf = "Código INEP" if "Código INEP" in df_inf_n.columns else df_inf_n.columns[0]
                self.infantil_ineps = df_inf_n[col_inf].astype(str).str.split('.').str[0].str.strip().tolist()
                save_cache(df_inf_n, CACHE_FILE_INFANTIL)
                
            self.frame.after(0, lambda: self.lbl_status.config(
                text="✅ ONLINE — Dados do cache sincronizados com o Sheets Google!", fg="#10b981"
            ))
        except Exception as e:
            log_message("ERROR", f"Falha na sincronizacao da automacao: {e}")
            if self.df_principal is not None:
                self.frame.after(0, lambda: self.lbl_status.config(
                    text="⚠️ Nuvem indisponivel — usando cache offline", fg="#ef4444"
                ))
            else:
                self.frame.after(0, lambda: self.lbl_status.config(
                    text="❌ Sem dados locais e falha de rede na sincronizacao", fg="#ef4444"
                ))
        self.frame.after(0, lambda: self.btn_nuvem.config(state=tk.NORMAL))

    def _forcar_atualizacao_nuvem(self):
        """Forca sincronizacao manual com a nuvem ao clicar no botão."""
        self.btn_nuvem.config(state=tk.DISABLED)
        self.lbl_status.config(text="⏳ Conectando a nuvem e baixando dados...", fg="#3b82f6")
        threading.Thread(target=self._sincronizar_nuvem_automacao, daemon=True).start()

    def upload_to_drive(self, local_path, filename):
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            drive_service = GOOGLE_CONNECTION["drive_service"]
            
        if not drive_service: 
            return False, "Google Drive Service nao autenticado."
            
        try:
            config = load_config()
            media = MediaFileUpload(local_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
            file = drive_service.files().create(
                body={'name': filename, 'parents': [config["GOOGLE_DRIVE_FOLDER_ID"]]}, 
                media_body=media, 
                fields='id', 
                supportsAllDrives=True
            ).execute()
            return True, f"ID: {file.get('id')}"
        except Exception as e: 
            return False, str(e)

    def upload_pdf_evidencia(self, local_pdf_path, inep_str):
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            drive_service = GOOGLE_CONNECTION["drive_service"]
            
        if not drive_service: 
            return False, "Google Drive Service nao autenticado."
            
        try:
            config = load_config()
            folder_id_evidencias = config["ID_PASTA_EVIDENCIAS_DRIVE"]
            
            # Buscar pasta do INEP
            q_f = f"name='{inep_str}' and '{folder_id_evidencias}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
            res_f = drive_service.files().list(q=q_f, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
            folders = res_f.get('files', [])
            
            if not folders:
                log_message("INFO", f"Pasta do INEP {inep_str} nao encontrada. Criando nova pasta...")
                file_metadata = {
                    'name': inep_str,
                    'mimeType': 'application/vnd.google-apps.folder',
                    'parents': [folder_id_evidencias]
                }
                folder = drive_service.files().create(body=file_metadata, fields='id', supportsAllDrives=True).execute()
                f_id = folder.get('id')
            else:
                f_id = folders[0]['id']
                
            # Upload do PDF local para a pasta do INEP
            filename = os.path.basename(local_pdf_path)
            file_metadata = {
                'name': filename,
                'parents': [f_id]
            }
            media = MediaFileUpload(local_pdf_path, mimetype='application/pdf', resumable=True)
            file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id',
                supportsAllDrives=True
            ).execute()
            
            log_message("SUCCESS", f"PDF de evidencia enviado para o Drive na pasta {inep_str} com sucesso (ID: {file.get('id')})")
            return True, f"✅ GOOGLE DRIVE (PDF): Upload realizado com sucesso."
        except Exception as e:
            log_message("ERROR", f"Falha no upload do PDF para o Drive: {e}")
            email_guia = "automacao-eace@automacao-eace.iam.gserviceaccount.com"
            try:
                creds_path = os.path.join(get_base_path(), "credentials.json")
                if os.path.exists(creds_path):
                    with open(creds_path, 'r', encoding='utf-8') as f:
                        creds_data = json.load(f)
                        email_guia = creds_data.get("client_email", email_guia)
            except:
                pass
            return False, f"❌ GOOGLE DRIVE (PDF): Falha de permissão. Compartilhe a pasta mãe de evidências no Drive com o e-mail '{email_guia}' como Editor."

    def mostrar_previa(self):
        import tempfile
        import webbrowser
        import os
        
        try:
            qtd_switch = int(self.ent_switch.get().strip() or "0")
        except: qtd_switch = 0
        try:
            qtd_rack = int(self.ent_rack.get().strip() or "0")
        except: qtd_rack = 0
        try:
            qtd_nobreak = int(self.ent_nobreak.get().strip() or "0")
        except: qtd_nobreak = 0
        
        adicionais = []
        if qtd_switch > 0:
            adicionais.append(f"{qtd_switch}x Switch")
        if qtd_rack > 0:
            adicionais.append(f"{qtd_rack}x Rack")
        if qtd_nobreak > 0:
            adicionais.append(f"{qtd_nobreak}x Nobreak")
        
        html_adicionais = ""
        if adicionais:
            texto_adic = ", ".join(adicionais)
            html_adicionais = f"""
            <p>Ressaltamos também que, para a viabilização da estrutura no ambiente, será necessária a implementação dos seguintes ativos adicionais: <b>{texto_adic}</b>.</p>
            """
            
        inep_str = self.ent_inep.get().strip()
        if not inep_str:
            messagebox.showwarning("Aviso", "Preencha o código INEP primeiro para gerar a prévia.")
            return
            
        dados_escola = None
        if hasattr(self, 'df_principal') and self.df_principal is not None:
            col = "Código INEP" if "Código INEP" in self.df_principal.columns else self.df_principal.columns[0]
            linha = self.df_principal[self.df_principal[col] == inep_str]
            if not linha.empty:
                dados_escola = linha.iloc[0].to_dict()
                
        if not dados_escola:
            messagebox.showwarning("Aviso", "Escola não encontrada no cache carregado. Impossível gerar a prévia do e-mail.")
            return
            
        nome_escola = self.get_field_safely(dados_escola, 'Nome da Escola')
        es_infantil = "SIM" if hasattr(self, 'infantil_ineps') and inep_str in self.infantil_ineps else "NÃO"
        
        aps_atuais = str(self.get_field_safely(dados_escola, 'Kit Wi-Fi'))
        novos_aps = self.ent_novos_aps.get().strip()
        if not novos_aps: novos_aps = "X"
        
        try: 
            qtd_atual = int(aps_atuais.split()[0])
            qtd_nova = int(novos_aps.split()[0])
            tipo = "UPGRADE" if qtd_nova > qtd_atual else "DOWNGRADE"
        except:
            tipo = "UPGRADE"
        
        to_email = self.ent_to.get().strip()
        cc_email = self.ent_cc.get("1.0", tk.END).strip().replace("\n", ",")
        
        html_body = f"""
        <html>
        <head><meta charset="utf-8"></head>
        <body style="font-family: Arial, sans-serif; font-size: 14px; color: #333333; background: #ffffff; padding: 20px;">
            <div style="background-color: #f8fafc; padding: 15px; border-bottom: 2px solid #e2e8f0; margin-bottom: 20px; font-family: monospace; font-size: 13px;">
                <p style="margin: 3px 0;"><b>De:</b> NOC - Núcleo de Operações de Rede</p>
                <p style="margin: 3px 0;"><b>Para:</b> {to_email}</p>
                <p style="margin: 3px 0;"><b>Cc:</b> {cc_email}</p>
                <p style="margin: 3px 0; color: #2563eb;"><b>Assunto:</b> ARAUJO E ALMEIDA - Pleito Alteração Solução RI - INEP {inep_str}</p>
            </div>
            <p>Prezados,</p>
            <p>Gostaria de solicitar a autorização da EACE para realização de <b>{tipo.lower()}</b> na quantidade de Access Points da escola abaixo:</p>
            <ul style="line-height: 1.6;">
                <li><b>Escola:</b> {nome_escola}</li>
                <li><b>Código INEP / Identificação:</b> {inep_str}</li>
                <li><b>Município:</b> {self.get_field_safely(dados_escola, 'Município')} - {self.get_field_safely(dados_escola, 'UF')}</li>
                <li><b>Endereço:</b> {self.get_field_safely(dados_escola, 'Endereço')}</li>
            </ul>
            <br>
            <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; text-align: center; font-size: 11px; width: 100%; border-color: #cccccc;">
                <tr style="background-color: #f2f2f2; font-weight: bold;">
                    <th>FASE</th><th>INEP</th><th>ESCOLA</th><th>ESTADO</th><th>CIDADE</th><th>KIT PREVISTO</th><th>PLEITO</th><th>KIT SUGERIDO</th><th>LATITUDE</th><th>LONGITUDE</th><th>ESCOLA INFANTIL</th><th>JUSTIFICATIVA</th>
                </tr>
                <tr>
                    <td>5</td><td>{inep_str}</td><td>{nome_escola}</td><td>{self.get_field_safely(dados_escola, 'UF')}</td><td>{self.get_field_safely(dados_escola, 'Município')}</td><td>{aps_atuais}</td><td style="font-weight: bold; color: {'#27ae60' if tipo == 'UPGRADE' else '#c0392b'};">{tipo}</td><td>{novos_aps}</td><td>{self.get_field_safely(dados_escola, 'Latitude')}</td><td>{self.get_field_safely(dados_escola, 'Longitude')}</td><td>{es_infantil}</td><td>Adequação à necessidade da escola</td>
                </tr>
            </table>
            <br>
            <p>Conforme consta na lista atual, a escola possui previsão de <b>{aps_atuais} APs</b>. Porém, após validação da necessidade do ambiente, identificamos que será necessário o <b>{tipo.lower()} para {novos_aps} APs</b>, a fim de atender de forma adequada a cobertura e o funcionamento da unidade escolar.</p>
            {html_adicionais}
            <p>Dessa forma, solicitamos a autorização formal da EACE para atualização da quantidade de equipamentos desta escola, passando de <b>{aps_atuais} AP</b> para <b>{novos_aps} APs</b>.</p>
            <p>Ficamos no aguardo da aprovação para prosseguir com o atendimento em campo.</p>
            <br>
            <p>Atenciosamente,</p>
            <p><b>NOC - Núcleo de Operações de Rede</b></p>
        </body></html>"""
        
        try:
            fd, temp_path = tempfile.mkstemp(suffix=".html", prefix="previa_email_")
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(html_body)
            
            # Delega para o Windows abrir no navegador padrão com a aba ativa
            if hasattr(os, 'startfile'):
                os.startfile(temp_path)
            else:
                webbrowser.open(temp_path, new=2)
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir a prévia no navegador: {e}")

    def atualizar_relatorio(self, dados):
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            gspread_client = GOOGLE_CONNECTION["gspread_client"]
            
        if gspread_client is None:
            log_message("WARNING", "Nao foi possivel registrar na planilha de relatorios: Google Sheets desconectado.")
            return False
            
        try:
            config = load_config()
            spreadsheet = gspread_client.open_by_key(config["SPREADSHEET_ID"])
            try: 
                ws = spreadsheet.worksheet("UP_DOWN")
            except: 
                ws = spreadsheet.add_worksheet(title="UP_DOWN", rows="1000", cols="10")
                ws.append_row(["INEP", "Escola", "Status", "Tipo", "Kit Previsto", "Kit Sugerido", "Observação", "OBS EACE"])
                
            inep = str(dados["INEP"]).strip()
            vals = ws.get_all_values()
            idx = -1
            for i, row in enumerate(vals):
                if i > 0 and len(row) > 0 and str(row[0]).strip() == inep: 
                    idx = i + 1
                    break
            if idx != -1:
                if messagebox.askyesno("Atualizar?", f"Codigo INEP {inep} ja existe na planilha de controle. Deseja sobrescrever os dados?"):
                    ws.update_cell(idx, 3, dados["Status"])
                    ws.update_cell(idx, 4, dados["Tipo"])
                    ws.update_cell(idx, 6, dados["Kit Sugerido"])
                    ws.update_cell(idx, 7, dados["Observação"])
                    ws.update_cell(idx, 8, dados["OBS EACE"])
                    log_message("INFO", f"Registro INEP {inep} atualizado na planilha 'UP_DOWN'.")
                    return True
                return False
            ws.append_row([dados["INEP"], dados["Escola"], dados["Status"], dados["Tipo"], dados["Kit Previsto"], dados["Kit Sugerido"], dados["Observação"], dados["OBS EACE"]])
            log_message("INFO", f"Novo registro INEP {inep} adicionado a planilha 'UP_DOWN'.")
            return True
        except Exception as e: 
            show_detailed_error("Erro no Relatorio Sheets", "Falha critica ao atualizar registro da planilha.", e)
            return False

    def verificar_duplicidade_gmail(self, inep):
        """Retorna True se um e-mail com este INEP ja foi enviado e barra o processo."""
        try:
            gmail_service = authenticate_gmail_readonly()
            inep_str = str(inep).strip()
            query = f"subject:{inep_str} in:sent"
            
            self.lbl_status.config(text="Verificando duplicidade no Gmail...")
            self.frame.update()
            
            results = gmail_service.users().messages().list(userId='me', q=query).execute()
            messages = results.get('messages', [])
            
            if messages:
                # Opcional: extrair a data exata da primeira mensagem
                msg_data = gmail_service.users().messages().get(userId='me', id=messages[0]['id'], format='metadata', metadataHeaders=['Date', 'Subject']).execute()
                headers = msg_data['payload']['headers']
                assunto = next((h['value'] for h in headers if h['name'] == 'Subject'), "Desconhecido")
                data_envio = next((h['value'] for h in headers if h['name'] == 'Date'), "Desconhecida")
                
                messagebox.showerror(
                    "❌ PLEITO JÁ ENVIADO!",
                    f"A automação foi BLOQUEADA.\n\n"
                    f"Você já enviou um pleito para este INEP ({inep_str}) no passado!\n\n"
                    f"📌 Assunto: {assunto}\n"
                    f"🕒 Data do Envio: {data_envio}",
                    parent=self.frame
                )
                self.btn_gerar.config(text="🚀 INICIAR AUTOMAÇÃO DOCX/EMAIL", state="normal", bg="#10b981")
                self.lbl_status.config(text="Bloqueado por duplicidade.", fg="#ef4444")
                return True
            return False
        except Exception as e:
            log_message("WARNING", f"Falha ao verificar duplicidade de e-mail: {e}")
            return False

    def enviar_email_automatico(self, inep, nome_escola, aps_atuais, novos_aps, tipo, es_infantil, dados_escola, temp_dir, pdf_local_path=None):
        """Logica robusta de e-mail e integracao do Drive Compartilhado."""
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            drive_service = GOOGLE_CONNECTION["drive_service"]
            
        try:
            gmail_service = authenticate_gmail()
            inep_str = str(inep).strip()
        except Exception as e: 
            log_message("ERROR", f"Falha de autenticacao no Gmail: {e}")
            return f"❌ ERRO GMAIL (Autenticacao): {str(e)}"

        pdf_path, pdf_name_drive, detalhe_busca = None, None, ""
        
        if pdf_local_path and os.path.exists(pdf_local_path):
            pdf_path = pdf_local_path
            pdf_name_drive = os.path.basename(pdf_local_path)
            detalhe_busca = "✅ PDF ANEXADO COM SUCESSO"
            log_message("INFO", f"Usando PDF local para o anexo do e-mail: {pdf_name_drive}")
        elif drive_service:
            try:
                config = load_config()
                folder_id_evidencias = config["ID_PASTA_EVIDENCIAS_DRIVE"]
                q_f = f"name='{inep_str}' and '{folder_id_evidencias}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
                res_f = drive_service.files().list(q=q_f, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
                folders = res_f.get('files', [])
                if not folders: 
                    detalhe_busca = "(Pasta INEP nao localizada no Drive de Evidencias)"
                    log_message("WARNING", f"Pasta do INEP {inep_str} nao encontrada no Google Drive de Evidencias.")
                else:
                    f_id = folders[0]['id']
                    q_p = f"'{f_id}' in parents and mimeType='application/pdf' and trashed=false"
                    res_p = drive_service.files().list(q=q_p, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
                    pdfs = res_p.get('files', [])
                    if not pdfs: 
                        detalhe_busca = "(PDF de evidencia nao localizado na pasta do INEP)"
                        log_message("WARNING", f"Nenhum PDF encontrado na pasta do INEP {inep_str}.")
                    else:
                        pdf_id, pdf_name_drive = pdfs[0]['id'], pdfs[0]['name']
                        pdf_path = os.path.join(temp_dir, pdf_name_drive)
                        log_message("INFO", f"Baixando PDF de evidencia do Drive: {pdf_name_drive}")
                        with io.FileIO(pdf_path, 'wb') as fh:
                            dl = MediaIoBaseDownload(fh, drive_service.files().get_media(fileId=pdf_id))
                            done = False
                            while not done: 
                                _, done = dl.next_chunk()
                        detalhe_busca = "✅ PDF ANEXADO COM SUCESSO"
            except Exception as e: 
                detalhe_busca = f"⚠️ Falha ao buscar PDF no Drive: {str(e)}"
                log_message("ERROR", f"Falha no download da evidencia PDF do Drive: {e}")
        else:
            detalhe_busca = "⚠️ Google Drive offline - PDF nao localizado"

        try:
            config = load_config()
            
            try:
                qtd_switch = int(self.ent_switch.get().strip() or "0")
            except: qtd_switch = 0
            try:
                qtd_rack = int(self.ent_rack.get().strip() or "0")
            except: qtd_rack = 0
            try:
                qtd_nobreak = int(self.ent_nobreak.get().strip() or "0")
            except: qtd_nobreak = 0
            
            adicionais = []
            if qtd_switch > 0:
                adicionais.append(f"{qtd_switch}x Switch")
            if qtd_rack > 0:
                adicionais.append(f"{qtd_rack}x Rack")
            if qtd_nobreak > 0:
                adicionais.append(f"{qtd_nobreak}x Nobreak")
            
            html_adicionais = ""
            if adicionais:
                texto_adic = ", ".join(adicionais)
                html_adicionais = f"""
                <p>Ressaltamos também que, para a viabilização da estrutura no ambiente, será necessária a implementação dos seguintes ativos adicionais: <b>{texto_adic}</b>.</p>
                """

            to_email = self.ent_to.get().strip()
            cc_email = self.ent_cc.get("1.0", tk.END).strip().replace("\n", ",")

            msg = EmailMessage()
            if to_email:
                msg['To'] = to_email
            if cc_email: 
                msg['Cc'] = cc_email
                
            msg['Subject'] = f"ARAUJO E ALMEIDA - Pleito Alteração Solução RI - INEP {inep_str}"
            html_body = f"""
            <html><body style="font-family: Arial, sans-serif; font-size: 14px; color: #333333;">
                <p>Prezados,</p>
                <p>Gostaria de solicitar a autorização da EACE para realização de <b>{tipo.lower()}</b> na quantidade de Access Points da escola abaixo:</p>
                <ul style="line-height: 1.6;">
                    <li><b>Escola:</b> {nome_escola}</li>
                    <li><b>Código INEP / Identificação:</b> {inep_str}</li>
                    <li><b>Município:</b> {self.get_field_safely(dados_escola, 'Município')} - {self.get_field_safely(dados_escola, 'UF')}</li>
                    <li><b>Endereço:</b> {self.get_field_safely(dados_escola, 'Endereço')}</li>
                </ul>
                <br>
                <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; text-align: center; font-size: 11px; width: 100%; border-color: #cccccc;">
                    <tr style="background-color: #f2f2f2; font-weight: bold;">
                        <th>FASE</th><th>INEP</th><th>ESCOLA</th><th>ESTADO</th><th>CIDADE</th><th>KIT PREVISTO</th><th>PLEITO</th><th>KIT SUGERIDO</th><th>LATITUDE</th><th>LONGITUDE</th><th>ESCOLA INFANTIL</th><th>JUSTIFICATIVA</th>
                    </tr>
                    <tr>
                        <td>5</td><td>{inep_str}</td><td>{nome_escola}</td><td>{self.get_field_safely(dados_escola, 'UF')}</td><td>{self.get_field_safely(dados_escola, 'Município')}</td><td>{aps_atuais}</td><td style="font-weight: bold; color: {'#27ae60' if tipo == 'UPGRADE' else '#c0392b'};">{tipo}</td><td>{novos_aps}</td><td>{self.get_field_safely(dados_escola, 'Latitude')}</td><td>{self.get_field_safely(dados_escola, 'Longitude')}</td><td>{es_infantil}</td><td>Adequação à necessidade da escola</td>
                    </tr>
                </table>
                <br>
                <p>Conforme consta na lista atual, a escola possui previsão de <b>{aps_atuais} APs</b>. Porém, após validação da necessidade do ambiente, identificamos que será necessário o <b>{tipo.lower()} para {novos_aps} APs</b>, a fim de atender de forma adequada a cobertura e o funcionamento da unidade escolar.</p>
                {html_adicionais}
                <p>Dessa forma, solicitamos a autorização formal da EACE para atualização da quantidade de equipamentos desta escola, passando de <b>{aps_atuais} AP</b> para <b>{novos_aps} APs</b>.</p>
                <p>Ficamos no aguardo da aprovação para prosseguir com o atendimento em campo.</p>
                <br>
                <p>Atenciosamente,</p>
                <p><b>NOC - Núcleo de Operações de Rede</b></p>
            </body></html>"""
            
            msg.add_alternative(html_body, subtype='html')
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as f: 
                    msg.add_attachment(f.read(), maintype='application', subtype='pdf', filename=pdf_name_drive)
                    
            gmail_service.users().messages().send(userId='me', body={'raw': base64.urlsafe_b64encode(msg.as_bytes()).decode()}).execute()
            log_message("SUCCESS", f"E-mail de pleito ({tipo}) enviado com sucesso para: {config['EMAIL_TO']}")
            return f"✅ E-MAIL ENVIADO COM SUCESSO | {detalhe_busca}"
        except Exception as e: 
            log_message("ERROR", f"Falha ao enviar e-mail: {e}")
            return f"❌ ERRO ENVIO EMAIL: {str(e)}"

    def gerar_documento(self):
        if self.df_principal is None: 
            messagebox.showwarning("Aviso", "Base de dados offline ou sem dados locais.")
            return
            
        inep = self.ent_inep.get().strip()
        novos_aps = self.ent_novos_aps.get().strip()
        
        if not inep or not novos_aps: 
            messagebox.showwarning("Campos vazios", "Informe o Código INEP e a nova quantidade de APs.", parent=self.frame)
            return
            
        self.btn_gerar.config(text="PROCESSANDO...", state="disabled", bg="#9ca3af")
        self.lbl_status.config(text="Iniciando processamento...")
        self.frame.update()
            
        # --- TRAVA ANTI-DUPLICIDADE ---
        if self.var_enviar_email.get():
            if self.verificar_duplicidade_gmail(inep):
                return
        # ------------------------------
            
        col = "Código INEP" if "Código INEP" in self.df_principal.columns else self.df_principal.columns[0]
        res = self.df_principal[self.df_principal[col] == inep]
        if res.empty: 
            messagebox.showerror("Erro", f"Código INEP '{inep}' não encontrado na base de dados.")
            return
            
        dados_escola = res.iloc[0]
        nome_escola = str(dados_escola.get('Nome da Escola', '')).strip()
        aps_atuais = str(dados_escola.get('Kit Wi-Fi', '0'))
        es_infantil = "SIM" if inep in self.infantil_ineps else "NÃO"
        
        try: 
            tipo = "UPGRADE" if int(novos_aps) > int(aps_atuais) else "DOWNGRADE"
        except: 
            tipo = "UPGRADE"
        log_message("INFO", f"Iniciando geracao de pleito {tipo} para escola: {nome_escola} (INEP: {inep})")
        
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        
        doc.add_paragraph("Prezados,")
        doc.add_paragraph()
        doc.add_paragraph(f"Gostaria de solicitar a autorização da EACE para realização de {tipo.lower()} na quantidade de Access Points da escola abaixo:")
        doc.add_paragraph(f"• Escola: {nome_escola}", style='List Bullet')
        doc.add_paragraph(f"• Código INEP / Identificação: {inep}", style='List Bullet')
        doc.add_paragraph(f"• Município: {self.get_field_safely(dados_escola, 'Município')} - {self.get_field_safely(dados_escola, 'UF')}", style='List Bullet')
        doc.add_paragraph(f"• Endereço: {self.get_field_safely(dados_escola, 'Endereço')}", style='List Bullet')
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=12)
        table.style = 'Table Grid'
        headers = ["FASE", "INEP", "ESCOLA", "ESTADO", "CIDADE", "KIT PREVISTO", "PLEITO", "KIT SUGERIDO", "LATITUDE", "LONGITUDE", "ESCOLA INFANTIL", "JUSTIFICATIVA"]
        for i, h in enumerate(headers): 
            table.cell(0, i).text = h
            
        cells = table.rows[1].cells
        vals = ["5", inep, nome_escola, self.get_field_safely(dados_escola, 'UF'), self.get_field_safely(dados_escola, 'Município'), aps_atuais, tipo, novos_aps, self.get_field_safely(dados_escola, 'Latitude'), self.get_field_safely(dados_escola, 'Longitude'), es_infantil, "Adequação à necessidade da escola"]
        for i, v in enumerate(vals): 
            cells[i].text = str(v)
            
        doc.add_paragraph()
        doc.add_paragraph(f"Conforme consta na lista atual, a escola possui previsão de {aps_atuais} APs. Porém, após validação da necessidade do ambiente, identificamos que será necessário {tipo.lower()} para {novos_aps} APs, a fim de atender de forma adequada a cobertura e o funcionamento da unidade.")
        doc.add_paragraph()
        doc.add_paragraph(f"Dessa forma, solicitamos a autorização da EACE para atualização da quantidade de equipamentos desta escola, passando de {aps_atuais} AP para {novos_aps} APs.")
        doc.add_paragraph()
        doc.add_paragraph("Ficamos no aguardo da aprovação para seguir com o atendimento.")
        doc.add_paragraph()
        doc.add_paragraph("Atenciosamente,")
        doc.add_paragraph("NOC - Núcleo de Operações de Rede")
        
        nome_limpo = "".join([c for c in nome_escola if c.isalnum() or c in (' ', '-', '_')]).rstrip()
        nome_arquivo = f"{inep}-{nome_limpo}.docx"
        
        temp_dir = os.path.join(get_base_path(), "temp_docx")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, nome_arquivo)
        
        try:
            doc.save(temp_path)
            save_status = []
            
            config = load_config()
            network_dir = config["NETWORK_PATH"]
            try:
                os.makedirs(network_dir, exist_ok=True)
                shutil.copy2(temp_path, os.path.join(network_dir, nome_arquivo))
                save_status.append("✅ REDE LOCAL: Salvo com sucesso.")
                log_message("INFO", f"Arquivo DOCX copiado para rede local: {network_dir}")
            except Exception as e_red: 
                save_status.append("❌ REDE LOCAL: Falha (Pasta temporariamente inacessivel).")
                log_message("WARNING", f"Falha ao salvar DOCX na pasta de rede ({network_dir}): {e_red}")
                
            ok, msg_drive = self.upload_to_drive(temp_path, nome_arquivo)
            if ok:
                save_status.append("✅ GOOGLE DRIVE: Upload realizado com sucesso.")
                log_message("INFO", f"Upload de DOCX no Google Drive: {msg_drive}")
            else:
                save_status.append(f"❌ GOOGLE DRIVE: Falha ({msg_drive})")
                log_message("ERROR", f"Falha no upload do DOCX para o Drive: {msg_drive}")
                
            try:
                os.makedirs(self.local_output_dir, exist_ok=True)
                shutil.copy2(temp_path, os.path.join(self.local_output_dir, nome_arquivo))
                save_status.append("✅ LOCAL (Documentos): Salvo com sucesso.")
                log_message("INFO", f"Arquivo DOCX copiado para pasta local: {self.local_output_dir}")
            except Exception as e_loc: 
                save_status.append("❌ LOCAL: Falha ao copiar arquivo.")
                log_message("ERROR", f"Falha ao salvar DOCX na pasta local: {e_loc}")
                
            self.atualizar_relatorio({
                "INEP": inep, 
                "Escola": nome_escola, 
                "Status": "PENDENTE", 
                "Tipo": tipo, 
                "Kit Previsto": aps_atuais, 
                "Kit Sugerido": novos_aps, 
                "Observação": "ENCAMINHADO", 
                "OBS EACE": "PENDENTE"
            })
            
            pdf_local = getattr(self, 'selected_pdf_path', '')
            
            if self.var_enviar_email.get():
                status_email = self.enviar_email_automatico(
                    inep, nome_escola, aps_atuais, novos_aps, tipo, es_infantil, dados_escola, temp_dir,
                    pdf_local_path=pdf_local if (pdf_local and os.path.exists(pdf_local)) else None
                )
                save_status.append(status_email)
                
            messagebox.showinfo("Status da Automação", "\n".join(save_status))
            
            self.ent_inep.delete(0, tk.END)
            self.ent_novos_aps.delete(0, tk.END)
            self.limpar_pdf_local()
            
        except Exception as e: 
            show_detailed_error("Erro de Automação", "Ocorreu uma falha critica na automacao de pleitos DOCX.", e)
        finally:
            self.btn_gerar.config(text="🚀 INICIAR AUTOMAÇÃO DOCX/EMAIL", state="normal", bg="#10b981")
            self.lbl_status.config(text="Pronto para automação.", fg="#94a3b8")
            if os.path.exists(temp_path): 
                os.remove(temp_path)
            try:
                for f in os.listdir(temp_dir):
                    fp = os.path.join(temp_dir, f)
                    if os.path.isfile(fp): 
                        os.remove(fp)
            except Exception as e_clean:
                log_message("WARNING", f"Erro ao limpar pasta temporaria: {e_clean}")

# --- CLASSE UNIFICADOR DE EVIDÊNCIAS ---
class UnificadorApp:
    def __init__(self, parent_frame, main_app_callback):
        self.parent_frame = parent_frame
        self.main_app_callback = main_app_callback
        self.frame = tk.Frame(parent_frame, bg="#121214")
        self.frame.pack(fill=tk.BOTH, expand=True)
        
        header_frame = tk.Frame(self.frame, bg="#121214")
        header_frame.pack(fill=tk.X, pady=(15, 10))
        
        tk.Label(header_frame, text="🖼️ UNIFICADOR DE EVIDÊNCIAS", font=("Arial", 16, "bold"), fg="#e2e8f0", bg="#121214").pack(side=tk.LEFT, padx=15)
        
        btn_voltar = tk.Button(
            header_frame, 
            text="⬅️ VOLTAR AO MENU", 
            command=self.main_app_callback, 
            bg="#4b5563", 
            fg="white", 
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=10,
            pady=5
        )
        btn_voltar.pack(side=tk.RIGHT, padx=15)
        btn_voltar.bind("<Enter>", lambda e: btn_voltar.config(bg="#374151"))
        btn_voltar.bind("<Leave>", lambda e: btn_voltar.config(bg="#4b5563"))
        
        self.config = load_config()
        self.default_save_path = self.config.get('unificador_save_path', os.path.join(os.path.expanduser("~"), "Documents"))
        
        self.selected_files = []
        
        mf = tk.Frame(self.frame, bg="#202024", padx=20, pady=15)
        mf.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        list_container = tk.Frame(mf, bg="#202024")
        list_container.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        scrollbar_y = ttk.Scrollbar(list_container)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.lb_files = tk.Listbox(
            list_container, 
            font=("Arial", 10), 
            bg="#1e1e24", 
            fg="#f4f4f5", 
            selectbackground="#3b82f6", 
            selectforeground="white",
            relief="flat",
            bd=3,
            yscrollcommand=scrollbar_y.set
        )
        self.lb_files.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_y.config(command=self.lb_files.yview)
        
        control_pane = tk.Frame(list_container, bg="#202024", padx=10)
        control_pane.pack(side=tk.RIGHT, fill=tk.Y)
        
        btn_add = tk.Button(
            control_pane, 
            text="➕ Adicionar...", 
            command=self.select_files, 
            bg="#3b82f6", 
            fg="white", 
            font=("Arial", 9, "bold"),
            width=14,
            relief="flat",
            cursor="hand2"
        )
        btn_add.pack(pady=5)
        btn_add.bind("<Enter>", lambda e: btn_add.config(bg="#2563eb"))
        btn_add.bind("<Leave>", lambda e: btn_add.config(bg="#3b82f6"))
        
        btn_remove = tk.Button(
            control_pane, 
            text="❌ Excluir", 
            command=self.remove_selected, 
            bg="#ef4444", 
            fg="white", 
            font=("Arial", 9, "bold"),
            width=14,
            relief="flat",
            cursor="hand2"
        )
        btn_remove.pack(pady=5)
        btn_remove.bind("<Enter>", lambda e: btn_remove.config(bg="#dc2626"))
        btn_remove.bind("<Leave>", lambda e: btn_remove.config(bg="#ef4444"))
        
        tk.Frame(control_pane, height=2, bg="#4b5563").pack(fill=tk.X, pady=10)
        
        btn_up = tk.Button(
            control_pane, 
            text="▲ Subir", 
            command=self.move_up, 
            bg="#4b5563", 
            fg="white", 
            font=("Arial", 9, "bold"),
            width=14,
            relief="flat",
            cursor="hand2"
        )
        btn_up.pack(pady=5)
        btn_up.bind("<Enter>", lambda e: btn_up.config(bg="#374151"))
        btn_up.bind("<Leave>", lambda e: btn_up.config(bg="#4b5563"))
        
        btn_down = tk.Button(
            control_pane, 
            text="▼ Descer", 
            command=self.move_down, 
            bg="#4b5563", 
            fg="white", 
            font=("Arial", 9, "bold"),
            width=14,
            relief="flat",
            cursor="hand2"
        )
        btn_down.pack(pady=5)
        btn_down.bind("<Enter>", lambda e: btn_down.config(bg="#374151"))
        btn_down.bind("<Leave>", lambda e: btn_down.config(bg="#4b5563"))
        
        path_frame = tk.Frame(mf, bg="#202024")
        path_frame.pack(fill=tk.X, pady=8)
        
        tk.Label(path_frame, text="Pasta de Destino:", font=("Arial", 10, "bold"), fg="#94a3b8", bg="#202024").pack(side=tk.LEFT)
        self.path_label = tk.Label(path_frame, text=self.default_save_path, font=("Arial", 9), fg="#3b82f6", bg="#202024", anchor="w")
        self.path_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=8)
        
        btn_alterar_caminho = tk.Button(
            path_frame, 
            text="Alterar...", 
            command=self.change_save_path, 
            bg="#4b5563", 
            fg="white",
            font=("Arial", 8, "bold"),
            relief="flat",
            cursor="hand2",
            padx=10
        )
        btn_alterar_caminho.pack(side=tk.RIGHT)
        btn_alterar_caminho.bind("<Enter>", lambda e: btn_alterar_caminho.config(bg="#374151"))
        btn_alterar_caminho.bind("<Leave>", lambda e: btn_alterar_caminho.config(bg="#4b5563"))
        
        name_frame = tk.Frame(mf, bg="#202024")
        name_frame.pack(fill=tk.X, pady=8)
        
        tk.Label(name_frame, text="Nome do PDF Final (sem .pdf):", font=("Arial", 10, "bold"), fg="#e2e8f0", bg="#202024").pack(side=tk.LEFT, padx=(0, 10))
        self.ent_pdf_name = tk.Entry(name_frame, font=("Arial", 11), bg="#1e1e24", fg="#f4f4f5", insertbackground="white", relief="flat", bd=2)
        self.ent_pdf_name.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=3)
        
        btn_gerar_pdf = tk.Button(
            mf, 
            text="⚡ GERAR PDF DE EVIDÊNCIAS COMPRIMIDO", 
            bg="#10b981", 
            fg="white", 
            font=("Arial", 12, "bold"), 
            height=2,
            relief="flat",
            cursor="hand2",
            command=self.process_files
        )
        btn_gerar_pdf.pack(fill=tk.X, pady=(15, 5))
        btn_gerar_pdf.bind("<Enter>", lambda e: btn_gerar_pdf.config(bg="#059669"))
        btn_gerar_pdf.bind("<Leave>", lambda e: btn_gerar_pdf.config(bg="#10b981"))

    def select_files(self):
        files_selected = filedialog.askopenfilenames(filetypes=[("Arquivos Suportados", "*.png *.jpg *.jpeg *.webp *.heic *.txt")])
        for file in files_selected:
            if file not in self.selected_files:
                self.selected_files.append(file)
                self.lb_files.insert(tk.END, os.path.basename(file))
        log_message("INFO", f"Arquivos adicionados para unificacao: {[os.path.basename(f) for f in files_selected]}")

    def remove_selected(self):
        selection = self.lb_files.curselection()
        if not selection:
            return
        idx = selection[0]
        file_removed = self.selected_files.pop(idx)
        self.lb_files.delete(idx)
        log_message("INFO", f"Arquivo removido da lista: {os.path.basename(file_removed)}")
        
        if self.lb_files.size() > 0:
            new_idx = min(idx, self.lb_files.size() - 1)
            self.lb_files.select_set(new_idx)
            self.lb_files.activate(new_idx)

    def move_up(self):
        selection = self.lb_files.curselection()
        if not selection:
            return
        idx = selection[0]
        if idx == 0:
            return
            
        self.selected_files[idx], self.selected_files[idx-1] = self.selected_files[idx-1], self.selected_files[idx]
        
        name = self.lb_files.get(idx)
        self.lb_files.delete(idx)
        self.lb_files.insert(idx - 1, name)
        
        self.lb_files.select_set(idx - 1)
        self.lb_files.activate(idx - 1)

    def move_down(self):
        selection = self.lb_files.curselection()
        if not selection:
            return
        idx = selection[0]
        if idx == self.lb_files.size() - 1:
            return
            
        self.selected_files[idx], self.selected_files[idx+1] = self.selected_files[idx+1], self.selected_files[idx]
        
        name = self.lb_files.get(idx)
        self.lb_files.delete(idx)
        self.lb_files.insert(idx + 1, name)
        
        self.lb_files.select_set(idx + 1)
        self.lb_files.activate(idx + 1)

    def change_save_path(self):
        new_path = filedialog.askdirectory(initialdir=self.default_save_path)
        if new_path:
            self.default_save_path = new_path
            self.path_label.config(text=self.default_save_path)
            self.config['unificador_save_path'] = new_path
            save_config(self.config)
            log_message("INFO", f"Pasta destino alterada para: {new_path}")
            messagebox.showinfo("Caminho Salvo", f"Nova pasta de destino padrao definida:\n{new_path}")

    def process_files(self):
        pdf_name = self.ent_pdf_name.get().strip()
        if not self.selected_files or not pdf_name:
            messagebox.showwarning("Aviso", "Adicione arquivos e defina o nome do PDF de saida.")
            return
            
        log_message("INFO", f"Compilando PDF '{pdf_name}.pdf' com {len(self.selected_files)} arquivos...")
        
        try:
            image_list = []
            a4_size = (1240, 1754)
            margin = 80
            font_size = 24
            max_resolution = 1920
            
            try:
                font = ImageFont.truetype("arial.ttf", font_size)
            except IOError:
                font = ImageFont.load_default()

            for file_path in self.selected_files:
                ext = os.path.splitext(file_path)[1].lower()
                
                if ext == ".txt":
                    with open(file_path, 'r', encoding='utf-8') as f_txt:
                        content = f_txt.read()
                    
                    lines = []
                    for line in content.splitlines():
                        lines.extend(textwrap.wrap(line, width=80))
                    
                    lines_per_page = (a4_size[1] - 2 * margin) // (font_size + 10)
                    
                    for i in range(0, len(lines), lines_per_page):
                        page_lines = lines[i:i + lines_per_page]
                        img = Image.new('RGB', a4_size, color=(255, 255, 255))
                        draw = ImageDraw.Draw(img)
                        y = margin
                        
                        for line in page_lines:
                            draw.text((margin, y), line, font=font, fill=(0, 0, 0))
                            y += font_size + 10
                            
                        image_list.append(img)
                else:
                    img_original = Image.open(file_path)
                    img_oriented = ImageOps.exif_transpose(img_original).convert('RGB')
                    
                    w, h = img_oriented.size
                    if w > max_resolution or h > max_resolution:
                        if w > h:
                            new_w = max_resolution
                            new_h = int(h * (max_resolution / w))
                        else:
                            new_h = max_resolution
                            new_w = int(w * (max_resolution / h))
                        img_oriented = img_oriented.resize((new_w, new_h), Image.Resampling.LANCZOS)
                        log_message("INFO", f"Imagem {os.path.basename(file_path)} otimizada de {w}x{h} para {new_w}x{new_h}")
                    
                    image_list.append(img_oriented)
            
            if image_list:
                output_pdf_path = os.path.join(self.default_save_path, f"{pdf_name}.pdf")
                
                image_list[0].save(
                    output_pdf_path, 
                    "PDF", 
                    save_all=True, 
                    append_images=image_list[1:],
                    quality=80,
                    optimize=True
                )
                
                tamanho_bytes = os.path.getsize(output_pdf_path)
                tamanho_mb = tamanho_bytes / (1024 * 1024)
                
                log_message("SUCCESS", f"PDF criado: {output_pdf_path} (Tamanho: {tamanho_mb:.2f} MB)")
                messagebox.showinfo("Sucesso", f"PDF de evidencias criado com sucesso!\n\nSalvo em:\n{output_pdf_path}\n\nTamanho final: {tamanho_mb:.2f} MB (Otimizado)")
                
                self.selected_files = []
                self.lb_files.delete(0, tk.END)
                self.ent_pdf_name.delete(0, tk.END)
            else:
                messagebox.showwarning("Aviso", "Nenhum arquivo valido processado.")
                
        except Exception as e:
            show_detailed_error("Erro no Unificador", "Nao foi possivel unificar os arquivos de evidencias.", e)

# --- CLASSE CENTRAL MAINAPP ---
class MainApp:
    def __init__(self, root):
        self.root = root
        self.root.title("NOC - Ferramentas de Suporte EACE v2.5")
        self.root.geometry("850x700")
        self.root.configure(bg="#121214")
        
        self.current_frame = None
        self.show_main_menu()
        self.iniciar_autenticacao_background()

    def clear_frame(self):
        if self.current_frame: 
            self.current_frame.destroy()

    def iniciar_autenticacao_background(self):
        """Inicia thread em segundo plano para autenticacao."""
        log_message("INFO", "Disparando thread de conexao Google APIs...")
        threading.Thread(target=thread_autenticar_google, args=(self.atualizar_status_conexao,), daemon=True).start()

    def atualizar_status_conexao(self):
        """Atualiza a UI com o status da conexao Google (thread-safe)."""
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            status = GOOGLE_CONNECTION["status"]
            err_msg = GOOGLE_CONNECTION["error_message"]
            
        self.root.after(0, lambda: self._atualizar_visual_status(status, err_msg))

    def _atualizar_visual_status(self, status, error_message):
        if not hasattr(self, "lbl_conexao") or not self.lbl_conexao.winfo_exists():
            return
            
        if status == "Conectando...":
            self.lbl_conexao.config(text="⏳ Conectando aos servicos Google em background...", fg="#eab308")
        elif status == "Conectado":
            self.lbl_conexao.config(text="🟢 Servicos Google Cloud: ONLINE (Autenticado)", fg="#10b981")
            log_message("INFO", "Status UI atualizado para: Conectado")
        else:
            self.lbl_conexao.config(text="🔴 Servicos Google offline. Usando modo cache local.", fg="#ef4444")
            log_message("WARNING", f"Status UI atualizado para: Offline ({error_message})")

    def show_main_menu(self):
        self.clear_frame()
        self.current_frame = tk.Frame(self.root, bg="#121214")
        self.current_frame.pack(fill=tk.BOTH, expand=True)
        
        title_frame = tk.Frame(self.current_frame, bg="#121214")
        title_frame.pack(pady=(45, 10))
        
        caminho_logo = os.path.join(get_base_path(), "logo_bitnet.png")
        if os.path.exists(caminho_logo):
            try:
                logo_img = Image.open(caminho_logo)
                logo_img.thumbnail((260, 90), Image.Resampling.LANCZOS)
                self.logo_tk = ImageTk.PhotoImage(logo_img)
                lbl_logo = tk.Label(title_frame, image=self.logo_tk, bg="#121214")
                lbl_logo.pack(pady=(0, 10))
            except Exception as e_logo:
                log_message("WARNING", f"Nao foi possivel carregar o logotipo: {e_logo}")
                
        tk.Label(title_frame, text="PAINEL DE FERRAMENTAS NOC", font=("Arial", 22, "bold"), fg="#e2e8f0", bg="#121214").pack()
        tk.Label(title_frame, text="Automações e Consultas Operacionais EACE", font=("Arial", 11), fg="#94a3b8", bg="#121214").pack(pady=5)
        
        menu_box = tk.Frame(self.current_frame, bg="#202024", padx=35, pady=30)
        menu_box.pack(pady=20, padx=50, fill=tk.X)
        
        btn_consulta = tk.Button(
            menu_box, 
            text="🔍 CONSULTA RÁPIDA DE INEP", 
            command=lambda: self.show_tool_frame(ConsultaEACE), 
            font=("Arial", 13, "bold"), 
            bg="#3b82f6", 
            fg="white", 
            height=2,
            relief="flat",
            cursor="hand2"
        )
        btn_consulta.pack(fill=tk.X, pady=8)
        btn_consulta.bind("<Enter>", lambda e: btn_consulta.config(bg="#2563eb"))
        btn_consulta.bind("<Leave>", lambda e: btn_consulta.config(bg="#3b82f6"))
        
        btn_docx = tk.Button(
            menu_box, 
            text="📝 AUTOMAÇÃO DE PLEITOS (DOCX & EMAIL)", 
            command=lambda: self.show_tool_frame(AutomacaoEACE), 
            font=("Arial", 13, "bold"), 
            bg="#10b981", 
            fg="white", 
            height=2,
            relief="flat",
            cursor="hand2"
        )
        btn_docx.pack(fill=tk.X, pady=8)
        btn_docx.bind("<Enter>", lambda e: btn_docx.config(bg="#059669"))
        btn_docx.bind("<Leave>", lambda e: btn_docx.config(bg="#10b981"))
        
        btn_unificador = tk.Button(
            menu_box, 
            text="🖼️ UNIFICADOR DE EVIDÊNCIAS (GERAR PDF)", 
            command=lambda: self.show_tool_frame(UnificadorApp), 
            font=("Arial", 13, "bold"), 
            bg="#eab308", 
            fg="#121214", 
            height=2,
            relief="flat",
            cursor="hand2"
        )
        btn_unificador.pack(fill=tk.X, pady=8)
        btn_unificador.bind("<Enter>", lambda e: btn_unificador.config(bg="#ca8a04"))
        btn_unificador.bind("<Leave>", lambda e: btn_unificador.config(bg="#eab308"))
        
        footer = tk.Frame(self.current_frame, bg="#121214")
        footer.pack(side=tk.BOTTOM, fill=tk.X, pady=25, padx=25)
        
        self.lbl_conexao = tk.Label(footer, text="⏳ Inicializando conexão com os servidores Google...", font=("Arial", 10, "bold"), fg="#eab308", bg="#121214", anchor="w")
        self.lbl_conexao.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        btn_config = tk.Button(
            footer, 
            text="⚙️ Configurações", 
            command=self.abrir_configuracoes, 
            font=("Arial", 9, "bold"), 
            bg="#4b5563", 
            fg="white", 
            relief="flat",
            cursor="hand2",
            padx=12,
            pady=5
        )
        btn_config.pack(side=tk.RIGHT)
        btn_config.bind("<Enter>", lambda e: btn_config.config(bg="#374151"))
        btn_config.bind("<Leave>", lambda e: btn_config.config(bg="#4b5563"))
        
        global GOOGLE_CONNECTION
        with GOOGLE_CONNECTION["lock"]:
            status = GOOGLE_CONNECTION["status"]
            err_msg = GOOGLE_CONNECTION["error_message"]
        self._atualizar_visual_status(status, err_msg)

    def show_tool_frame(self, ToolClass):
        self.clear_frame()
        self.current_frame = ToolClass(self.root, self.show_main_menu).frame

    def abrir_configuracoes(self):
        ConfiguracoesDialog(self.root)

# --- EXECUÇÃO DO APLICATIVO ---
if __name__ == "__main__":
    multiprocessing.freeze_support()
    
    log_message("INFO", "=============================================")
    log_message("INFO", "Iniciando aplicativo NOC Ferramentas EACE v2.5")
    log_message("INFO", "=============================================")
    
    root = tk.Tk()
    
    # Customizacao visual avançada de widgets do ttk
    style = ttk.Style(root)
    style.theme_use('clam')
    
    style.configure(
        "TScrollbar",
        gripcount=0,
        background="#4b5563",
        darkcolor="#202024",
        lightcolor="#4b5563",
        troughcolor="#1e1e24",
        bordercolor="#202024"
    )
    style.map(
        "TScrollbar",
        background=[('active', '#374151')]
    )
    
    app = MainApp(root)
    root.mainloop()
